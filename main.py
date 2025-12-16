#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Crypto Exam Generator (v1.8c)

Novinky v 1.8c
- Přidána aplikace ikona (icon/icon.png). Pokud existuje, nastaví se na QApplication i hlavní okno.
- Oprava chyby: chybějící metoda MainWindow._import_from_docx (menu „Import z DOCX…“).
- Import DOCX: jemné vylepšení – přenáší informaci o typu číslování (decimal / lowerLetter / upperLetter / bullet).

Pozn.: Word numbering (numbering.xml) je mapováno pouze na výsledné vizuální <ol>/<ul> v HTML, bez úprav numbering.xml
(vizualně věrné, minimální zásah).
"""
from __future__ import annotations

import hashlib
import secrets

import subprocess

import json
import sys
import uuid as _uuid
import re
import os
import html as _html
import zipfile
from xml.etree import ElementTree as ET
from dataclasses import dataclass, asdict, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from html.parser import HTMLParser

from PySide6.QtCore import Qt, QSize, QSaveFile, QByteArray, QTimer, QDateTime, QPoint, QRect, QTime
from PySide6.QtGui import (
    QAction,
    QActionGroup,
    QKeySequence,
    QTextCharFormat,
    QTextCursor,
    QTextListFormat,
    QTextBlockFormat,
    QColor,
    QPalette, QTextDocument,
    QFont, QPen, 
    QPixmap, QImage, QImageReader, QPainter, QIcon, QBrush, QPainterPath
)
from PySide6.QtWidgets import (
    QApplication,
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QTreeWidget,
    QTreeWidgetItem,
    QSplitter,
    QToolBar,
    QTextEdit,
    QFileDialog,
    QMessageBox, QProgressBar,
    QLineEdit,
    QPushButton,
    QFormLayout,
    QSpinBox,
    QDoubleSpinBox,
    QComboBox,
    QColorDialog,
    QAbstractItemView,
    QDialog,
    QDialogButtonBox,
    QLabel,
    QStyle,
    QScrollArea,
    QWizard,
    QWizardPage,
    QDateTimeEdit,
    # Nové importy pro v4.0 UI
    QGroupBox,
    QTableWidget, QListWidget,
    QTableWidgetItem, QFrame, QStyledItemDelegate,
    QHeaderView, QCheckBox, QGridLayout,
    QTreeWidgetItemIterator, QButtonGroup,
    QHeaderView, QMenu, QTabWidget, QRadioButton,
    QTreeWidget, QTreeWidgetItem, QSizePolicy
)

APP_VERSION = "8.1.0"
APP_NAME = f"Správce zkouškových testů (v{APP_VERSION})"

# ---------------------------------------------------------------------------
# Globální pomocné funkce
# ---------------------------------------------------------------------------
def generate_colored_icon(text: str, color: QColor, shape: str = "circle") -> QIcon:
    """Vygeneruje jednoduchou ikonu s textem/symbolem (Global Helper)."""
    pix = QPixmap(16, 16)
    pix.fill(Qt.transparent)
    painter = QPainter(pix)
    painter.setRenderHint(QPainter.Antialiasing)
    
    painter.setBrush(color)
    painter.setPen(Qt.NoPen)
    
    if shape == "circle":
        painter.drawEllipse(1, 1, 14, 14)
    elif shape == "star":
        path = QPainterPath()
        path.moveTo(8, 0)
        path.lineTo(16, 8)
        path.lineTo(8, 16)
        path.lineTo(0, 8)
        path.closeSubpath()
        painter.drawPath(path)
    else:
        painter.drawRect(1, 1, 14, 14)
        
    painter.setPen(QColor("black")) 
    font = painter.font()
    font.setBold(True)
    font.setPointSize(9)
    painter.setFont(font)
    painter.drawText(pix.rect(), Qt.AlignCenter, text)
    painter.end()
    return QIcon(pix)

def parse_html_to_paragraphs(html: str) -> List[dict]:
    if not html:
        return []
    
    parser = HTMLToDocxParser()
    try:
        parser.feed(html)
        parser._end_paragraph() # Flush
    except Exception as e:
        print(f"Chyba při parsování HTML: {e}")
        return [{'align': 'left', 'runs': [{'text': html, 'b': False, 'i': False, 'u': False, 'color': None}], 'prefix': ''}]
    
    # Filtrování prázdných odstavců na začátku a konci (Trim)
    res = parser.paragraphs
    
    # Remove empty start
    while res and not any(r['text'].strip() for r in res[0]['runs']):
        res.pop(0)
    # Remove empty end
    while res and not any(r['text'].strip() for r in res[-1]['runs']):
        res.pop()
        
    if not res:
         # Pokud po promazání nic nezbylo, vrátíme aspoň jeden prázdný (nebo nic, podle logiky)
         # Pro insert_rich chceme raději nic, než prázdný řádek navíc
         return []
         
    return res

# --------------------------- Datové typy ---------------------------

# Přidat do importů (pokud tam není QTableWidget atd., ale v 5.9.2 byly):
# from PySide6.QtWidgets import ..., QTableWidget, QTableWidgetItem, QAbstractItemView

@dataclass
class RootData:
    groups: List[Group]
    trash: List[dict] = field(default_factory=list)

@dataclass
class FunnyAnswer:
    text: str
    author: str
    date: str
    # Nové pole – uložený zdrojový dokument (cesta k souboru, nebo prázdný string)
    source_doc: str = ""

@dataclass
class Question:
    id: str
    type: str  # "classic" nebo "bonus"
    text_html: str
    title: str
    points: int = 1
    bonus_correct: float = 0.0
    bonus_wrong: float = 0.0
    created_at: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    # Nová pole
    correct_answer: str = ""
    funny_answers: List[FunnyAnswer] = field(default_factory=list)
    image_path: str = ""  # cesta k obrázku (volitelné)
    image_width_cm: float = 0.0  # cílová šířka vloženého obrázku v DOCX (cm), 0 = default
    image_height_cm: float = 0.0  # cílová výška vloženého obrázku v DOCX (cm), 0 = default/auto
    image_keep_aspect: bool = True  # pokud True, UI udržuje poměr stran (šířka/výška) při editaci rozměrů

    @staticmethod
    def new_default(q_type: str = "classic") -> "Question":
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        if q_type == "bonus":
            return Question(
                id=str(_uuid.uuid4()),
                type="bonus",
                text_html="<p>Znění bonusové otázky...</p>",
                title="BONUS otázka",
                points=0,
                bonus_correct=1.0,
                bonus_wrong=0.0,
                created_at=now,
                correct_answer="",
                funny_answers=[]
            )
        return Question(
            id=str(_uuid.uuid4()),
            type="classic",
            text_html="<p>Znění otázky...</p>",
            title="Otázka",
            points=1,
            bonus_correct=0.0,
            bonus_wrong=0.0,
            created_at=now,
            correct_answer="",
            funny_answers=[]
        )

@dataclass
class Subgroup:
    id: str
    name: str
    subgroups: List["Subgroup"] = field(default_factory=list)
    questions: List[Question] = field(default_factory=list)

@dataclass
class Group:
    id: str
    name: str
    subgroups: List[Subgroup]

@dataclass
class RootData:
    groups: List[Group]

# --------------------------- Utility: Dark theme ---------------------------

def apply_dark_theme(app: QApplication) -> None:
    QApplication.setStyle("Fusion")
    palette = QPalette()
    palette.setColor(QPalette.Window, QColor(37, 37, 38))
    palette.setColor(QPalette.WindowText, Qt.white)
    palette.setColor(QPalette.Base, QColor(30, 30, 30))
    palette.setColor(QPalette.AlternateBase, QColor(45, 45, 45))
    palette.setColor(QPalette.ToolTipBase, Qt.black)
    palette.setColor(QPalette.ToolTipText, Qt.white)
    palette.setColor(QPalette.Text, Qt.white)
    palette.setColor(QPalette.Button, QColor(45, 45, 48))
    palette.setColor(QPalette.ButtonText, Qt.white)
    palette.setColor(QPalette.BrightText, Qt.red)
    palette.setColor(QPalette.Highlight, QColor(10, 132, 255))
    palette.setColor(QPalette.HighlightedText, Qt.black)
    app.setPalette(palette)


class MultiSourceDialog(QDialog):
    """Dialog se stromem a checkboxy pro výběr více zdrojů s počítadlem otázek."""
    def __init__(self, owner: "MainWindow", selected_data: list) -> None:
        super().__init__(owner)
        self.setWindowTitle("Vyberte zdroje otázek")
        self.resize(600, 700)
        self.owner = owner
        
        # Ukládáme si vybraná ID
        self.selected_ids = {item['id'] for item in selected_data} if selected_data else set()
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(8, 8, 8, 8)
        
        layout.addWidget(QLabel("Zaškrtněte skupiny nebo podskupiny, ze kterých se mají náhodně vybírat otázky.\n(Výběr skupiny automaticky zahrne všechny její podskupiny)"))
        
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Název zdroje"])
        self.tree.header().hide()
        # Signál pro změnu checkboxu
        self.tree.itemChanged.connect(self._on_item_changed)
        layout.addWidget(self.tree)
        
        # Label pro celkový součet
        self.lbl_total = QLabel("Celkem vybráno otázek: 0")
        self.lbl_total.setStyleSheet("font-weight: bold; color: #42a5f5; font-size: 14px; margin-top: 5px;")
        layout.addWidget(self.lbl_total)
        
        # Tlačítka
        bb = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        bb.accepted.connect(self.accept)
        bb.rejected.connect(self.reject)
        layout.addWidget(bb)
        
        # Styly
        self.icon_group = owner._generate_icon("S", QColor("#ff5252"), "rect")
        self.icon_sub = owner._generate_icon("P", QColor("#ff8a80"), "rect")
        self.color_group = QBrush(QColor("#ff5252"))
        self.color_subgroup = QBrush(QColor("#ff8a80"))

        # Naplnění stromu
        self._is_populating = True # Flag aby se nespouštěly signály při plnění
        self._populate_tree(owner.root.groups)
        self._is_populating = False
        
        # Prvotní přepočet
        self._recalculate_total()

    def _get_classic_count(self, node):
        """Rekurzivně spočítá klasické otázky v uzlu."""
        count = 0
        if hasattr(node, "questions"):
            count += len([q for q in node.questions if q.type == "classic"])
        if hasattr(node, "subgroups") and node.subgroups:
            for sub in node.subgroups:
                count += self._get_classic_count(sub)
        return count

    def _populate_tree(self, groups):
        def add_sub_recursive(parent_item, subs):
            for sg in subs:
                q_count = self._get_classic_count(sg)
                
                # Zobrazíme jen pokud má nějaké klasické otázky (přímo nebo v dětech)
                if q_count > 0:
                    display_text = f"{sg.name} ({q_count})"
                    
                    item = QTreeWidgetItem([display_text])
                    item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
                    item.setFlags(item.flags() & ~Qt.ItemIsAutoTristate)
                    
                    check_state = Qt.Checked if sg.id in self.selected_ids else Qt.Unchecked
                    item.setCheckState(0, check_state)
                    
                    item.setIcon(0, self.icon_sub)
                    item.setForeground(0, self.color_subgroup)
                    item.setData(0, Qt.UserRole, {"id": sg.id, "type": "subgroup", "count": q_count})
                    
                    parent_item.addChild(item)
                    
                    if sg.subgroups:
                        add_sub_recursive(item, sg.subgroups)
                        item.setExpanded(True)

        for g in groups:
            q_count = self._get_classic_count(g)
            
            # Zobrazíme jen pokud má > 0 klasických otázek
            if q_count > 0:
                display_text = f"{g.name} ({q_count})"
                
                item = QTreeWidgetItem([display_text])
                item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
                item.setFlags(item.flags() & ~Qt.ItemIsAutoTristate)
                
                check_state = Qt.Checked if g.id in self.selected_ids else Qt.Unchecked
                item.setCheckState(0, check_state)
                
                item.setIcon(0, self.icon_group)
                item.setForeground(0, self.color_group)
                f = item.font(0); f.setBold(True); item.setFont(0, f)
                item.setData(0, Qt.UserRole, {"id": g.id, "type": "group", "count": q_count})
                
                self.tree.addTopLevelItem(item)
                add_sub_recursive(item, g.subgroups)
                item.setExpanded(True)

    def _on_item_changed(self, item, column):
        """Logika pro kaskádové zaškrtávání."""
        if self._is_populating: return
        
        state = item.checkState(0)
        self._is_populating = True # Blokujeme rekurzivní volání signálů
        
        # 1. Aplikovat na všechny potomky (směr dolů)
        self._set_check_state_recursive(item, state)
        
        # 2. Aktualizovat rodiče (směr nahoru - pokud jsou všechny děti checked, rodič checked)
        # (Tohle je volitelné, někdy je lepší nechat rodiče nezávislého, 
        # ale pro "výběr skupiny vybere podskupiny" je směr dolů klíčový).
        
        self._is_populating = False
        self._recalculate_total()

    def _set_check_state_recursive(self, item, state):
        for i in range(item.childCount()):
            child = item.child(i)
            child.setCheckState(0, state)
            self._set_check_state_recursive(child, state)

    def _recalculate_total(self):
        """Spočítá unikátní otázky ve vybraných uzlech."""
        # Musíme být opatrní, abychom nepočítali duplicitně.
        # Nejjednodušší: Pokud je zaškrtnutý rodič, započítáme jeho 'count'.
        # Ale pozor: 'count' u rodiče už obsahuje děti.
        # Takže: Pokud je rodič zaškrtnutý, přičteme ho a děti už ignorujeme (protože jsou v jeho součtu).
        # Pokud rodič není, koukneme na děti.
        
        total = 0
        
        # Projdeme top-level items
        for i in range(self.tree.topLevelItemCount()):
            item = self.tree.topLevelItem(i)
            total += self._count_checked_item(item)
            
        self.lbl_total.setText(f"Celkem vybráno otázek: {total}")

    def _count_checked_item(self, item):
        # Pokud je tento item zaškrtnutý, vezmeme jeho 'count' (který zahrnuje vše pod ním)
        # a už neprocházíme děti (abychom je nepřičetli znovu).
        if item.checkState(0) == Qt.Checked:
            data = item.data(0, Qt.UserRole)
            return data.get("count", 0)
        
        # Pokud není zaškrtnutý, musíme se podívat do dětí, jestli tam něco není
        sub_total = 0
        for i in range(item.childCount()):
            child = item.child(i)
            sub_total += self._count_checked_item(child)
        return sub_total

    def get_selected_items(self) -> list:
        """Vrátí seznam slovníků pro ExportWizard.
        Vracíme jen ty položky, které jsou checked.
        Pokud je checked rodič, vracíme i jeho ID (což v logice wizardu znamená 'vše pod ním').
        Víme, že wizard si pak umí vybrat unikátní otázky.
        """
        res = []
        it = QTreeWidgetItemIterator(self.tree)
        while it.value():
            item = it.value()
            if item.checkState(0) == Qt.Checked:
                data = item.data(0, Qt.UserRole)
                if data:
                    res.append({"id": data["id"], "type": data["type"], "name": item.text(0), "count": data["count"]})
            it += 1
        return res
    
    def get_total_selected_count(self):
        # Rychlý hack - text labelu parse
        txt = self.lbl_total.text()
        try:
            return int(txt.split(": ")[1])
        except:
            return 0


# ---------------------- Custom Delegate pro Combo (Export Wizard) ----------------------
class ComboGroupDelegate(QStyledItemDelegate):
    """Delegate pro obarvení textu položek v ComboBoxu podle typu (Skupina/Podskupina)."""
    def paint(self, painter, option, index):
        # Získání dat položky
        data = index.data(Qt.UserRole)
        
        # Určení barvy podle typu
        if data and isinstance(data, dict):
            item_type = data.get("type")
            if item_type == "group":
                text_color = QColor("#ff5252")
            elif item_type == "subgroup":
                text_color = QColor("#ff8a80")
            else:
                text_color = option.palette.color(QPalette.Text)
        else:
            text_color = option.palette.color(QPalette.Text)
        
        # Kreslení pozadí (výběr/hover)
        if option.state & QStyle.State_Selected:
            painter.fillRect(option.rect, option.palette.highlight())
        else:
            # Volitelné: lehce odlišit pozadí řádku
            pass
        
        # Kreslení ikony (pokud existuje)
        icon = index.data(Qt.DecorationRole)
        if icon:
            # Standardní velikost ikony
            icon_size = option.decorationSize
            # Vycentrování ikony vertikálně
            icon_y = option.rect.top() + (option.rect.height() - icon_size.height()) // 2
            icon_rect = QRect(option.rect.left() + 4, icon_y, icon_size.width(), icon_size.height())
            icon.paint(painter, icon_rect)
        
        # Kreslení textu s barvou
        text = index.data(Qt.DisplayRole)
        if text:
            # Odsazení za ikonu (cca 28px)
            text_rect = option.rect.adjusted(28, 0, -4, 0)
            painter.setPen(text_color)
            
            # Tučné písmo
            font = option.font
            font.setBold(True)
            painter.setFont(font)
            
            painter.drawText(text_rect, Qt.AlignLeft | Qt.AlignVCenter, text)


# --------------------------- DnD Tree ---------------------------

class DnDTree(QTreeWidget):
    """QTreeWidget s podporou drag&drop, po přesunu synchronizuje datový model."""

    def __init__(self, owner: "MainWindow") -> None:
        super().__init__()
        self.owner = owner
        self.setHeaderLabels(["Název", "Typ / body"])
        
        # Nastavení chování hlavičky pro správné roztažení
        header = self.header()
        # 0. sloupec (Název) se roztáhne do zbytku
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        # 1. sloupec (Typ/body) se přizpůsobí obsahu
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        
        # Důležité: StretchLastSection musí být False, jinak přebije naše nastavení
        header.setStretchLastSection(False)

        self.setContextMenuPolicy(Qt.CustomContextMenu)
        self.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.setDragEnabled(True)
        self.setAcceptDrops(True)
        self.setDropIndicatorShown(True)
        self.setDefaultDropAction(Qt.MoveAction)
        self.setDragDropMode(QAbstractItemView.InternalMove)

    def dropEvent(self, event) -> None:
        ids_before = self.owner._selected_question_ids()

        # Pokud uživatel pustí drag mimo položky (na prázdnou plochu / mimo seznam),
        # upozorníme, že tím položku "vyhodí" ze seznamu. Po potvrzení ji opravdu smažeme.
        try:
            pos = event.position().toPoint()  # Qt6
        except Exception:
            pos = event.pos()  # fallback

        target = self.itemAt(pos)
        if target is None and self.dropIndicatorPosition() == QAbstractItemView.OnViewport:
            mb = QMessageBox(self)
            mb.setIcon(QMessageBox.Warning)
            mb.setWindowTitle("Pozor")
            mb.setText("Přesouváš položku mimo seznam.")
            mb.setInformativeText("Tím se položka odstraní ze seznamu. Chceš pokračovat?")

            mb.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
            mb.setDefaultButton(QMessageBox.No)

            # Tlačítka česky
            yes_btn = mb.button(QMessageBox.Yes)
            no_btn = mb.button(QMessageBox.No)
            if yes_btn is not None:
                yes_btn.setText("Ano")
            if no_btn is not None:
                no_btn.setText("Ne")

            res = mb.exec()
            if res != QMessageBox.Yes:
                event.ignore()
                self.owner.statusBar().showMessage("Přesun zrušen.", 3000)
                return

            # Uživatel potvrdil -> vybrané položky opravdu smažeme ze stromu (včetně skupin).
            selected = list(self.selectedItems())
            selected_set = set(selected)

            def has_selected_ancestor(it) -> bool:
                p = it.parent()
                while p is not None:
                    if p in selected_set:
                        return True
                    p = p.parent()
                return False

            def depth(it) -> int:
                d = 0
                p = it.parent()
                while p is not None:
                    d += 1
                    p = p.parent()
                return d

            # Neodstraňovat položky, které už budou odstraněny s rodičem
            selected = [it for it in selected if not has_selected_ancestor(it)]
            # Odstraňovat od nejhlubších (bezpečnější)
            selected.sort(key=depth, reverse=True)

            for it in selected:
                parent = it.parent()
                if parent is not None:
                    idx = parent.indexOfChild(it)
                    if idx >= 0:
                        parent.takeChild(idx)
                else:
                    idx = self.indexOfTopLevelItem(it)
                    if idx >= 0:
                        self.takeTopLevelItem(idx)

            self.owner._sync_model_from_tree()
            self.owner._refresh_tree()
            self.owner.save_data()
            self.owner.statusBar().showMessage("Položka odstraněna ze seznamu (uloženo).", 3000)
            return

        # Necháme QTreeWidget provést interní DnD
        super().dropEvent(event)

        def kind_of(it) -> str:
            meta = it.data(0, Qt.UserRole) or {}
            if isinstance(meta, dict):
                return meta.get("kind") or ""
            return ""

        def iter_all_items():
            stack = []
            for ti in range(self.topLevelItemCount()):
                stack.append(self.topLevelItem(ti))
            while stack:
                it = stack.pop()
                yield it
                for ci in range(it.childCount() - 1, -1, -1):
                    stack.append(it.child(ci))

        def top_level_group_of(it):
            cur = it
            while cur is not None:
                if kind_of(cur) == "group" and cur.parent() is None:
                    return cur
                cur = cur.parent()
            return None

        def nearest_group_for_top_level_index(idx: int):
            # najdi nejbližší skupinu nad idx, jinak pod idx
            for j in range(idx - 1, -1, -1):
                it = self.topLevelItem(j)
                if it is not None and kind_of(it) == "group":
                    return it, j
            for j in range(idx + 1, self.topLevelItemCount()):
                it = self.topLevelItem(j)
                if it is not None and kind_of(it) == "group":
                    return it, j
            return None, -1

        changed = True
        while changed:
            changed = False

            # 1) NIC nesmí být dítě "question" (cokoliv puštěné na otázku).
            #    Nové chování:
            #    - subgroup puštěná na otázku => vnořit do nadřazené podskupiny (tj. podskupiny, která obsahuje cílovou otázku)
            #    - question puštěná na otázku => zůstane sourozencem (pod cílovou otázkou) ve stejné podskupině
            for it in iter_all_items():
                if kind_of(it) != "question":
                    continue
                if it.childCount() <= 0:
                    continue

                moved = it.takeChild(0)
                if moved is None:
                    continue

                ck = kind_of(moved)
                q_parent = it.parent()  # typicky subgroup

                if ck == "question":
                    if q_parent is not None:
                        insert_at = q_parent.indexOfChild(it) + 1
                        q_parent.insertChild(insert_at, moved)
                    else:
                        self.insertTopLevelItem(self.topLevelItemCount(), moved)

                elif ck == "subgroup":
                    # PODSKUPINA NA OTÁZKU => vnořit do nadřazené podskupiny (q_parent)
                    if q_parent is not None:
                        q_parent.insertChild(q_parent.childCount(), moved)
                    else:
                        # fallback (nemáme nadřazenou podskupinu) - dáme na top-level
                        self.insertTopLevelItem(self.topLevelItemCount(), moved)

                elif ck == "group":
                    # skupina musí být top-level
                    tg = top_level_group_of(it)
                    insert_at = self.indexOfTopLevelItem(tg) + 1 if tg is not None else self.topLevelItemCount()
                    self.insertTopLevelItem(insert_at, moved)

                else:
                    if q_parent is not None:
                        insert_at = q_parent.indexOfChild(it) + 1
                        q_parent.insertChild(insert_at, moved)
                    else:
                        self.insertTopLevelItem(self.topLevelItemCount(), moved)

                changed = True
                break

            if changed:
                continue

            # 2) Top-level podskupina -> přesun do nejbližší skupiny (aby nemizela)
            for ti in range(self.topLevelItemCount()):
                top = self.topLevelItem(ti)
                if top is None:
                    continue
                if kind_of(top) != "subgroup":
                    continue

                grp, grp_idx = nearest_group_for_top_level_index(ti)
                if grp is None:
                    continue

                moved = self.takeTopLevelItem(ti)
                if moved is None:
                    continue

                if grp_idx >= 0 and ti == grp_idx + 1:
                    grp.insertChild(0, moved)
                else:
                    grp.insertChild(grp.childCount(), moved)

                changed = True
                break

            if changed:
                continue

            # 3) Skupina nesmí být nikdy vnořená (group musí být top-level)
            for it in iter_all_items():
                if kind_of(it) != "group":
                    continue
                if it.parent() is None:
                    continue

                parent = it.parent()
                idx = parent.indexOfChild(it)
                if idx < 0:
                    continue
                moved = parent.takeChild(idx)
                if moved is None:
                    continue

                self.insertTopLevelItem(self.topLevelItemCount(), moved)
                changed = True
                break

        self.owner._sync_model_from_tree()
        self.owner._refresh_tree()
        self.owner._reselect_questions(ids_before)
        self.owner.save_data()
        self.owner.statusBar().showMessage("Přesun dokončen (uloženo).", 3000)

class BonusQuestionSelectorDialog(QDialog):
    """Dialog se stromem a checkboxy pro výběr konkrétních bonusových otázek."""

    def __init__(self, owner: "MainWindow", selected_ids: set) -> None:
        super().__init__(owner)
        self.setWindowTitle("Vyberte bonusové otázky")
        self.resize(700, 800)
        self.owner = owner
        self.selected_ids = selected_ids.copy()

        layout = QVBoxLayout(self)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.addWidget(QLabel("Zaškrtněte konkrétní bonusové otázky, které se mají použít v exportu.\n(Pokud nevyberete nic, použijí se náhodné bonusové otázky ze všech dostupných)"))

        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Název zdroje / Otázka"])
        self.tree.header().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.tree.itemChanged.connect(self._on_item_changed)
        layout.addWidget(self.tree)

        self.lbl_total = QLabel(f"Celkem vybráno otázek: {len(self.selected_ids)}")
        self.lbl_total.setStyleSheet("font-weight: bold; color: #ffea00; font-size: 14px; margin-top: 5px;")
        layout.addWidget(self.lbl_total)

        bb = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        bb.accepted.connect(self.accept)
        bb.rejected.connect(self.reject)
        layout.addWidget(bb)

        self.icon_group = owner._generate_icon("S", QColor("#ff5252"), "rect")
        self.icon_sub = owner._generate_icon("P", QColor("#ff8a80"), "rect")
        self.icon_bonus = owner._generate_icon("B", QColor("#ffea00"), "star")

        self._is_populating = True
        self._populate_tree()
        self._is_populating = False
        self._recalculate_total()

    def _populate_tree(self):
        for g in self.owner.root.groups:
            bonus_questions_in_group = self._collect_bonus_questions(g)
            if not bonus_questions_in_group:
                continue

            g_item = QTreeWidgetItem([g.name])
            g_item.setFlags(g_item.flags() | Qt.ItemIsUserCheckable)
            g_item.setFlags(g_item.flags() & ~Qt.ItemIsAutoTristate)
            g_item.setCheckState(0, Qt.Unchecked)
            g_item.setIcon(0, self.icon_group)
            f = g_item.font(0); f.setBold(True); g_item.setFont(0, f)
            self.tree.addTopLevelItem(g_item)

            self._add_subgroups_recursive(g_item, g.subgroups)
            g_item.setExpanded(True)

    def _add_subgroups_recursive(self, parent_item, subgroups):
        for sg in subgroups:
            bonus_questions_in_subgroup = self._collect_bonus_questions(sg)
            if not bonus_questions_in_subgroup:
                continue

            sg_item = QTreeWidgetItem([sg.name])
            sg_item.setFlags(sg_item.flags() | Qt.ItemIsUserCheckable)
            sg_item.setFlags(sg_item.flags() & ~Qt.ItemIsAutoTristate)
            sg_item.setCheckState(0, Qt.Unchecked)
            sg_item.setIcon(0, self.icon_sub)
            parent_item.addChild(sg_item)

            for q in bonus_questions_in_subgroup:
                q_item = QTreeWidgetItem([q.title or "Bonusová otázka"])
                q_item.setFlags(q_item.flags() | Qt.ItemIsUserCheckable)
                q_item.setData(0, Qt.UserRole, {"id": q.id, "type": "question"})
                q_item.setIcon(0, self.icon_bonus)
                
                check_state = Qt.Checked if q.id in self.selected_ids else Qt.Unchecked
                q_item.setCheckState(0, check_state)
                sg_item.addChild(q_item)

            if sg.subgroups:
                self._add_subgroups_recursive(sg_item, sg.subgroups)

            sg_item.setExpanded(True)

    def _collect_bonus_questions(self, node) -> list:
        questions = []
        if hasattr(node, "questions"):
            questions.extend([q for q in node.questions if q.type == "bonus"])
        if hasattr(node, "subgroups"):
            for sub in node.subgroups:
                questions.extend(self._collect_bonus_questions(sub))
        return questions

    def _on_item_changed(self, item, column):
        if self._is_populating: return
        state = item.checkState(0)
        self._is_populating = True
        if not item.data(0, Qt.UserRole): 
            self._set_check_state_recursive(item, state)
        self._is_populating = False
        self._recalculate_total()

    def _set_check_state_recursive(self, item, state):
        for i in range(item.childCount()):
            child = item.child(i)
            child.setCheckState(0, state)
            self._set_check_state_recursive(child, state)

    def _recalculate_total(self):
        count = len(self.get_selected_ids())
        self.lbl_total.setText(f"Celkem vybráno otázek: {count}")

    def get_selected_ids(self) -> set:
        selected = set()
        it = QTreeWidgetItemIterator(self.tree)
        while it.value():
            item = it.value()
            if item.checkState(0) == Qt.Checked:
                data = item.data(0, Qt.UserRole)
                if data and data.get("type") == "question":
                    selected.add(data["id"])
            it += 1
        return selected


# ---------------------- Dialog pro výběr cíle ----------------------
class MoveTargetDialog(QDialog):
    """Dialog pro výběr cílové skupiny/podskupiny pomocí stromu."""
    def __init__(self, owner: "MainWindow") -> None:
        super().__init__(owner)
        self.setWindowTitle("Vyberte cílovou skupinu/podskupinu")
        self.resize(600, 700)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(8)
        
        self.info = QLabel("Vyberte podskupinu (nebo skupinu – vytvoří se Default).")
        layout.addWidget(self.info)
        
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Název", "Typ"])
        
        # Nastavení sloupců
        header = self.tree.header()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        header.setStretchLastSection(False)

        layout.addWidget(self.tree, 1)
        
        bb = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        bb.accepted.connect(self.accept)
        bb.rejected.connect(self.reject)
        layout.addWidget(bb)

        # --- Generování ikon a barev (konzistentní s hlavním stromem) ---
        
        # Barvy (stejné jako v _refresh_tree)
        color_group = QBrush(QColor("#ff5252"))
        color_subgroup = QBrush(QColor("#ff8a80"))
        
        # Ikony (použijeme ownerovu metodu _generate_icon)
        icon_group = owner._generate_icon("S", QColor("#ff5252"), "rect")
        
        for g in owner.root.groups:
            g_item = QTreeWidgetItem([g.name, "Skupina"])
            g_item.setData(0, Qt.UserRole, {"kind": "group", "id": g.id})
            
            # Aplikace stylu
            g_item.setIcon(0, icon_group)
            g_item.setForeground(0, color_group)
            g_item.setForeground(1, color_group)
            f = g_item.font(0); f.setBold(True); f.setPointSize(13); g_item.setFont(0, f)
            
            self.tree.addTopLevelItem(g_item)
            g_item.setExpanded(True)
            
            if g.subgroups:
                self._add_subs(owner, g_item, g.id, g.subgroups)

    def _add_subs(self, owner: "MainWindow", parent_item: QTreeWidgetItem, gid: str, subs: List[Subgroup]) -> None:
        color_subgroup = QBrush(QColor("#ff8a80"))
        icon_subgroup = owner._generate_icon("P", QColor("#ff8a80"), "rect")

        for sg in subs:
            it = QTreeWidgetItem([sg.name, "Podskupina"])
            it.setData(0, Qt.UserRole, {"kind": "subgroup", "id": sg.id, "parent_group_id": gid})
            
            # Aplikace stylu
            it.setIcon(0, icon_subgroup)
            it.setForeground(0, color_subgroup)
            it.setForeground(1, color_subgroup)
            f = it.font(0); f.setBold(True); it.setFont(0, f)
            
            parent_item.addChild(it)
            
            if sg.subgroups:
                self._add_subs(owner, it, gid, sg.subgroups)
                it.setExpanded(True)


    def selected_target(self) -> tuple[str, Optional[str]]:
        items = self.tree.selectedItems()
        if not items:
            return "", None
        meta = items[0].data(0, Qt.UserRole) or {}
        if meta.get("kind") == "subgroup":
            return meta.get("parent_group_id"), meta.get("id")
        elif meta.get("kind") == "group":
            return meta.get("id"), None
        return "", None


# --------------------------- Export Wizard ---------------------------

def cz_day_of_week(dt: datetime) -> str:
    days = ["pondělí","úterý","středa","čtvrtek","pátek","sobota","neděle"]
    return days[dt.weekday()]

def round_dt_to_10m(dt: QDateTime) -> QDateTime:
    m = dt.time().minute()
    rounded = (m // 10) * 10
    return QDateTime(dt.date(), dt.time().addSecs((rounded - m) * 60))


# ---- HTML -> jednoduché mezireprezentace pro DOCX ----

class HTMLToDocxParser(HTMLParser):
    """
    Převádí HTML na seznam odstavců pro DOCX.
    Podporuje vnořené styly, seznamy a odsazení (indentation).
    """
    def __init__(self) -> None:
        super().__init__()
        self.paragraphs: List[dict] = []
        self._stack: List[dict] = [] 
        self._list_stack: List[dict] = [] 
        self._current_runs: List[dict] = []
        self._current_align: str = "left"
        self._current_indent: int = 0
        self._in_body = False
        self._has_body_tag = False
        self._ignore_content = False

    def _start_paragraph(self, prefix: str = "", indent: int = 0) -> None:
        if self._current_runs:
            self._end_paragraph()
        self._current_runs = []
        self._current_align = "left"
        self._current_indent = indent
        for item in reversed(self._stack):
            if item['tag'] in ('p', 'div', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                if item['styles'].get('align'):
                    self._current_align = item['styles']['align']
                break
        self._current_prefix = prefix

    def _end_paragraph(self) -> None:
        if not self._current_runs and not hasattr(self, '_current_prefix'): return
        merged = []
        for r in self._current_runs:
            if r['text'] == "": continue
            if merged and all(merged[-1][k] == r[k] for k in ('b','i','u','color')):
                merged[-1]['text'] += r['text']
            else:
                merged.append(r)
        prefix = getattr(self, '_current_prefix', '')
        if merged or prefix:
            self.paragraphs.append({
                'align': self._current_align,
                'runs': merged,
                'prefix': prefix,
                'indent': self._current_indent
            })
        self._current_runs = []
        if hasattr(self, '_current_prefix'): del self._current_prefix

    def _append_text(self, text: str):
        b = any(item['styles'].get('b') for item in self._stack)
        i = any(item['styles'].get('i') for item in self._stack)
        u = any(item['styles'].get('u') for item in self._stack)
        color = None
        for item in reversed(self._stack):
            if item['styles'].get('color'):
                color = item['styles']['color']
                break
        self._current_runs.append({'text': text, 'b': b, 'i': i, 'u': u, 'color': color})

    def feed(self, data: str) -> None:
        if "<body" in data.lower():
            self._has_body_tag = True; self._in_body = False
        else:
            self._has_body_tag = False; self._in_body = True
        super().feed(data)

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        attrs_d = dict(attrs)
        if tag == 'body': self._in_body = True; return
        if tag in ('head', 'style', 'script', 'meta', 'title', 'html', '!doctype'): self._ignore_content = True; return
        if not self._in_body or self._ignore_content: return

        styles = self._parse_style(attrs_d.get('style', ''))
        if attrs_d.get('align'): styles['align'] = attrs_d['align'].lower()
        if tag in ('b', 'strong'): styles['b'] = True
        if tag in ('i', 'em'): styles['i'] = True
        if tag == 'u': styles['u'] = True
        
        # Explicitní margin - ZVÝŠENÁ CITLIVOST (20px = 1 level)
        margin_left = styles.get('margin-left', '0px')
        explicit_indent_val = 0
        if 'px' in margin_left:
            try:
                px = float(margin_left.replace('px', '').strip())
                explicit_indent_val = int(px / 20) # Changed from 30 to 20
            except: pass

        parent_style = self._stack[-1]['styles'] if self._stack else {}
        merged_styles = parent_style.copy()
        merged_styles.update(styles)
        self._stack.append({'tag': tag, 'attrs': attrs_d, 'styles': merged_styles})

        list_nesting = len(self._list_stack)
        base_indent = max(0, list_nesting - 1) if tag == 'li' else max(0, list_nesting)
        final_indent = base_indent + explicit_indent_val

        if tag in ('p', 'div'):
            self._start_paragraph(indent=final_indent)
        elif tag == 'br':
            self._append_text("\n")
        elif tag in ('ul', 'ol'):
            t = attrs_d.get('type', '1')
            self._list_stack.append({'tag': tag, 'type': t, 'count': 0})
        elif tag == 'li':
            prefix = self._get_list_prefix()
            self._start_paragraph(prefix=prefix, indent=final_indent)

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag == 'body': self._in_body = False; return
        if tag in ('head', 'style', 'script', 'meta', 'title', 'html'): self._ignore_content = False; return
        if not self._in_body or self._ignore_content: return
        if tag in ('p', 'div', 'li'): self._end_paragraph()
        if tag in ('ul', 'ol'):
            if self._list_stack: self._list_stack.pop()
        for i in range(len(self._stack) - 1, -1, -1):
            if self._stack[i]['tag'] == tag:
                del self._stack[i:]
                break

    def handle_data(self, data):
        if not self._in_body or self._ignore_content: return
        if not data: return
        in_block = False
        current_indent = 0
        list_nesting = len(self._list_stack)
        current_indent = max(0, list_nesting)
        for item in reversed(self._stack):
            if item['tag'] in ('p', 'div', 'li'):
                in_block = True
                break
        if not in_block:
            if not data.strip(): return
            self._start_paragraph(indent=current_indent)
        self._append_text(data)

    def _parse_style(self, style_str: str) -> dict:
        res = {}
        if not style_str: return res
        for part in style_str.split(';'):
            if ':' in part:
                k,v = part.split(':', 1)
                k=k.strip().lower(); v=v.strip().lower()
                if k=='color':
                    m = re.search(r'#?([0-9a-f]{6})', v)
                    if m: res['color']=m.group(1)
                    else:
                         m2 = re.search(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', v)
                         if m2: r,g,b=[int(x) for x in m2.groups()]; res['color']=f"{r:02X}{g:02X}{b:02X}"
                elif k=='text-align': res['align']=v
                elif k=='margin-left': res['margin-left']=v
                elif k=='font-weight':
                    if 'bold' in v or (re.search(r'\d+',v) and int(re.search(r'\d+',v).group(0))>=600): res['b']=True
                elif k=='font-style' and 'italic' in v: res['i']=True
                elif k=='text-decoration' and 'underline' in v: res['u']=True
        return res

    def _get_list_prefix(self) -> str:
        if not self._list_stack: return ""
        L = self._list_stack[-1]
        if L['tag'] == 'ul': return "•\t"
        elif L['tag'] == 'ol':
            L['count'] += 1; n = L['count']; t = L.get('type', '1')
            val = f"{n}."
            if t == 'a': val = f"{self._to_alpha(n, False)}."
            elif t == 'A': val = f"{self._to_alpha(n, True)}."
            elif t == 'i': val = f"{self._to_roman(n, False)}."
            elif t == 'I': val = f"{self._to_roman(n, True)}."
            return f"{val}\t"
        return ""
    def _to_alpha(self, n, upper):
        s=""; n-=1
        while n>=0: s=chr(97+n%26)+s; n=n//26-1
        return s.upper() if upper else s
    def _to_roman(self, n, upper):
        val=[(50,'L'),(40,'XL'),(10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]; res=""
        for v,r in val:
            while n>=v: res+=r; n-=v
        return res if upper else res.lower()

# ---- Tvorba WordprocessingML prvků (w:p, w:r) ----

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {'w': W_NS}

def w_tag(name: str) -> str:
    return f"{{{W_NS}}}{name}"

def make_w_run(text: str, b: bool=False, i: bool=False, u: bool=False, color: Optional[str]=None) -> ET.Element:
    r = ET.Element(w_tag("r"))
    rpr = ET.SubElement(r, w_tag("rPr"))
    if b:
        ET.SubElement(rpr, w_tag("b"))
    if i:
        ET.SubElement(rpr, w_tag("i"))
    if u:
        uel = ET.SubElement(rpr, w_tag("u"))
        uel.set(w_tag("val"), "single")
    if color:
        cel = ET.SubElement(rpr, w_tag("color"))
        cel.set(w_tag("val"), color.upper())
    t = ET.SubElement(r, w_tag("t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text or ""
    return r

def make_w_paragraph(align: str, runs: List[dict], prefix: str="") -> ET.Element:
    p = ET.Element(w_tag("p"))
    ppr = ET.SubElement(p, w_tag("pPr"))
    jc = ET.SubElement(ppr, w_tag("jc"))
    jc.set(w_tag("val"), {"left":"left","center":"center","right":"right","justify":"both"}.get(align, "left"))
    if prefix:
        p.append(make_w_run(prefix))
    for r in runs:
        p.append(make_w_run(r.get('text',''), r.get('b'), r.get('i'), r.get('u'), r.get('color')))
    return p


# --------------------------- Export Wizard ---------------------------

class ExportWizard(QWizard):
    def __init__(self, owner: "MainWindow") -> None:
        super().__init__(owner)
        self.setWindowTitle("Export DOCX – Průvodce")
        self.setWizardStyle(QWizard.ModernStyle)
        self.owner = owner
        self.resize(1600, 1200)

        # Lokalizace
        self.setButtonText(QWizard.BackButton, "< Zpět")
        self.setButtonText(QWizard.NextButton, "Další >")
        self.setButtonText(QWizard.CommitButton, "Dokončit")
        self.setButtonText(QWizard.FinishButton, "Dokončit")
        self.setButtonText(QWizard.CancelButton, "Zrušit")

        # IDs
        self.ID_PAGE1 = 0
        self.ID_PAGE2 = 1
        self.ID_PAGE3 = 2

        # Data
        self.template_path: Optional[Path] = None
        self.output_path: Optional[Path] = None
        self.output_changed_manually = False
        self.placeholders_q = []
        self.placeholders_b = []
        self.selection_map = {}
        self.cached_hash = ""
        self.has_datum_cas = False
        self.has_pozn = False
        self.has_min_max = False
        
        # NOVÉ: Ukládáme seznam vybraných zdrojů pro Multi Export [{'id':..., 'type':...}]
        self.multi_selected_sources = []
        # NOVÉ: Ukládáme seznam ID vybraných bonusových otázek
        self.multi_selected_bonus_ids = set()

        # Načtení uložených cest
        self.settings_file = self.owner.project_root / "data" / "export_settings.json"
        self.stored_settings = self._load_settings()

        # Cesty (Default)
        default_templates_dir = self.owner.project_root / "data" / "Šablony"
        default_output_dir = self.owner.project_root / "data" / "Vygenerované testy"
        default_print_dir = self.owner.project_root / "data" / "Tisk"
        default_templates_dir.mkdir(parents=True, exist_ok=True)
        default_output_dir.mkdir(parents=True, exist_ok=True)
        default_print_dir.mkdir(parents=True, exist_ok=True)
        self.templates_dir = Path(self.stored_settings.get("templates_dir", default_templates_dir))
        self.output_dir = Path(self.stored_settings.get("output_dir", default_output_dir))
        self.print_dir = Path(self.stored_settings.get("print_dir", default_print_dir))

        last_template = self.stored_settings.get("last_template")
        if last_template and Path(last_template).exists():
            self.default_template = Path(last_template)
        else:
            self.default_template = self.templates_dir / "template.docx"

        # --- INIT STRÁNEK ---
        self.page1 = QWizardPage()
        self.page2 = QWizardPage()
        self.page3 = QWizardPage()
        self._build_page1_content()
        self._build_page2_content()
        self._build_page3_content()

        self.setPage(self.ID_PAGE1, self.page1)
        self.setPage(self.ID_PAGE2, self.page2)
        self.setPage(self.ID_PAGE3, self.page3)

        self.setStartId(self.ID_PAGE1)

        if self.default_template.exists():
            self.le_template.setText(str(self.default_template))
            self.template_path = self.default_template
            QTimer.singleShot(100, self._scan_placeholders)
        else:
            self.le_template.setText(str(self.default_template))
            self._update_path_indicators()

        self._update_default_output()
        self._update_path_indicators()

        
    def _load_settings(self) -> dict:
        if self.settings_file.exists():
            try:
                with open(self.settings_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                pass
        return {}

    def _save_settings(self):
        data = {
            "templates_dir": str(self.templates_dir),
            "output_dir": str(self.output_dir),
            "print_dir": str(self.print_dir),
            "last_template": self.le_template.text()
        }
        try:
            with open(self.settings_file, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2)
        except Exception as e:
            print(f"Chyba ukládání nastavení exportu: {e}")

    def _update_path_indicators(self):
        # Šablona
        t_path = Path(self.le_template.text())
        if t_path.exists() and t_path.is_file():
            self.lbl_status_template.setText("✅ OK")
            self.lbl_status_template.setStyleSheet("color: #81c784; font-weight: bold;")
        else:
            self.lbl_status_template.setText("❌ Chybí")
            self.lbl_status_template.setStyleSheet("color: #ef5350; font-weight: bold;")
            
        # Výstupní složka
        if self.output_dir.exists() and self.output_dir.is_dir():
            self.lbl_status_out_dir.setText("✅ OK")
            self.lbl_status_out_dir.setStyleSheet("color: #81c784; font-weight: bold;")
        else:
            self.lbl_status_out_dir.setText("❌ Chybí")
            self.lbl_status_out_dir.setStyleSheet("color: #ef5350; font-weight: bold;")

        # Tisková složka
        if self.print_dir.exists() and self.print_dir.is_dir():
            self.lbl_status_print_dir.setText("✅ OK")
            self.lbl_status_print_dir.setStyleSheet("color: #81c784; font-weight: bold;")
        else:
            self.lbl_status_print_dir.setText("❌ Chybí")
            self.lbl_status_print_dir.setStyleSheet("color: #ef5350; font-weight: bold;")


    # --- Build Content Methods ---

    def _build_page1_content(self):
        self.page1.setTitle("Krok 1: Výběr šablony a nastavení cest")
        l1 = QVBoxLayout(self.page1)
        
        # GroupBox: Parametry
        gb_params = QGroupBox("Parametry testu")
        form_params = QFormLayout()
        
        self.le_prefix = QLineEdit("MůjTest")
        self.le_prefix.textChanged.connect(self._update_default_output)
        form_params.addRow("Prefix verze:", self.le_prefix)
        
        self.dt_edit = QDateTimeEdit(QDateTime.currentDateTime())
        self.dt_edit.setDisplayFormat("dd.MM.yyyy HH:mm")
        self.dt_edit.setCalendarPopup(True)
        self.dt_edit.dateTimeChanged.connect(self._update_default_output)
        form_params.addRow("Datum testu:", self.dt_edit)
        
        gb_params.setLayout(form_params)
        l1.addWidget(gb_params)

        # GroupBox: Soubory a složky
        gb_files = QGroupBox("Cesty k souborům")
        grid_files = QGridLayout()
        grid_files.setColumnStretch(1, 1) # Input stretch
        
        # 1. Šablona
        grid_files.addWidget(QLabel("Šablona:"), 0, 0)
        self.le_template = QLineEdit()
        self.le_template.textChanged.connect(self._on_templ_change)
        grid_files.addWidget(self.le_template, 0, 1)
        
        btn_t = QPushButton("Vybrat...")
        btn_t.clicked.connect(self._choose_template)
        grid_files.addWidget(btn_t, 0, 2)
        
        self.lbl_status_template = QLabel("?")
        self.lbl_status_template.setFixedWidth(80)
        grid_files.addWidget(self.lbl_status_template, 0, 3)

        # 2. Výstup DOCX (Složka)
        grid_files.addWidget(QLabel("Složka pro testy:"), 1, 0)
        self.le_out_dir = QLineEdit(str(self.output_dir))
        self.le_out_dir.setReadOnly(True) # Editace jen přes dialog pro bezpečí
        grid_files.addWidget(self.le_out_dir, 1, 1)
        
        btn_od = QPushButton("Změnit...")
        btn_od.clicked.connect(self._choose_output_dir)
        grid_files.addWidget(btn_od, 1, 2)
        
        self.lbl_status_out_dir = QLabel("?")
        grid_files.addWidget(self.lbl_status_out_dir, 1, 3)
        
        # 3. Výstup PDF Tisk (Složka)
        grid_files.addWidget(QLabel("Složka pro tisk:"), 2, 0)
        self.le_print_dir = QLineEdit(str(self.print_dir))
        self.le_print_dir.setReadOnly(True)
        grid_files.addWidget(self.le_print_dir, 2, 1)
        
        btn_pd = QPushButton("Změnit...")
        btn_pd.clicked.connect(self._choose_print_dir)
        grid_files.addWidget(btn_pd, 2, 2)
        
        self.lbl_status_print_dir = QLabel("?")
        grid_files.addWidget(self.lbl_status_print_dir, 2, 3)
        
        # 4. Konkrétní soubor (náhled názvu)
        grid_files.addWidget(QLabel("Název souboru:"), 3, 0)
        self.le_output = QLineEdit()
        self.le_output.textChanged.connect(self._on_output_text_changed)
        grid_files.addWidget(self.le_output, 3, 1, 1, 2) # Span 2 sloupce
        
        gb_files.setLayout(grid_files)
        l1.addWidget(gb_files)
        
        self.lbl_scan_info = QLabel("Info: Čekám na načtení šablony...")
        self.lbl_scan_info.setStyleSheet("color: gray; font-style: italic; margin-top: 10px;")
        l1.addWidget(self.lbl_scan_info)

    def _build_page2_content(self):
        self.page2.setTitle("Krok 2: Přiřazení otázek do šablony")
        self.page2.initializePage = self._init_page2
        main_layout = QVBoxLayout(self.page2)

        # 1. Info Panel
        self.info_box_p2 = QGroupBox("Kontext exportu")
        self.info_box_p2.setStyleSheet("QGroupBox { font-weight: bold; border: 1px solid #555; margin-top: 6px; } QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 3px; }")
        l_info = QFormLayout(self.info_box_p2)
        self.lbl_templ_p2 = QLabel("-")
        self.lbl_out_p2 = QLabel("-")
        l_info.addRow("Vstupní šablona:", self.lbl_templ_p2)
        l_info.addRow("Výstupní soubor:", self.lbl_out_p2)
        main_layout.addWidget(self.info_box_p2)

        # 2. Volba režimu
        self.mode_box = QGroupBox("Režim exportu")
        self.mode_box.setStyleSheet("""
            QGroupBox { font-weight: bold; margin-top: 10px; }
            QRadioButton { font-size: 14px; padding: 5px; border-radius: 4px; }
            QRadioButton::indicator { width: 16px; height: 16px; }
            QRadioButton:checked { color: #61dafb; font-weight: bold; background-color: #3a3a3c; }
            QRadioButton:hover { background-color: #2d2d30; }
        """)
        l_mode = QHBoxLayout(self.mode_box)
        self.rb_mode_single = QRadioButton("Jednotlivý export (Standardní)")
        self.rb_mode_single.setToolTip("Vytvoří jeden soubor s ručně vybranými otázkami.")
        self.rb_mode_single.setChecked(True)
        self.rb_mode_multi = QRadioButton("Hromadný export (Generátor variant)")
        self.rb_mode_multi.setToolTip("Vytvoří více kopií testu. Otázky 1-10 budou vybrány náhodně pro každou kopii.")
        self.mode_group = QButtonGroup(self)
        self.mode_group.addButton(self.rb_mode_single, 0)
        self.mode_group.addButton(self.rb_mode_multi, 1)
        self.mode_group.buttonToggled.connect(self._on_mode_toggled)
        l_mode.addWidget(self.rb_mode_single)
        l_mode.addWidget(self.rb_mode_multi)
        l_mode.addStretch()
        main_layout.addWidget(self.mode_box)

        # 3. Nastavení pro hromadný export
        self.widget_multi_options = QWidget()
        self.widget_multi_options.setVisible(False)
        self.widget_multi_options.setStyleSheet("background-color: #2d2d30; border-radius: 4px; padding: 10px; border: 1px solid #444;")
        l_multi = QGridLayout(self.widget_multi_options)
        l_multi.setContentsMargins(5, 5, 5, 5)
        
        l_multi.addWidget(QLabel("Počet kopií:"), 0, 0)
        self.spin_multi_count = QSpinBox()
        self.spin_multi_count.setRange(2, 50)
        self.spin_multi_count.setValue(2)
        self.spin_multi_count.setStyleSheet("padding: 4px;")
        l_multi.addWidget(self.spin_multi_count, 0, 1)
        
        l_multi.addWidget(QLabel("Zdroje otázek (pro |Otázka|1-10):"), 1, 0)
        self.btn_select_sources = QPushButton("Vybrat zdroje...")
        self.btn_select_sources.setCursor(Qt.PointingHandCursor)
        self.btn_select_sources.setStyleSheet("""
            QPushButton { background-color: #0d47a1; color: white; border: none; padding: 6px 12px; border-radius: 4px; font-weight: bold; }
            QPushButton:hover { background-color: #1565c0; }
            QPushButton:pressed { background-color: #0d47a1; }
        """)
        self.btn_select_sources.clicked.connect(self._on_select_sources_clicked)
        self.lbl_selected_sources = QLabel("Nevybráno (použijí se všechny otázky)")
        self.lbl_selected_sources.setStyleSheet("color: #aaa; font-style: italic; margin-left: 8px;")
        self.lbl_selected_sources.setWordWrap(True)
        src_container = QWidget(); src_layout = QHBoxLayout(src_container); src_layout.setContentsMargins(0,0,0,0)
        src_layout.addWidget(self.btn_select_sources); src_layout.addWidget(self.lbl_selected_sources, 1)
        l_multi.addWidget(src_container, 1, 1)
        
        l_multi.addWidget(QLabel("Bonusové otázky (pro |Bonus|...):"), 2, 0)
        self.btn_select_bonus = QPushButton("Vybrat bonusové...")
        self.btn_select_bonus.setCursor(Qt.PointingHandCursor)
        self.btn_select_bonus.setStyleSheet("""
            QPushButton { background-color: #b08d00; color: white; border: none; padding: 6px 12px; border-radius: 4px; font-weight: bold; }
            QPushButton:hover { background-color: #c9a300; }
            QPushButton:pressed { background-color: #b08d00; }
        """)
        self.btn_select_bonus.clicked.connect(self._on_select_bonus_clicked)
        self.lbl_selected_bonus = QLabel("Nevybráno (použijí se náhodné)")
        self.lbl_selected_bonus.setStyleSheet("color: #aaa; font-style: italic; margin-left: 8px;")
        self.lbl_selected_bonus.setWordWrap(True)
        bonus_container = QWidget(); bonus_layout = QHBoxLayout(bonus_container); bonus_layout.setContentsMargins(0,0,0,0)
        bonus_layout.addWidget(self.btn_select_bonus); bonus_layout.addWidget(self.lbl_selected_bonus, 1)
        l_multi.addWidget(bonus_container, 2, 1)
        main_layout.addWidget(self.widget_multi_options)

        # 4. Hlavní obsah (Dva sloupce: Strom | Sloty)
        columns_layout = QHBoxLayout()
        self.widget_left_panel = QWidget()
        left_layout = QVBoxLayout(self.widget_left_panel)
        left_layout.setContentsMargins(0,0,0,0)
        left_layout.addWidget(QLabel("<b>Dostupné otázky:</b>"))
        self.tree_source = QTreeWidget()
        self.tree_source.setHeaderLabels(["Struktura otázek"])
        self.tree_source.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.tree_source.setContextMenuPolicy(Qt.CustomContextMenu)
        self.tree_source.customContextMenuRequested.connect(self._on_tree_source_context_menu)
        if hasattr(self, "_on_tree_selection"):
            self.tree_source.itemSelectionChanged.connect(self._on_tree_selection)
        left_layout.addWidget(self.tree_source)
        
        self.btn_assign_multi = QPushButton(">> Přiřadit vybrané na volné pozice >>")
        self.btn_assign_multi.setToolTip("Doplní vybrané otázky zleva na první volná místa v šabloně vpravo.")
        self.btn_assign_multi.clicked.connect(self._assign_selected_multi)
        left_layout.addWidget(self.btn_assign_multi)
        columns_layout.addWidget(self.widget_left_panel, 4)
        
        # Pravý panel - ZDE JE OPRAVA (QScrollArea + layout_slots)
        right_layout = QVBoxLayout()
        right_header = QHBoxLayout()
        right_header.addWidget(QLabel("<b>Sloty v šabloně:</b>"))
        right_header.addStretch()
        self.btn_clear_all = QPushButton("Vyprázdnit vše")
        self.btn_clear_all.clicked.connect(self._clear_all_assignments)
        right_header.addWidget(self.btn_clear_all)
        right_layout.addLayout(right_header)
        
        self.scroll_slots = QScrollArea()
        self.scroll_slots.setWidgetResizable(True)
        self.widget_slots = QWidget()
        # Tady definujeme self.layout_slots, který vaše _init_page2 potřebuje
        self.layout_slots = QVBoxLayout(self.widget_slots)
        self.layout_slots.setSpacing(6)
        self.layout_slots.addStretch()
        self.scroll_slots.setWidget(self.widget_slots)
        
        right_layout.addWidget(self.scroll_slots)
        columns_layout.addLayout(right_layout, 6)
        main_layout.addLayout(columns_layout, 3)

        # 5. Náhled (self.preview_box)
        self.preview_box = QGroupBox("Náhled vybrané otázky")
        preview_layout = QVBoxLayout(self.preview_box)
        preview_layout.setContentsMargins(5,5,5,5)
        self.text_preview_q = QTextEdit()
        self.text_preview_q.setReadOnly(True)
        self.text_preview_q.setMaximumHeight(120)
        self.text_preview_q.setStyleSheet("QTextEdit { background-color: #2e2e2e; color: #ffffff; font-size: 14px; border: 1px solid #555; padding: 5px; }")
        preview_layout.addWidget(self.text_preview_q)
        main_layout.addWidget(self.preview_box, 1)

    def _refresh_slots_list(self):
        """Vykreslí seznam slotů. V režimu Multi ukáže jen placeholdery."""
        self.slots_list_widget.clear()
        is_multi = (self.mode_group.checkedId() == 1)

        # Helper
        def add_slot_item(placeholder, is_bonus):
            item = QListWidgetItem()
            
            if is_multi:
                # Režim Hromadný export -> Ignorujeme výběr, vše je "Náhodné"
                if is_bonus:
                    txt = f"{placeholder} : [Náhodný výběr BONUS]"
                    color = QColor("#d4a017") # Zlatavá pro bonus
                else:
                    txt = f"{placeholder} : [Náhodný výběr z variant]"
                    color = QColor("#757575") # Šedá
                
                item.setText(txt)
                item.setForeground(QBrush(color))
                # Neaktivní item (nejde na něj kliknout)
                item.setFlags(Qt.NoItemFlags) 
                
            else:
                # Režim Jednotlivý export -> Ukazujeme výběr z self.selection_map
                qid = self.selection_map.get(placeholder)
                q = self.owner._find_question_by_id(qid) if qid else None
                
                if q:
                    title = q.title or "Otázka"
                    item.setText(f"{placeholder} : {title}")
                    if is_bonus:
                        item.setForeground(QBrush(QColor("#ffea00"))) # Jasná žlutá
                    else:
                        item.setForeground(QBrush(Qt.white))
                    item.setToolTip(q.text_html)
                    item.setData(Qt.UserRole, qid) # Uložíme ID pro kliknutí
                else:
                    item.setText(f"{placeholder} : --- NEVYPLNĚNO ---")
                    item.setForeground(QBrush(QColor("#888888")))
                    
                item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)

            self.slots_list_widget.addItem(item)

        # 1. Klasické sloty
        for ph in self.placeholders_q:
            add_slot_item(ph, False)
            
        # 2. Bonusové sloty
        for ph in self.placeholders_b:
            add_slot_item(ph, True)


    def _on_select_sources_clicked(self):
        # Zjistíme počet potřebných otázek ze šablony
        needed_count = len(self.placeholders_q)
        
        dlg = MultiSourceDialog(self.owner, self.multi_selected_sources)
        if dlg.exec():
            self.multi_selected_sources = dlg.get_selected_items()
            total_available = dlg.get_total_selected_count() # Počet unikátních dostupných otázek
            
            names = [item['name'].split(" (")[0] for item in self.multi_selected_sources]
            count_sources = len(self.multi_selected_sources)
            
            if count_sources == 0:
                # Pokud nic nevybral -> bere se "všechno".
                # Zde bychom mohli spočítat celkový počet v DB pro přesnost,
                # ale pro jednoduchost napíšeme:
                self.lbl_selected_sources.setText("Nevybráno (použijí se všechny otázky)")
                self.lbl_selected_sources.setStyleSheet("color: #aaa; font-style: italic; margin-left: 5px;")
            else:
                short_list = ", ".join(names[:2])
                if len(names) > 2:
                    short_list += f" a {len(names)-2} dalších"
                
                # -- KONTROLA DOSTATKU OTÁZEK --
                if total_available < needed_count:
                    # NEDOSTATEK -> ČERVENĚ + VAROVÁNÍ
                    text = f"<b>VYBRÁNO MÁLO: {total_available} z {needed_count}</b> (Zdroj: {short_list})"
                    style = "color: #ff5252; font-weight: bold; margin-left: 5px;"
                    QMessageBox.warning(self, "Nedostatek otázek", 
                                        f"Pozor! Vybrané zdroje obsahují pouze {total_available} otázek,\n"
                                        f"ale šablona vyžaduje {needed_count}.\n\n"
                                        "Některé otázky se budou opakovat nebo zůstanou prázdné.")
                else:
                    # DOSTATEK -> MODŘE/ZELENĚ
                    text = f"<b>{total_available} otázek</b> (Potřeba: {needed_count}) z: {short_list}"
                    style = "color: #81c784; font-weight: bold; margin-left: 5px;" # Zelená
                
                self.lbl_selected_sources.setText(text)
                self.lbl_selected_sources.setStyleSheet(style)

    def _on_select_bonus_clicked(self):
        dlg = BonusQuestionSelectorDialog(self.owner, self.multi_selected_bonus_ids)
        if dlg.exec() == QDialog.Accepted:
            self.multi_selected_bonus_ids = dlg.get_selected_ids()
            count = len(self.multi_selected_bonus_ids)
            if count > 0:
                self.lbl_selected_bonus.setText(f"Vybráno konkrétních: {count}")
                self.lbl_selected_bonus.setStyleSheet("color: #ffea00; font-weight: bold;")
            else:
                self.lbl_selected_bonus.setText("Nevybráno (použijí se náhodné)")
                self.lbl_selected_bonus.setStyleSheet("color: #aaa; font-style: italic;")

    def _on_mode_toggled(self, btn, checked):
        if not checked: 
            return
        
        is_multi = (self.mode_group.checkedId() == 1)
        
        # 1. Viditelnost nastavení pro multi
        self.widget_multi_options.setVisible(is_multi)
        
        # 2. SKRYTÍ LEVÉHO PANELU (Dostupné otázky) v Multi režimu
        #    (V Single režimu je vidět, v Multi ne)
        self.widget_left_panel.setVisible(not is_multi)
        
        # 3. Skrytí náhledu v multi režimu
        if hasattr(self, "preview_box"):
            self.preview_box.setVisible(not is_multi)
            if is_multi:
                 self.text_preview_q.clear()
        
        # 4. Zablokování tlačítek (pro jistotu, i když jsou skrytá)
        self.btn_assign_multi.setEnabled(not is_multi)
        self.btn_clear_all.setEnabled(not is_multi)
        
        # 5. Překreslení slotů (aktivní/neaktivní)
        self._update_slots_visuals(is_multi)


    def _update_slots_visuals(self, is_multi: bool):
        """Projteruje všechny sloty v layoutu a nastaví stav tlačítek podle režimu."""
        # Procházíme widgety v layoutu (pozor na stretch a labely)
        for i in range(self.layout_slots.count()):
            item = self.layout_slots.itemAt(i)
            w = item.widget()
            if not w: continue
            
            # Hledáme naše řádky (mají property 'placeholder')
            ph = w.property("placeholder")
            if not ph: continue
            
            is_bonus_slot = w.property("is_bonus") # Získali jsme při vytváření
            
            # Získáme tlačítka z layoutu řádku
            # Layout je: 0:Label, 1:BtnAssign, 2:BtnClear
            layout = w.layout()
            if not layout or layout.count() < 3: continue
            
            btn_assign = layout.itemAt(1).widget()
            btn_clear = layout.itemAt(2).widget()
            
            if is_multi:
                # REŽIM MULTI: Tlačítka neaktivní, text indikuje automatiku
                btn_assign.setEnabled(False)
                btn_clear.setEnabled(False)
                btn_clear.setVisible(False) # Skryjeme křížek, nedává smysl
                
                if is_bonus_slot:
                    btn_assign.setText("[Náhodný výběr BONUS]")
                    # Bonusové tlačítko 'disabled' styl
                    btn_assign.setStyleSheet("color: #888; background-color: #333; border: 1px solid #444;")
                else:
                    btn_assign.setText("[Náhodný výběr z VARIANT]")
                    btn_assign.setStyleSheet("color: #888; background-color: #333; border: 1px solid #444;")
            
            else:
                # REŽIM SINGLE: Tlačítka aktivní, text podle výběru
                btn_assign.setEnabled(True)
                btn_clear.setEnabled(True)
                btn_clear.setVisible(True)
                btn_assign.setStyleSheet("") # Reset stylu
                
                # Obnovíme text podle aktuální selection_map
                qid = self.selection_map.get(ph)
                if qid:
                    q = self.owner._find_question_by_id(qid)
                    if q:
                         btn_assign.setText(q.title)
                    else:
                         btn_assign.setText("??? (Nenalezeno)")
                else:
                    btn_assign.setText("--- Volné ---")

    def _clear_all_assignments(self) -> None:
        """Vymaže všechna přiřazení otázek ve slotech."""
        if not self.selection_map:
            return

        if QMessageBox.question(self, "Vyprázdnit", "Opravdu zrušit všechna přiřazení?") != QMessageBox.Yes:
            return

        self.selection_map.clear()
        self._init_page2() # Obnoví UI
        self.owner.statusBar().showMessage("Všechna přiřazení byla zrušena.", 3000)

    def _build_page3_content(self):
        self.page3.setTitle("Krok 3: Kontrola a Export")
        self.page3.initializePage = self._init_page3
        
        main_layout = QVBoxLayout(self.page3)
        
        # Info Panel
        self.info_box_p3 = QGroupBox("Kontext exportu")
        self.info_box_p3.setStyleSheet("QGroupBox { font-weight: bold; border: 1px solid #555; margin-top: 6px; } QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 3px; }")
        l_info = QFormLayout(self.info_box_p3)
        self.lbl_templ_p3 = QLabel("-")
        self.lbl_out_p3 = QLabel("-")
        l_info.addRow("Vstupní šablona:", self.lbl_templ_p3)
        l_info.addRow("Výstupní soubor:", self.lbl_out_p3)
        main_layout.addWidget(self.info_box_p3)
        
        # NOVÉ: Exportní volby
        self.options_box = QGroupBox("Exportní volby")
        self.options_box.setStyleSheet("font-weight: bold; margin-top: 6px;")
        l_opts = QVBoxLayout(self.options_box)
        
        self.chk_export_pdf = QCheckBox("Exportovat do PDF pro tisk (složka /data/Tisk/)")
        self.chk_export_pdf.setChecked(True)
        self.chk_export_pdf.setToolTip("DOCX se automaticky převede na PDF. V hromadném režimu budou všechny varianty spojeny do jednoho PDF.")
        l_opts.addWidget(self.chk_export_pdf)
        
        main_layout.addWidget(self.options_box)

        # --- NOVÉ: Progress Bar a Status Label ---
        status_layout = QVBoxLayout()
        self.lbl_status_final = QLabel("Připraveno k exportu.")
        self.lbl_status_final.setStyleSheet("font-weight: bold; color: #aaa;")
        self.lbl_status_final.setAlignment(Qt.AlignCenter)
        status_layout.addWidget(self.lbl_status_final)

        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 2px solid #555;
                border-radius: 4px;
                text-align: center;
                background-color: #2d2d30;
                color: white;
            }
            QProgressBar::chunk {
                background-color: #4caf50;
                width: 10px;
            }
        """)
        status_layout.addWidget(self.progress_bar)
        main_layout.addLayout(status_layout)
        # -----------------------------------------
        
        # Náhled
        lbl_prev = QLabel("<b>Náhled obsahu testu:</b>")
        main_layout.addWidget(lbl_prev)
        
        self.preview_edit = QTextEdit()
        self.preview_edit.setReadOnly(True)
        # Dark Theme CSS pro QTextEdit
        self.preview_edit.setStyleSheet("QTextEdit { background-color: #252526; color: #e0e0e0; border: 1px solid #3e3e42; }")
        main_layout.addWidget(self.preview_edit)

        # Hash label
        self.lbl_hash_preview = QLabel("Hash: -")
        self.lbl_hash_preview.setWordWrap(True)
        self.lbl_hash_preview.setStyleSheet("color: #555; font-family: Consolas, Monaco, monospace; font-size: 10px; margin-top: 5px;")
        self.lbl_hash_preview.setTextInteractionFlags(Qt.TextSelectableByMouse)
        main_layout.addWidget(self.lbl_hash_preview)


    # --- Helpers & Logic ---

    def _update_default_output(self):
        if self.output_changed_manually and self.sender() == self.le_output:
            # Pokud uživatel edituje ručně celou cestu, necháme ho
            self.output_path = Path(self.le_output.text())
            return

        prefix = self.le_prefix.text().strip()
        # Nahrazení nepovolených znaků
        prefix = re.sub(r'[\\/*?:"<>|]', "_", prefix)
        
        dt = self.dt_edit.dateTime()
        # 2. část: YYYY-MM-DD_HHMM
        date_time_str = dt.toString("yyyy-MM-dd_HHmm")
        
        # 3. část: Timestamp (aktuální čas v sekundách)
        timestamp = str(int(datetime.now().timestamp()))
        
        # Sestavení názvu: PREFIX_DATETIME_TIMESTAMP.docx
        filename = f"{prefix}_{date_time_str}_{timestamp}.docx"
        
        self.output_path = self.output_dir / filename
        
        # Blokujeme signál, abychom necyklili přes _on_output_text_changed
        self.le_output.blockSignals(True)
        self.le_output.setText(str(self.output_path))
        self.le_output.blockSignals(False)
        
        self.page1.completeChanged.emit()

    def _on_output_text_changed(self, text):
        self.output_changed_manually = True
        self.output_path = Path(text)

    def _choose_template(self):
        start_dir = str(self.templates_dir) if self.templates_dir.exists() else str(self.owner.project_root)
        path, _ = QFileDialog.getOpenFileName(self, "Vybrat šablonu", start_dir, "*.docx")
        if path:
            self.le_template.setText(path)
            self.templates_dir = Path(path).parent
            # Uložit nastavení ihned
            self._save_settings()

    def _choose_output(self):
        path, _ = QFileDialog.getSaveFileName(self, "Cíl exportu", str(self.default_output), "*.docx")
        if path:
            self.le_output.setText(path)
            
    def _choose_output_dir(self):
        d = QFileDialog.getExistingDirectory(self, "Vybrat složku pro testy", str(self.output_dir))
        if d:
            self.output_dir = Path(d)
            self.le_out_dir.setText(d)
            # Uložit nastavení ihned
            self._save_settings()
            
            self._update_path_indicators()
            self._update_default_output() # Přegenerovat cestu k souboru

    def _choose_print_dir(self):
        d = QFileDialog.getExistingDirectory(self, "Vybrat složku pro tisk", str(self.print_dir))
        if d:
            self.print_dir = Path(d)
            self.le_print_dir.setText(d)
            # Uložit nastavení ihned
            self._save_settings()
            
            self._update_path_indicators()

    def _on_templ_change(self, text: str):
        self.template_path = Path(text)
        self._update_path_indicators()
        
        # Uložit nastavení ihned
        self._save_settings()
        
        if self.template_path.exists() and self.template_path.suffix.lower() == '.docx':
             self._scan_placeholders()
        else:
             self.lbl_scan_info.setText("Šablona nenalezena nebo není .docx")
        
        self.page1.completeChanged.emit()


    def _scan_placeholders(self):
        if not self.template_path or not self.template_path.exists():
            return

        try:
            doc = docx.Document(self.template_path)
            full_text = ""
            # Načteme text z celého dokumentu pro regex
            for p in doc.paragraphs: full_text += p.text + "\n"
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs: full_text += p.text + "\n"
            
            import re
            # Regex pro <Placeholder> i {Placeholder}
            placeholders = re.findall(r'[<{]([A-Za-z0-9ÁČĎÉĚÍŇÓŘŠŤÚŮÝŽáčďéěíňóřšťúůýž]+)[>}]', full_text)
            placeholders = sorted(list(set(placeholders)))
            
            self.placeholders_q = [p for p in placeholders if re.match(r'^Otázka\d+$', p)]
            self.placeholders_q.sort(key=lambda x: int(re.findall(r'\d+', x)[0]))
            
            self.placeholders_b = [p for p in placeholders if re.match(r'^BONUS\d+$', p)]
            self.placeholders_b.sort(key=lambda x: int(re.findall(r'\d+', x)[0]))
            
            self.has_datumcas = any(x in placeholders for x in ['DatumČas', 'DatumCas', 'DATUMCAS'])
            self.has_pozn = any(x in placeholders for x in ['PoznamkaVerze', 'POZNAMKAVERZE'])
            self.has_minmax = (any('MinBody' in x for x in placeholders), any('MaxBody' in x for x in placeholders))
            
            msg = f"Nalezeno: {len(self.placeholders_q)}x Otázka, {len(self.placeholders_b)}x BONUS."
            if self.has_minmax[0]: msg += " (S body)."
            self.lbl_scan_info.setText(msg)
            
        except Exception as e:
            self.lbl_scan_info.setText(f"Chyba čtení šablony: {e}")

    # --- Page Initializers ---
    
    def _on_tree_selection(self):
        sel = self.tree_source.selectedItems()
        # Náhled funguje jen pokud je vybrána přesně jedna položka
        if not sel or len(sel) != 1:
            self.text_preview_q.clear()
            return

        item = sel[0]
        data = item.data(0, Qt.UserRole)

        # Očekáváme buď přímo ID otázky, nebo dict s metadaty (kind/id/...)
        if not data:
            self.text_preview_q.setText("--- (Vyberte konkrétní otázku pro náhled) ---")
            return

        if isinstance(data, dict):
            # Náhled má smysl jen pro položku typu 'question'
            if data.get("kind") != "question":
                self.text_preview_q.setText("--- (Vyberte konkrétní otázku pro náhled) ---")
                return
            qid = data.get("id")
        else:
            qid = data

        if not qid:
            self.text_preview_q.setText("--- (Vyberte konkrétní otázku pro náhled) ---")
            return

        q = self.owner._find_question_by_id(qid)
        if q:
            import re
            import html

            html_content = q.text_html or ""

            # 1. Odstranění <style>...</style> a <head>...</head> i s obsahem
            # Flag re.DOTALL zajistí, že . matchuje i newlines
            clean = re.sub(r'<style.*?>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            clean = re.sub(r'<head.*?>.*?</head>', '', clean, flags=re.DOTALL | re.IGNORECASE)

            # 2. Odstranění všech ostatních tagů <...>
            clean = re.sub(r'<[^>]+>', ' ', clean)

            # 3. Decode entities (&nbsp;, &lt;...)
            clean = html.unescape(clean)

            # 4. Squeeze whitespace – odstranění prázdných řádků
            lines = [line.strip() for line in clean.splitlines() if line.strip()]
            final_text = "\n".join(lines)

            self.text_preview_q.setText(final_text)
        else:
            self.text_preview_q.clear()

    def _init_page2(self):
        try:
            # Info update
            t_name = self.template_path.name if self.template_path else "Nevybráno"
            o_name = self.output_path.name if self.output_path else "Nevybráno"
            self.lbl_templ_p2.setText(t_name)
            self.lbl_out_p2.setText(o_name)

            # Rescan
            if not self.placeholders_q and not self.placeholders_b:
                self._scan_placeholders()

            # 1. Clear Tree
            self.tree_source.blockSignals(True)
            self.tree_source.clear()
            
            # 2. Clear Slots
            while self.layout_slots.count():
                item = self.layout_slots.takeAt(0)
                if item.widget(): item.widget().deleteLater()
            self.layout_slots.addStretch()
            
            # Barvy a ikony
            color_group = QBrush(QColor("#ff5252"))
            color_subgroup = QBrush(QColor("#ff8a80"))
            
            # Helper pro generování ikon (musí být dostupný self.owner._generate_icon)
            def generate_colored_icon(text, color, shape="rect"):
                 return self.owner._generate_icon(text, color, shape)

            icon_group = generate_colored_icon("S", QColor("#ff5252"), "rect")
            icon_sub = generate_colored_icon("P", QColor("#ff8a80"), "rect")
            
            # 4. Populate Tree (Seřazeno abecedně)
            def add_subgroup_recursive(parent_item, subgroup_list, parent_gid):
                # SEŘAZENÍ PODSKUPIN
                sorted_subs = sorted(subgroup_list, key=lambda s: s.name.lower())
                
                for sg in sorted_subs:
                    # Tree Item
                    sg_item = QTreeWidgetItem([sg.name])
                    sg_item.setIcon(0, icon_sub)
                    sg_item.setForeground(0, color_subgroup)
                    f = sg_item.font(0); f.setBold(True); sg_item.setFont(0, f)
                    sg_item.setData(0, Qt.UserRole, {
                        "kind": "subgroup", 
                        "id": sg.id, 
                        "parent_group_id": parent_gid
                    })
                    parent_item.addChild(sg_item)
                    
                    # SEŘAZENÍ OTÁZEK
                    sorted_qs = sorted(sg.questions, key=lambda q: (q.title or "").lower())
                    
                    for q in sorted_qs:
                        label_type = "Klasická" if str(q.type).lower() != "bonus" else "BONUS"
                        is_bonus = (label_type == "BONUS")
                        info = f"({q.points} b)" if not is_bonus else f"(Bonus: {q.bonus_correct})"
                        label_text = f"{q.title} {info}"
                        q_item = QTreeWidgetItem([label_text])
                        q_item.setData(0, Qt.UserRole, {
                            "kind": "question",
                            "id": q.id,
                            "parent_group_id": parent_gid,
                            "parent_subgroup_id": sg.id
                        })
                        
                        # Vizualizace (Ikona + Obrázek)
                        has_img = bool(getattr(q, "image_path", "") and os.path.exists(q.image_path))
                        # Voláme metodu ownera pro vizuál
                        self.owner._apply_question_item_visuals(q_item, q.type, has_image=has_img)
                        
                        sg_item.addChild(q_item)
                    
                    if sg.subgroups:
                        add_subgroup_recursive(sg_item, sg.subgroups, parent_gid)

            # SEŘAZENÍ HLAVNÍCH SKUPIN
            groups = sorted(self.owner.root.groups, key=lambda g: g.name.lower())
            
            for g in groups:
                g_item = QTreeWidgetItem([g.name])
                g_item.setIcon(0, icon_group)
                g_item.setForeground(0, color_group)
                f = g_item.font(0); f.setBold(True); g_item.setFont(0, f)
                g_item.setData(0, Qt.UserRole, {
                    "kind": "group", 
                    "id": g.id
                })
                self.tree_source.addTopLevelItem(g_item)
                add_subgroup_recursive(g_item, g.subgroups, g.id)
                
            self.tree_source.expandAll()
            self.tree_source.blockSignals(False)

            # 5. Create Slots
            def create_slot_row(ph, is_bonus):
                row_w = QWidget()
                row_l = QHBoxLayout(row_w)
                row_l.setContentsMargins(0,0,0,0)
                lbl_name = QLabel(f"{ph}:")
                lbl_name.setFixedWidth(120)
                if is_bonus: lbl_name.setStyleSheet("color: #ffea00;")
                else: lbl_name.setStyleSheet("color: #42a5f5;")
                
                btn_assign = QPushButton("Vybrat...")
                # Načteme existující přiřazení (pro Single mode)
                qid = self.selection_map.get(ph)
                if qid:
                    q = self.owner._find_question_by_id(qid)
                    if q:
                        btn_assign.setText(q.title)
                        btn_assign.setToolTip(f"{q.text_html[:100]}...")
                    else:
                        btn_assign.setText("???")
                else:
                    btn_assign.setText("--- Volné ---")
                    
                btn_assign.clicked.connect(lambda checked, p=ph: self._on_slot_assign_clicked(p))
                
                btn_clear = QPushButton("X")
                btn_clear.setFixedWidth(30)
                btn_clear.clicked.connect(lambda checked, p=ph: self._on_slot_clear_clicked(p))
                
                # Uložíme do layoutu
                row_l.addWidget(lbl_name)       # index 0
                row_l.addWidget(btn_assign, 1)  # index 1 (stretch)
                row_l.addWidget(btn_clear)      # index 2
                
                row_w.setProperty("placeholder", ph)
                # Uložíme si informaci, zda je to bonusový slot (pro update_visuals)
                row_w.setProperty("is_bonus", is_bonus)
                
                # Vložíme před stretch
                self.layout_slots.insertWidget(self.layout_slots.count()-1, row_w)

            if self.placeholders_q:
                lbl = QLabel("--- KLASICKÉ OTÁZKY ---")
                lbl.setStyleSheet("font-weight:bold; color:#4da6ff; margin-top:5px;")
                self.layout_slots.insertWidget(self.layout_slots.count()-1, lbl)
                for ph in self.placeholders_q: create_slot_row(ph, False)

            if self.placeholders_b:
                lbl = QLabel("--- BONUSOVÉ OTÁZKY ---")
                lbl.setStyleSheet("font-weight:bold; color:#ffcc00; margin-top:10px;")
                self.layout_slots.insertWidget(self.layout_slots.count()-1, lbl)
                for ph in self.placeholders_b: create_slot_row(ph, True)
            
            is_multi = (self.mode_group.checkedId() == 1)
            
            # Zavoláme metodu pro aktualizaci stavu tlačítek podle režimu
            self._update_slots_visuals(is_multi)
            
            # Pokud existuje metoda pro refresh stromu (v minulé verzi možná byla), zavoláme ji
            if hasattr(self, "_refresh_tree_visuals"): self._refresh_tree_visuals()
            
            # Update labelů pro vybrané zdroje
            if self.multi_selected_sources:
                count_src = len(self.multi_selected_sources)
                self.lbl_selected_sources.setText(f"Vybráno {count_src} zdrojů.")
                self.lbl_selected_sources.setStyleSheet("color: #42a5f5; margin-left: 5px;")
            
            # Update labelů pro bonusové zdroje (pokud jsou vybrány)
            if self.multi_selected_bonus_ids:
                count_bonus = len(self.multi_selected_bonus_ids)
                self.lbl_selected_bonus.setText(f"Vybráno konrétních: {count_bonus}")
                self.lbl_selected_bonus.setStyleSheet("color: #ffea00; margin-left: 5px;")

        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Chyba", f"Chyba při inicializaci stránky 2:\n{e}")


    def _on_slot_assign_clicked(self, ph: str) -> None:
        # Jednoduchý výběr: Otevře dialog se seznamem dostupných otázek
        # Pro jednoduchost: jen zrušíme výběr, aby se uvolnilo místo? Ne, má to přiřadit.
        # V rámci "Minimal intervention" bych neměl přidávat komplexní dialogy, pokud tam nebyly.
        # Ale tlačítko tam je.
        # Pokud tam metoda byla, je to ok. Pokud ne, přidám basic logiku.
        # Zkusíme předpokládat, že tam nebyla a přidáme ji.
        
        # Dialog pro výběr konkrétní otázky do slotu
        # (Může využít existující výběr ve stromu nebo nový dialog)
        
        # Pro teď: Pokud je vybrána otázka ve stromu, přiřadí ji.
        items = self.tree_source.selectedItems()
        if not items:
            QMessageBox.information(self, "Výběr", "Vyberte otázku ve stromu vlevo a pak klikněte sem.")
            return
            
        meta = items[0].data(0, Qt.UserRole) or {}
        if meta.get("kind") != "question":
            return
            
        q = self.owner._find_question(meta["parent_group_id"], meta["parent_subgroup_id"], meta["id"])
        if q:
            self._assign_question_to_slot(ph, q)

    def _on_slot_clear_clicked(self, ph: str) -> None:
        if ph in self.selection_map:
            del self.selection_map[ph]
            self._init_page2() # Refresh

    def _on_tree_source_context_menu(self, pos):
        # Kontextové menu nad stromem zdrojových otázek (Krok 2)
        
        # 1. Zjistíme položku pod myší (pro single action)
        item_under_mouse = self.tree_source.itemAt(pos)
        
        # 2. Zjistíme všechny vybrané položky (pro multi action)
        selected_items = self.tree_source.selectedItems()
        
        menu = QMenu(self)
        has_action = False

        # A) AKCE PRO KONKRÉTNÍ OTÁZKU (JEDNOTLIVÝ VÝBĚR)
        # (Ponecháme stávající logiku, pokud je vybrána jen jedna otázka nebo kliknuto na otázku)
        if item_under_mouse:
            data = item_under_mouse.data(0, Qt.UserRole)
            if data and data.get("kind") == "question":
                qid = data.get("id")
                q = self.owner._find_question_by_id(qid)
                if q:
                    # Submenu "Přiřadit k..."
                    menu_assign = menu.addMenu("Přiřadit k...")
                    
                    # Najdeme volné sloty
                    free_slots = []
                    if q.type == "classic":
                        for ph in self.placeholders_q:
                            if ph not in self.selection_map: free_slots.append(ph)
                    else: # bonus
                        for ph in self.placeholders_b:
                            if ph not in self.selection_map: free_slots.append(ph)
                    
                    if not free_slots:
                        a = menu_assign.addAction("(Žádné volné sloty)")
                        a.setEnabled(False)
                    else:
                        for slot in free_slots:
                            action = QAction(slot, self.tree_source)
                            action.triggered.connect(lambda checked=False, s_slot=slot, q_id=qid: self._assign_from_context(s_slot, q_id))
                            menu_assign.addAction(action)
                    has_action = True

        # B) AKCE PRO MULTI VÝBĚR OTÁZEK (PŘIŘADIT NA PRVNÍ VOLNÉ)
        # (Pokud je vybráno více otázek)
        questions_selected = [it for it in selected_items if it.data(0, Qt.UserRole).get("kind") == "question"]
        if len(questions_selected) > 0:
            act_multi = QAction(f"Přiřadit vybrané ({len(questions_selected)}) na volné pozice", self)
            act_multi.triggered.connect(self._assign_selected_multi)
            menu.addAction(act_multi)
            has_action = True

        # C) NOVÉ: AKCE PRO MULTI VÝBĚR SKUPIN/PODSKUPIN (NÁHODNÝ VÝBĚR Z NICH)
        # Filtrujeme vybrané položky, které jsou group nebo subgroup
        containers_selected = [it for it in selected_items if it.data(0, Qt.UserRole).get("kind") in ("group", "subgroup")]
        
        if len(containers_selected) > 0:
            # Pokud je vybráno více kontejnerů (nebo i jeden, ale chceme sjednotit logiku)
            label = "této větve" if len(containers_selected) == 1 else "vybraných větví"
            act_random_multi = QAction(f"Naplnit volné pozice náhodně z {label}", self)
            act_random_multi.triggered.connect(lambda: self._assign_random_from_multiple_nodes(containers_selected))
            menu.addAction(act_random_multi)
            has_action = True

        if has_action:
            menu.exec(self.tree_source.viewport().mapToGlobal(pos))
            
    # --- POMOCNÉ METODY PRO MULTI-SELECT VE STROMU (JEDNOTLIVÝ EXPORT) ---

    def _find_group_in_root(self, gid: str):
        """Najde Group podle ID v rootu."""
        for g in self.owner.root.groups:
            if g.id == gid:
                return g
        return None

    def _find_subgroup_in_group(self, group, sgid: str):
        """Najde Subgroup podle ID v dané Group (rekurzivně)."""
        def search(subs):
            for s in subs:
                if s.id == sgid:
                    return s
                if s.subgroups:
                    found = search(s.subgroups)
                    if found:
                        return found
            return None
        
        # Bezpečný přístup k subgroups
        subs = getattr(group, "subgroups", []) or []
        return search(subs)

    def _collect_all_questions_in_branch(self, node) -> List["Question"]:
        """Rekurzivně vrátí všechny otázky v uzlu a jeho podskupinách."""
        res: List["Question"] = []
        
        # 1. Otázky přímo v uzlu
        if hasattr(node, "questions") and node.questions:
            res.extend(node.questions)
        
        # 2. Otázky v podskupinách
        if hasattr(node, "subgroups") and node.subgroups:
            for sub in node.subgroups:
                res.extend(self._collect_all_questions_in_branch(sub))
        
        return res

    def _assign_random_from_multiple_nodes(self, items: List[QTreeWidgetItem]) -> None:
        """Naplní volné klasické sloty náhodnými otázkami ze všech vybraných větví."""
        
        # 1. Posbírat všechny klasické otázky ze všech vybraných větví
        all_available_qs: List["Question"] = []
        
        for item in items:
            meta = item.data(0, Qt.UserRole) or {}
            kind = meta.get("kind")
            
            if kind not in ("group", "subgroup"):
                continue
            
            node = None
            if kind == "group":
                gid = meta.get("id")
                node = self._find_group_in_root(gid)
                
            elif kind == "subgroup":
                gid = meta.get("parent_group_id")
                sgid = meta.get("id")
                if gid and sgid:
                    group = self._find_group_in_root(gid)
                    if group:
                        node = self._find_subgroup_in_group(group, sgid)
            
            if node:
                qs_in_branch = self._collect_all_questions_in_branch(node)
                # Filtrujeme jen klasické otázky
                classic_qs = [q for q in qs_in_branch if getattr(q, "type", "") == "classic"]
                all_available_qs.extend(classic_qs)

        # 2. Odstranit duplicity (podle ID)
        unique_qs = {}
        for q in all_available_qs:
            unique_qs[q.id] = q
        pool = list(unique_qs.values())
        
        if not pool:
            QMessageBox.information(self, "Info", "Ve vybraných větvích nejsou žádné klasické otázky.")
            return

        # 3. Volné sloty pro klasické otázky
        free_slots = [ph for ph in self.placeholders_q if ph not in self.selection_map]
        if not free_slots:
            QMessageBox.information(self, "Info", "Všechny sloty jsou již obsazené.")
            return

        # 4. Odstranit z poolu otázky, které už jsou někde použité
        used_ids = set(self.selection_map.values())
        pool = [q for q in pool if q.id not in used_ids]
        
        if not pool:
            QMessageBox.information(self, "Info", "Všechny dostupné otázky z vybraných větví už jsou použity.")
            return

        import random
        needed = len(free_slots)
        assigned_count = 0

        if len(pool) >= needed:
            picked = random.sample(pool, needed)
            for idx, ph in enumerate(free_slots):
                self._assign_question_to_slot(ph, picked[idx])
                assigned_count += 1
        else:
            # Máme méně otázek než slotů -> naplníme, co jde, bez opakování
            random.shuffle(pool)
            for idx, q in enumerate(pool):
                if idx >= len(free_slots):
                    break
                self._assign_question_to_slot(free_slots[idx], q)
                assigned_count += 1
            
            QMessageBox.warning(
                self,
                "Nedostatek otázek",
                f"Doplněno pouze {assigned_count} otázek.\n"
                "Ve vybraných větvích nebylo dostatek unikátních otázek pro všechny sloty."
            )

        if assigned_count > 0:
            try:
                self.owner.statusBar().showMessage(f"Doplněno {assigned_count} otázek z vybraných větví.", 3000)
            except Exception:
                pass


    # V ExportWizard nebo použít owner metodu
    def _find_subgroup_helper(self, parent_gid, sub_id):
        # Musíme najít skupinu
        group = None
        for g in self.owner.root.groups:
            if g.id == parent_gid:
                group = g; break
        
        if not group: return None
        
        # Rekurzivní hledání v group
        def search(subs):
            for s in subs:
                if s.id == sub_id: return s
                if s.subgroups:
                    found = search(s.subgroups)
                    if found: return found
            return None
            
        return search(group.subgroups)


    def _refresh_tree_visuals(self) -> None:
        """Aktualizuje vizuální stav položek ve stromu (vybrané vs volné)."""
        iterator = QTreeWidgetItemIterator(self.tree_source)
        used_ids = set(self.selection_map.values())
        
        c_classic = QColor("#42a5f5")
        c_bonus = QColor("#ffea00")
        c_used = QColor("#666666") # Šedá pro vybrané
        
        while iterator.value():
            item = iterator.value()
            meta = item.data(0, Qt.UserRole) or {}
            
            if meta.get("kind") == "question":
                qid = meta.get("id")
                txt = item.text(0)
                clean_txt = txt.replace(" [VYBRÁNO]", "")
                
                # Zjištění typu
                q = self.owner._find_question_by_id(qid)
                is_bonus = False
                if q:
                    is_bonus = (str(q.type).lower() == "bonus" or q.type == 1)
                
                if qid in used_ids:
                    # Vybráno -> Šedá
                    item.setText(0, clean_txt + " [VYBRÁNO]")
                    item.setForeground(0, QBrush(c_used))
                    f = item.font(0); f.setItalic(True); f.setBold(False)
                    item.setFont(0, f)
                else:
                    # Volné -> Barva podle typu
                    item.setText(0, clean_txt)
                    target_color = c_bonus if is_bonus else c_classic
                    item.setForeground(0, QBrush(target_color))
                    
                    f = item.font(0); f.setItalic(False)
                    if is_bonus: f.setBold(True)
                    else: f.setBold(False)
                    item.setFont(0, f)
            
            iterator += 1

    def _assign_single_question_from_context(self, meta: dict) -> None:
        """Přiřadí jednu otázku z kontextového menu."""
        q = self.owner._find_question(meta["parent_group_id"], meta["parent_subgroup_id"], meta["id"])
        if not q:
            return
            
        # Využijeme logiku z multi-selectu (je to jako select 1 item)
        # Ale pro jednoduchost přímo:
        if q.id in self.selection_map.values():
            self.owner.statusBar().showMessage("Otázka již je přiřazena.", 2000)
            return
            
        target_ph = None
        if q.type == "classic":
            for ph in self.placeholders_q:
                if ph not in self.selection_map:
                    target_ph = ph; break
        elif q.type == "bonus":
            for ph in self.placeholders_b:
                if ph not in self.selection_map:
                    target_ph = ph; break
                    
        if target_ph:
            self._assign_question_to_slot(target_ph, q)
            self.owner.statusBar().showMessage("Otázka přiřazena.", 2000)
        else:
            QMessageBox.information(self, "Plno", "Není volné místo pro tento typ otázky.")


    def _assign_random_from_context(self, meta: dict) -> None:
        """Vybere náhodné otázky z dané větve a doplní je na volná místa."""
        # 1. Zjistit volné sloty
        free_slots = []
        # Musíme iterovat přes layout slotů, ale nemáme přímý list.
        # Můžeme projít placeholders_q a placeholders_b a zkontrolovat selection_map.
        
        # Spojíme seznamy placeholderů (nejprve klasické, pak bonusové, nebo jak chceme plnit)
        # Obvykle chceme plnit klasické klasickými a bonusové bonusovými? 
        # Zadání specifikuje "otázky", ale v aplikaci je rozdělení na typy.
        # Pro jednoduchost a konzervativnost: 
        # Pokud je slot pro klasickou otázku (v placeholders_q), hledáme klasické otázky.
        # Pokud je slot pro bonus (v placeholders_b), hledáme bonusové.
        
        # Sběr všech otázek ve větvi
        all_questions_in_branch = []
        
        def collect_recursive(gid, sgid):
            if sgid:
                sg = self.owner._find_subgroup(gid, sgid)
                if sg:
                    all_questions_in_branch.extend(sg.questions)
                    for sub in sg.subgroups:
                        collect_recursive(gid, sub.id)
            else:
                g = self.owner._find_group(gid)
                if g:
                    for sub in g.subgroups:
                        collect_recursive(gid, sub.id)

        gid = meta.get("id") if meta.get("kind") == "group" else meta.get("parent_group_id")
        sgid = None if meta.get("kind") == "group" else meta.get("id")
        
        collect_recursive(gid, sgid)
        
        if not all_questions_in_branch:
            QMessageBox.information(self, "Info", "V této větvi nejsou žádné otázky.")
            return

        # Rozdělení dostupných otázek podle typu
        available_classic = [q for q in all_questions_in_branch if q.type == "classic"]
        available_bonus = [q for q in all_questions_in_branch if q.type == "bonus"]
        
        # Promíchat pro náhodnost
        import random
        random.shuffle(available_classic)
        random.shuffle(available_bonus)
        
        # 2. Plnění slotů
        # (Iterujeme přes placeholdery a pokud je volný, vezmeme otázku)
        
        assigned_count = 0
        
        # Klasické sloty
        for ph in self.placeholders_q:
            if ph not in self.selection_map: # Volný slot
                # Najít otázku, která ještě NENÍ použita v mapě
                for q in available_classic:
                    if q.id not in self.selection_map.values():
                        self._assign_question_to_slot(ph, q)
                        available_classic.remove(q) # Odebrat, aby se neopakovala
                        assigned_count += 1
                        break
        
        # Bonusové sloty
        for ph in self.placeholders_b:
            if ph not in self.selection_map:
                for q in available_bonus:
                    if q.id not in self.selection_map.values():
                        self._assign_question_to_slot(ph, q)
                        available_bonus.remove(q)
                        assigned_count += 1
                        break
                        
        if assigned_count > 0:
            self.owner.statusBar().showMessage(f"Doplněno {assigned_count} otázek.", 3000)
        else:
            QMessageBox.information(self, "Info", "Nebylo možné doplnit žádné další otázky (buď nejsou volná místa, nebo došly unikátní otázky).")

    def _assign_selected_multi(self) -> None:
        """Přiřadí vybrané otázky ve stromu na první volná místa."""
        items = self.tree_source.selectedItems()
        selected_questions = []
        
        for it in items:
            meta = it.data(0, Qt.UserRole) or {}
            if meta.get("kind") == "question":
                q = self.owner._find_question(meta["parent_group_id"], meta["parent_subgroup_id"], meta["id"])
                if q:
                    selected_questions.append(q)
        
        if not selected_questions:
            QMessageBox.information(self, "Info", "Vyberte ve stromu alespoň jednu otázku.")
            return

        assigned_count = 0
        
        for q in selected_questions:
            if q.id in self.selection_map.values():
                continue
                
            target_ph = None
            if q.type == "classic":
                for ph in self.placeholders_q:
                    if ph not in self.selection_map:
                        target_ph = ph; break
            elif q.type == "bonus":
                for ph in self.placeholders_b:
                    if ph not in self.selection_map:
                        target_ph = ph; break
            
            if target_ph:
                self.selection_map[target_ph] = q.id # Přímý zápis
                assigned_count += 1
        
        if assigned_count > 0:
            self._init_page2() # Hromadný refresh na konci (aktualizuje sloty i strom)
            self.owner.statusBar().showMessage(f"Přiřazeno {assigned_count} otázek.", 3000)
        else:
            self.owner.statusBar().showMessage("Nebylo co přiřadit (vše plné nebo vybrané už použité).", 3000)

    def _assign_question_to_slot(self, ph: str, q: Question) -> None:
        self.selection_map[ph] = q.id
        self._init_page2() # Refresh UI
        
        # Aktualizace UI (tlačítka slotu)
        # Musíme najít widget odpovídající tomuto placeholderu v layoutu
        # Protože nemáme přímou referenci ph -> widget, projdeme layout.
        # (Nebo si můžeme držet mapu ph -> button při vytváření, ale "Maintenance mode" velí neměnit init příliš).
        
        count = self.layout_slots.count()
        for i in range(count):
            item = self.layout_slots.itemAt(i)
            w = item.widget()
            if w and hasattr(w, "property") and w.property("placeholder") == ph:
                # Našli jsme widget (SlotRowWidget nebo podobný)
                # Předpokládám, že má metodu set_question nebo update_ui
                # Pokud neznám vnitřní strukturu SlotRowWidget, musím ji odhadnout z _init_page2 loopu.
                # Ale jelikož nemám kód SlotRowWidget (pokud existuje), udělám to přes refresh celé stránky nebo chytřeji.
                
                # Z logiky _init_page2 (kterou jsem neviděl celou, ale bývá to loop):
                # Obvykle se sloty generují znovu.
                # Pro jednoduchost a robustnost: Zavoláme refresh slotů.
                # Ale to by bylo pomalé.
                
                # Zkusíme najít tlačítko/label v tom widgetu.
                # Předpoklad: w je nějaký container.
                
                # Nejčistší v rámci "Black Box":
                # Znovu vygenerovat sloty je jistota.
                pass
        
        # Protože nemám detailní přístup k widgetům slotů, zavolám obnovu UI slotů.
        # Toto je sice méně efektivní, ale bezpečné.
        self._refresh_slots_ui()

    def _refresh_slots_ui(self) -> None:
        """Znovu vykreslí pravý panel se sloty (volá _init_page2 logiku pro sloty)."""
        # Protože _init_page2 dělá clear a populate, můžeme ji zavolat, 
        # ALE musíme dát pozor, aby nesmazala selection_map (což _init_page2 obvykle nedělá, ta ji čte).
        # Pokud _init_page2 RESCANUJE placeholdery, mohlo by to vadit.
        # V _init_page2 je: if not self.placeholders_q ... scan. Takže ok.
        
        # Zavoláme část _init_page2, která kreslí sloty.
        # Nebo jednoduše celou _init_page2, pokud je idempotentní.
        # Z kódu výše: _init_page2 maže tree a sloty a plní je.
        # To je trochu heavy (refresh tree zruší výběr).
        # Takže raději jen sloty.
        
        # Implementace refresh slotů (zkopírováno/vytaženo z _init_page2 logiky):
        
        # 1. Clear Slots
        while self.layout_slots.count():
            item = self.layout_slots.takeAt(0)
            if item.widget(): item.widget().deleteLater()
        self.layout_slots.addStretch()
        
        # 2. Re-populate
        # (Tuto logiku nemám k dispozici ve snippetu, musím ji "získat" nebo napsat znovu podle logiky aplikace)
        # Pokud nemám kód pro plnění slotů (loop over placeholders), nemohu to napsat.
        
        # ŘEŠENÍ: Zavolám self._init_page2() s tím, že se smířím se zrušením výběru ve stromu.
        # Uživatel právě klikl na tlačítko nebo menu, takže akce skončila.
        # Obnova stránky je akceptovatelná.
        self._init_page2()


    def _add_slot_widget(self, placeholder_name, allowed_type):
        w = QWidget()
        l = QHBoxLayout(w)
        l.setContentsMargins(0,2,0,2)
        
        lbl_ph = QLabel(f"{placeholder_name}:")
        lbl_ph.setFixedWidth(80)
        
        btn_sel = QPushButton("Vybrat...")
        btn_sel.setFixedWidth(80)
        
        current_qid = self.selection_map.get(placeholder_name)
        current_title = "(nevybráno)"
        if current_qid:
            q = self.owner._find_question_by_id(current_qid)
            if q: current_title = q.title

        lbl_val = QLabel(current_title)
        lbl_val.setStyleSheet("color: gray; font-style: italic;" if not current_qid else "color: white; font-weight: bold;")
        
        btn_clr = QPushButton("X")
        btn_clr.setFixedWidth(30)
        btn_clr.setEnabled(bool(current_qid))
        
        l.addWidget(lbl_ph)
        l.addWidget(lbl_val, 1)
        l.addWidget(btn_sel)
        l.addWidget(btn_clr)
        
        def select_current():
            sel = self.tree_source.selectedItems()
            if not sel:
                QMessageBox.information(self, "Info", "Nejprve označte otázku v levém seznamu.")
                return
            item = sel[0]
            qid = item.data(0, Qt.UserRole)
            if not qid: return
            
            q = self.owner._find_question_by_id(qid)
            if not q: return
            
            if q.type != allowed_type:
                QMessageBox.warning(self, "Typ nesedí", f"Do slotu {placeholder_name} nelze vložit otázku typu {q.type}.")
                return
            
            old_qid = self.selection_map.get(placeholder_name)
            if old_qid: self._show_tree_item(old_qid)
            
            self.selection_map[placeholder_name] = qid
            lbl_val.setText(q.title)
            lbl_val.setStyleSheet("color: white; font-weight: bold;")
            btn_clr.setEnabled(True)
            item.setHidden(True)

        def clear_current():
            old_qid = self.selection_map.get(placeholder_name)
            if old_qid:
                del self.selection_map[placeholder_name]
                self._show_tree_item(old_qid)
            lbl_val.setText("(nevybráno)")
            lbl_val.setStyleSheet("color: gray; font-style: italic;")
            btn_clr.setEnabled(False)

        btn_sel.clicked.connect(select_current)
        btn_clr.clicked.connect(clear_current)
        self.layout_slots.insertWidget(self.layout_slots.count()-1, w)

    def _show_tree_item(self, qid):
        it = QTreeWidgetItemIterator(self.tree_source)
        while it.value():
            if it.value().data(0, Qt.UserRole) == qid:
                it.value().setHidden(False)
                break
            it += 1

    def _show_context_menu(self, position):
        item = self.tree_source.itemAt(position)
        if not item: return
        qid = item.data(0, Qt.UserRole)
        if not qid: return
        
        q = self.owner._find_question_by_id(qid)
        if not q: return

        from PySide6.QtWidgets import QMenu
        from PySide6.QtGui import QAction
        
        menu = QMenu()
        menu_assign = menu.addMenu("Přiřadit k...")
        
        free_slots = []
        if q.type == 'classic':
            for ph in self.placeholders_q:
                if ph not in self.selection_map: free_slots.append(ph)
        else:
            for ph in self.placeholders_b:
                if ph not in self.selection_map: free_slots.append(ph)
        
        if not free_slots:
            a = menu_assign.addAction("(Žádné volné sloty)")
            a.setEnabled(False)
        else:
            for slot in free_slots:
                action = QAction(slot, self.tree_source)
                action.triggered.connect(lambda checked=False, s=slot, q_id=qid: self._assign_from_context(s, q_id))
                menu_assign.addAction(action)

        menu.exec(self.tree_source.viewport().mapToGlobal(position))

    def _assign_from_context(self, slot_name, qid):
        # 1. Uložit přiřazení
        self.selection_map[slot_name] = qid
        
        # 2. Aktualizovat UI Slotu (pravý panel)
        found_widget = False
        for i in range(self.layout_slots.count()):
            w = self.layout_slots.itemAt(i).widget()
            if w and isinstance(w, QWidget):
                children = w.findChildren(QLabel)
                if children and children[0].text() == f"{slot_name}:":
                    # Našli jsme správný slot widget
                    
                    # Index 1 = Tlačítko s názvem (prostřední)
                    lbl_val = w.layout().itemAt(1).widget()
                    # Index 2 = Tlačítko Clear (napravo) - OPRAVENO z 3 na 2
                    btn_clr = w.layout().itemAt(2).widget()
                    
                    q = self.owner._find_question_by_id(qid)
                    if q:
                        lbl_val.setText(q.title)
                        # Nastavíme styl tlačítka ve slotu (aby bylo vidět, že je plné)
                        lbl_val.setStyleSheet("color: white; font-weight: bold;") 
                        btn_clr.setEnabled(True)
                        btn_clr.setVisible(True)
                    found_widget = True
                    break
        
        # 3. Aktualizovat vizuál stromu pomocí existující metody
        # (Tato metoda projde strom a obarví vše, co je v selection_map)
        if hasattr(self, "_refresh_tree_visuals"):
            self._refresh_tree_visuals()
        else:
            # Fallback, pokud metoda neexistuje (pro jistotu, ale měla by tam být)
            print("Warning: _refresh_tree_visuals method not found in ExportWizard.")
            # Zde případně fallback na ruční obarvení, pokud by metoda chyběla

    def _init_page3(self):
        try:
            # 1. Generování hashe
            ts = str(datetime.now().timestamp())
            salt = secrets.token_hex(16)
            data_to_hash = f"{ts}{salt}"
            self._cached_hash = hashlib.sha3_256(data_to_hash.encode("utf-8")).hexdigest()
            
            if hasattr(self, "lbl_hash_preview"):
                self.lbl_hash_preview.setText(f"SHA3-256 Hash:\n{self._cached_hash}")

            t_name = self.template_path.name if self.template_path else "Nevybráno"
            o_name = self.output_path.name if self.output_path else "Nevybráno"
            self.lbl_templ_p3.setText(t_name)
            self.lbl_out_p3.setText(o_name)

            total_bonus_points = 0.0
            min_loss = 0.0
            
            bg_color = "#252526"; text_color = "#e0e0e0"; border_color = "#555555"
            sec_q_bg = "#2d3845"; sec_b_bg = "#453d2d"; sec_s_bg = "#2d452d"
            
            prefix = self.le_prefix.text().strip()
            today = datetime.now().strftime("%Y-%m-%d")
            verze_preview = f"{prefix} {today}" 
            
            is_multi = (self.mode_group.checkedId() == 1)
            
            multi_info = ""
            if is_multi:
                multi_count = self.spin_multi_count.value()
                multi_info = f"<tr><td colspan='2' style='color: #ffcc00; font-weight: bold;'>⚡ Hromadný export: {multi_count} verzí (stejný hash)</td></tr>"

            html = f"""
            <html>
            <body style='font-family: Arial, sans-serif; background-color: {bg_color}; color: {text_color};'>
            <h2 style='color: #61dafb; border-bottom: 2px solid #61dafb;'>Souhrn testu</h2>
            <table width='100%' style='margin-bottom: 20px; color: {text_color};'>
                <tr>
                    <td><b>Verze:</b> {verze_preview}</td>
                    <td align='right'><b>Datum:</b> {self.dt_edit.dateTime().toString("dd.MM.yyyy HH:mm")}</td>
                </tr>
                {multi_info}
                <tr>
                    <td colspan='2' style='font-size: 10px; color: #888;'><b>Hash:</b> {self._cached_hash[:32]}...</td>
                </tr>
            </table>
            """

           # Klasické
            html += f"<h3 style='background-color: {sec_q_bg}; padding: 5px; border-left: 4px solid #4da6ff;'>1. Klasické otázky</h3>"
            html += f"<table width='100%' border='0' cellspacing='0' cellpadding='5' style='color: {text_color};'>"
            for ph in self.placeholders_q:
                # Pokud je multi, ignorujeme selection_map a rovnou píšeme, že je to náhodné
                if is_multi:
                     html += f"<tr><td width='100' style='color:#888;'>{ph}:</td><td colspan='2' style='color:#ffcc00;'>[Náhodný výběr pro každou verzi]</td></tr>"
                else:
                    qid = self.selection_map.get(ph)
                    if qid:
                        q = self.owner._find_question_by_id(qid)
                        if q:
                            title_clean = re.sub(r'<[^>]+>', '', q.title)
                            html += f"<tr><td width='100' style='color:#888;'>{ph}:</td><td><b>{title_clean}</b></td><td align='right'>({q.points} b)</td></tr>"
                    else:
                        html += f"<tr><td width='100' style='color:#ff5555;'>{ph}:</td><td colspan='2' style='color:#ff5555;'>--- NEVYPLNĚNO ---</td></tr>"
            html += "</table>"

            # Bonusy
            html += f"<h3 style='background-color: {sec_b_bg}; padding: 5px; border-left: 4px solid #ffcc00;'>2. Bonusové otázky</h3>"
            html += f"<table width='100%' border='0' cellspacing='0' cellpadding='5' style='color: {text_color};'>"
            for ph in self.placeholders_b:
                if is_multi:
                     html += f"<tr><td width='100' style='color:#888;'>{ph}:</td><td colspan='2' style='color:#ffcc00;'>[Náhodný výběr BONUS pro každou verzi]</td></tr>"
                else:
                    qid = self.selection_map.get(ph)
                    if qid:
                        q = self.owner._find_question_by_id(qid)
                        if q:
                            total_bonus_points += float(q.bonus_correct)
                            min_loss += float(q.bonus_wrong)
                            title_clean = re.sub(r'<[^>]+>', '', q.title)
                            html += f"<tr><td width='100' style='color:#888;'>{ph}:</td><td><b>{title_clean}</b></td><td align='right' style='color:#81c784;'>+{q.bonus_correct} / <span style='color:#e57373;'>{q.bonus_wrong}</span></td></tr>"
                    else:
                        html += f"<tr><td width='100' style='color:#ff5555;'>{ph}:</td><td colspan='2' style='color:#ff5555;'>--- NEVYPLNĚNO ---</td></tr>"
            html += "</table>"

            # 3. Klasifikace - LOGIKA PRO ZOBRAZENÍ
            html += f"<h3 style='background-color: {sec_s_bg}; padding: 5px; border-left: 4px solid #66bb6a;'>3. Klasifikace</h3>"
            
            if is_multi:
                # REŽIM MULTI: Zobrazíme upozornění a obecnou tabulku
                html += """
                <p style='color: #ffcc00; font-weight: bold; border: 1px solid #ffcc00; padding: 5px;'>
                ⚠ UPOZORNĚNÍ: V hromadném exportu jsou bonusové otázky voleny náhodně.<br>
                Hodnoty Max. bodů a Min. bodů (a tím i intervaly známek) se budou lišit pro každou variantu.<br>
                Tabulka níže je pouze orientační pro základ (10 bodů).</p>
                """
                max_txt = "10 + (bonus)"
                min_txt = "(variabilní)"
                val_A_top = "Max"
                val_F_bot = "Min"
                # Pro tabulku použijeme base 10, aby se vygenerovala čísla, ale horní mez označíme textem
                max_body_val = 10.0 
            else:
                # REŽIM SINGLE: Standardní výpočet
                max_body_val = 10.0 + total_bonus_points
                max_txt = f"{max_body_val:.2f} (10 + {total_bonus_points:.2f})"
                min_txt = f"{min_loss:.2f}"
                val_A_top = f"{max_body_val:.2f}"
                val_F_bot = f"{min_loss:.2f}"

            html += f"""
            <p><b>Max. bodů:</b> {max_txt} &nbsp;&nbsp;|&nbsp;&nbsp; <b>Min. bodů (penalizace):</b> {min_txt}</p>
            <table width='60%' border='1' cellspacing='0' cellpadding='5' style='border-collapse: collapse; border: 1px solid {border_color}; color: {text_color};'>
                <tr style='background-color: #333;'><th>Známka</th><th>Interval</th></tr>
                <tr><td align='center' style='color:#81c784'><b>A</b></td><td>&lt; 9.2 ; <b>{val_A_top}</b> &gt;</td></tr>
                <tr><td align='center' style='color:#a5d6a7'><b>B</b></td><td>&lt; 8.4 ; 9.2 )</td></tr>
                <tr><td align='center' style='color:#c8e6c9'><b>C</b></td><td>&lt; 7.6 ; 8.4 )</td></tr>
                <tr><td align='center' style='color:#fff59d'><b>D</b></td><td>&lt; 6.8 ; 7.6 )</td></tr>
                <tr><td align='center' style='color:#ffcc80'><b>E</b></td><td>&lt; 6.0 ; 6.8 )</td></tr>
                <tr><td align='center' style='color:#ef5350'><b>F</b></td><td>&lt; <b>{val_F_bot}</b> ; 6.0 )</td></tr>
            </table>
            """
            html += "</body></html>"
            self.preview_edit.setHtml(html)
            
        except Exception as e:
            print(f"CRITICAL ERROR in _init_page3: {e}")
            import traceback
            traceback.print_exc()
            self.preview_edit.setText(f"Chyba při generování náhledu: {e}")


    def _round_dt_to_10m(self, dt: QDateTime) -> QDateTime:
        m = dt.time().minute()
        rounded = m % 10
        if rounded < 5:
            m -= rounded
        else:
            m += (10 - rounded)
        
        # QDateTime je immutable, musíme vytvořit nový
        # Jednodušší hack: převést na python datetime a zpět, nebo manipulovat s časem
        # Ale Qt metody pro addSecs jsou nejlepší
        # Zde původní implementace (dle kontextu) asi vracela upravený QDateTime
        # Uděláme to robustně:
        
        time = dt.time()
        total_minutes = time.hour() * 60 + time.minute()
        
        # Zaokrouhlení na 10 minut
        remainder = total_minutes % 10
        if remainder >= 5:
            total_minutes += (10 - remainder)
        else:
            total_minutes -= remainder
            
        # Ošetření přetečení dne (24:00)
        # Pro jednoduchost v rámci dnešního dne:
        new_h = (total_minutes // 60) % 24
        new_m = total_minutes % 60
        
        new_time = QTime(new_h, new_m)
        return QDateTime(dt.date(), new_time)

    def _cz_day_of_week(self, dt: QDateTime) -> str:
        # dt.date().dayOfWeek() vrací 1 (Mon) až 7 (Sun)
        days = ["pondělí", "úterý", "středa", "čtvrtek", "pátek", "sobota", "neděle"]
        idx = dt.date().dayOfWeek() - 1
        return days[idx]

    def accept(self) -> None:
        self._save_settings()
        
        if not self.template_path or not self.output_path:
            return

        # Kontrolní Hash
        k_hash = getattr(self, "_cached_hash", "")
        if not k_hash:
            ts = str(datetime.now().timestamp())
            salt = secrets.token_hex(16)
            data_to_hash = f"{ts}{salt}"
            k_hash = hashlib.sha3_256(data_to_hash.encode("utf-8")).hexdigest()

        is_multi = (self.mode_group.checkedId() == 1)
        count = self.spin_multi_count.value() if is_multi else 1
        do_pdf_export = self.chk_export_pdf.isChecked()
        
        # --- PŘÍPRAVA POOLU OTÁZEK ---
        question_pool = []
        # --- NOVÉ: Pool pro bonusové otázky ---
        question_pool_bonus = []
        
        if is_multi:
            def collect_questions(group_id, is_subgroup, target_type="classic"):
                qs = []
                nodes_to_visit = list(self.owner.root.groups)
                target_node = None
                while nodes_to_visit:
                    curr = nodes_to_visit.pop(0)
                    if curr.id == group_id:
                        target_node = curr
                        break
                    if hasattr(curr, "subgroups") and curr.subgroups:
                        nodes_to_visit.extend(curr.subgroups)
                
                if target_node:
                    def extract_q(node):
                        valid_qs = []
                        if hasattr(node, "questions"):
                            valid_qs.extend([q.id for q in node.questions if q.type == target_type])
                        if hasattr(node, "subgroups") and node.subgroups:
                            for sub in node.subgroups:
                                valid_qs.extend(extract_q(sub))
                        return valid_qs
                    qs = extract_q(target_node)
                return qs

            # 1. Zdroje pro KLASICKÉ otázky
            sources_to_process = self.multi_selected_sources
            if not sources_to_process:
                sources_to_process = [{"id": g.id, "type": "group"} for g in self.owner.root.groups]

            for source in sources_to_process:
                sid = source["id"]
                stype = source["type"]
                is_sub = (stype == "subgroup")
                pool_ids = collect_questions(sid, is_sub, "classic")
                question_pool.extend(pool_ids)
            
            question_pool = list(set(question_pool))
            
            # 2. Zdroje pro BONUSOVÉ otázky
            # Pokud uživatel vybral konkrétní, použijeme ty. Jinak sebereme VŠECHNY dostupné bonusy.
            if self.multi_selected_bonus_ids:
                question_pool_bonus = list(self.multi_selected_bonus_ids)
            else:
                # Pokud není vybráno, vezmeme bonusy ze všech skupin
                all_groups = [{"id": g.id, "type": "group"} for g in self.owner.root.groups]
                for g_src in all_groups:
                    b_ids = collect_questions(g_src["id"], False, "bonus")
                    question_pool_bonus.extend(b_ids)
                question_pool_bonus = list(set(question_pool_bonus))

        base_output_path = self.output_path
        success_count = 0
        generated_docx_files = []
        generated_pdf_files = []

        print_folder = self.print_dir
        if do_pdf_export:
            print_folder.mkdir(parents=True, exist_ok=True)

        # --- UI SETUP ---
        # Zablokujeme tlačítka, aby uživatel neklikal znovu
        self.button(QWizard.FinishButton).setEnabled(False)
        self.button(QWizard.BackButton).setEnabled(False)
        
        if hasattr(self, "progress_bar"):
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, count)
            self.progress_bar.setValue(0)
            self.lbl_status_final.setText("Generuji DOCX soubory...")
            QApplication.processEvents()

        # --- LOOP GENEROVÁNÍ ---
        try:
            for i in range(count):
                current_selection = self.selection_map.copy()
                
                if is_multi and question_pool:
                    import random
                    # Pro multi režim ignorujeme ruční výběr
                    current_selection = {} 
                    
                    # A. Klasické otázky
                    targets = self.placeholders_q
                    needed = len(targets)
                    
                    if len(question_pool) >= needed:
                        picked = random.sample(question_pool, needed)
                        for idx, ph in enumerate(targets):
                            current_selection[ph] = picked[idx]
                    else:
                        if len(question_pool) > 0:
                            for ph in targets:
                                current_selection[ph] = random.choice(question_pool)
                    
                    # B. Bonusové otázky
                    targets_b = self.placeholders_b
                    needed_b = len(targets_b)
                    
                    if len(question_pool_bonus) >= needed_b:
                        picked_b = random.sample(question_pool_bonus, needed_b)
                        for idx, ph in enumerate(targets_b):
                            current_selection[ph] = picked_b[idx]
                    else:
                        if len(question_pool_bonus) > 0:
                             for idx, ph in enumerate(targets_b):
                                 # Pokud nemáme dost unikátních, musíme opakovat, nebo vzít co je
                                 # Zde fallback: pokud dojdou, bereme random s opakováním
                                 current_selection[ph] = random.choice(question_pool_bonus)

                
                repl_plain: Dict[str, str] = {}
                
                # --- Datum a Čas (Inline logika) ---
                raw_dt = self.dt_edit.dateTime()
                # Zaokrouhlení na 10 min
                time = raw_dt.time()
                total_minutes = time.hour() * 60 + time.minute()
                remainder = total_minutes % 10
                if remainder >= 5: total_minutes += (10 - remainder)
                else: total_minutes -= remainder
                new_h = (total_minutes // 60) % 24
                new_m = total_minutes % 60
                rounded_time = QTime(new_h, new_m)
                rounded_dt = QDateTime(raw_dt.date(), rounded_time)
                
                # Den v týdnu
                days_cz = ["pondělí", "úterý", "středa", "čtvrtek", "pátek", "sobota", "neděle"]
                day_idx = rounded_dt.date().dayOfWeek() - 1
                day_str = days_cz[day_idx]
                
                dt_str = f"{day_str} {rounded_dt.toString('dd.MM.yyyy HH:mm')}"
                repl_plain["DatumČas"] = dt_str
                repl_plain["DatumCas"] = dt_str
                repl_plain["DATUMCAS"] = dt_str
                
                prefix = self.le_prefix.text().strip()
                today = datetime.now().strftime("%Y-%m-%d")
                version_suffix = f"-{i+1}" if is_multi else ""
                verze_str = f"{prefix}{version_suffix} {today}"
                
                repl_plain["PoznamkaVerze"] = verze_str
                repl_plain["POZNAMKAVERZE"] = verze_str
                repl_plain["KontrolniHash"] = k_hash
                repl_plain["KONTROLNIHASH"] = k_hash
                
                total_bonus = 0.0
                min_loss = 0.0
                for qid in current_selection.values():
                    q = self.owner._find_question_by_id(qid)
                    if not q: continue
                    if q.type == 'bonus':
                        total_bonus += float(q.bonus_correct)
                        min_loss += float(q.bonus_wrong)
                
                # Base points - zde fixně 10, nebo spočítat z klasických
                max_body = 10.0 + total_bonus
                
                repl_plain["MaxBody"] = f"{max_body:.2f}"
                repl_plain["MAXBODY"] = f"{max_body:.2f}"
                repl_plain["MinBody"] = f"{min_loss:.2f}"
                repl_plain["MINBODY"] = f"{min_loss:.2f}"

                # --- ÚPRAVA: Podpora obrázků (tuple) ---
                rich_map: Dict[str, object] = {}
                for ph, qid in current_selection.items():
                    q = self.owner._find_question_by_id(qid)
                    if q:
                        img_path = getattr(q, "image_path", None)
                        rich_map[ph] = (q.text_html, img_path, float(getattr(q, "image_width_cm", 0.0) or 0.0), float(getattr(q, "image_height_cm", 0.0) or 0.0))
                    else:
                        rich_map[ph] = ("", None)
                # ----------------------------------------

                if is_multi:
                    p = Path(base_output_path)
                    new_name = f"{p.stem}_v{i+1}{p.suffix}"
                    target_path = p.parent / new_name
                else:
                    target_path = base_output_path

                try:
                    self.owner._generate_docx_from_template(self.template_path, target_path, repl_plain, rich_map)
                    success_count += 1
                    generated_docx_files.append(target_path)
                except Exception as e:
                    QMessageBox.critical(self, "Export", f"Chyba při exportu verze {i+1}:\n{e}")
                    if not is_multi: 
                        # Povolit tlačítka zpět
                        self.button(QWizard.FinishButton).setEnabled(True)
                        self.button(QWizard.BackButton).setEnabled(True)
                        return
                
                # Update Progress
                if hasattr(self, "progress_bar"):
                    self.progress_bar.setValue(i + 1)
                    self.lbl_status_final.setText(f"Generuji DOCX {i+1}/{count}...")
                    QApplication.processEvents()

            # Historie
            if is_multi:
                record_name = f"Balík {count} verzí: {base_output_path.name}"
                self.owner.register_export(record_name, k_hash)
            else:
                self.owner.register_export(base_output_path.name, k_hash)

            # PDF Export
            pdf_success_msg = ""
            if do_pdf_export and generated_docx_files:
                # Reset progress bar
                if hasattr(self, "progress_bar"):
                    self.progress_bar.setRange(0, len(generated_docx_files))
                    self.progress_bar.setValue(0)
                    self.lbl_status_final.setText("Převádím DOCX na PDF...")
                    QApplication.processEvents()
                    
                try:
                    for idx, docx_file in enumerate(generated_docx_files):
                        if hasattr(self, "lbl_status_final"):
                            self.lbl_status_final.setText(f"PDF Konverze: {docx_file.name}")
                            QApplication.processEvents()
                            
                        pdf_file = self.owner._convert_docx_to_pdf(docx_file)
                        if pdf_file and pdf_file.exists():
                            generated_pdf_files.append(pdf_file)
                        
                        if hasattr(self, "progress_bar"):
                            self.progress_bar.setValue(idx + 1)

                    if generated_pdf_files:
                        if is_multi and len(generated_pdf_files) > 1:
                            if hasattr(self, "lbl_status_final"):
                                self.lbl_status_final.setText("Slučuji PDF soubory...")
                                self.progress_bar.setRange(0, 0) 
                                QApplication.processEvents()
                                
                            merged_name = f"{base_output_path.stem}_merged.pdf"
                            final_pdf = print_folder / merged_name
                            
                            if self.owner._merge_pdfs(generated_pdf_files, final_pdf, cleanup=True):
                                pdf_success_msg = f"\n\nPDF pro tisk (sloučené) uloženo do:\n{final_pdf}"
                            else:
                                pdf_success_msg = f"\n\nPOZOR: Slučování selhalo. Jednotlivá PDF jsou v:\n{print_folder}"
                                import shutil
                                for tmp_pdf in generated_pdf_files:
                                    if tmp_pdf.exists():
                                        try:
                                            dest = print_folder / tmp_pdf.name
                                            shutil.move(str(tmp_pdf), str(dest))
                                        except: pass
                        else:
                            import shutil
                            final_pdf = print_folder / generated_pdf_files[0].name
                            shutil.move(str(generated_pdf_files[0]), str(final_pdf))
                            pdf_success_msg = f"\n\nPDF pro tisk uloženo do:\n{final_pdf}"
                                
                except Exception as e:
                    QMessageBox.warning(self, "PDF Export", f"Kritická chyba při exportu PDF:\n{e}")
                    import traceback; traceback.print_exc()
            
            if hasattr(self, "progress_bar"):
                self.progress_bar.setVisible(False)
                self.lbl_status_final.setText("Hotovo.")

            if is_multi:
                msg = f"Hromadný export dokončen.\nVygenerováno {success_count} souborů DOCX.{pdf_success_msg}"
                QMessageBox.information(self, "Export", msg)
            else:
                msg = f"Export dokončen.\nSoubor uložen:\n{base_output_path}{pdf_success_msg}"
                QMessageBox.information(self, "Export", msg)
                
            super().accept()
        
        except Exception as e:
            # Pokud nastane chyba, odblokujeme tlačítka
            self.button(QWizard.FinishButton).setEnabled(True)
            self.button(QWizard.BackButton).setEnabled(True)
            QMessageBox.critical(self, "Kritická chyba", f"Neočekávaná chyba:\n{e}")

# --------------------------- Hlavní okno (UI + logika) ---------------------------

class FunnyAnswerDialog(QDialog):
    """Dialog pro přidání nové vtipné odpovědi."""
    def __init__(self, parent=None, project_root: Optional[Path] = None) -> None:
        super().__init__(parent)
        self.setWindowTitle("Přidat vtipnou odpověď")
        self.resize(600, 400)
        
        self.project_root = project_root
        
        layout = QVBoxLayout(self)
        form = QFormLayout()
        
        # Výběr zdrojové písemky
        self.combo_source = QComboBox()
        self.combo_source.addItem("(Ruční zadání / Bez vazby na soubor)", None)
        self.combo_source.currentIndexChanged.connect(self._on_source_changed)
        
        self.text_edit = QTextEdit()
        self.text_edit.setPlaceholderText("Znění vtipné odpovědi...")
        
        self.author_edit = QLineEdit()
        self.author_edit.setPlaceholderText("Např. Student, Anonym...")
        
        self.date_edit = QDateTimeEdit(QDateTime.currentDateTime())
        self.date_edit.setDisplayFormat("dd.MM.yyyy HH:mm") # Přidán i čas pro kontrolu
        self.date_edit.setCalendarPopup(True)

        form.addRow("Zdroj písemky:", self.combo_source)
        form.addRow("Odpověď:", self.text_edit)
        form.addRow("Autor:", self.author_edit)
        form.addRow("Datum:", self.date_edit)
        
        layout.addLayout(form)
        
        bb = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        bb.accepted.connect(self.accept)
        bb.rejected.connect(self.reject)
        layout.addWidget(bb)
        
        # Načtení souborů
        self._load_files()

    def _load_files(self) -> None:
        if not self.project_root:
            return
            
        base_dir = self.project_root / "data" / "Vygenerované testy"
        if not base_dir.exists():
            return
            
        # Rekurzivní vyhledání všech .docx
        found_files = sorted(base_dir.rglob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
        
        for p in found_files:
            # Zobrazíme relativní cestu vůči složce vygenerovaných testů pro přehlednost
            try:
                rel_path = p.relative_to(base_dir)
                display_text = str(rel_path)
            except ValueError:
                display_text = p.name
            
            self.combo_source.addItem(display_text, str(p))

    def _on_source_changed(self, index: int) -> None:
        data = self.combo_source.itemData(index)
        if not data:
            return
            
        path = Path(data)
        filename = path.name
        
        # Očekávaný formát: Prefix_YYYY-MM-DD_HHMM_...
        parts = filename.split('_')
        
        if len(parts) >= 3:
            date_str = parts[1]  # YYYY-MM-DD
            time_str = parts[2]  # HHMM
            
            # Validace a parsování
            try:
                # Datum
                qdate = QDateTime.fromString(date_str, "yyyy-MM-dd").date()
                
                # Čas (HHMM)
                qtime = QDateTime.fromString(time_str, "HHmm").time()
                
                if qdate.isValid() and qtime.isValid():
                    dt = QDateTime(qdate, qtime)
                    self.date_edit.setDateTime(dt)
            except Exception:
                # Pokud parsování selže, neděláme nic (necháme aktuální)
                pass
            
    def set_data(self, text: str, date_str: str, author: str, source_doc: Optional[str] = None) -> None:
        """Naplní formulář daty pro editaci."""
        self.text_edit.setText(text)
        self.author_edit.setText(author)

        # Pokusíme se parsovat datum (očekáváme dd.MM.yyyy nebo dd.MM.yyyy HH:mm)
        dt = QDateTime.fromString(date_str, "dd.MM.yyyy HH:mm")
        if not dt.isValid():
            dt = QDateTime.fromString(date_str, "dd.MM.yyyy")

        if dt.isValid():
            self.date_edit.setDateTime(dt)

        # Nastavení zdrojového dokumentu v comboboxu
        if source_doc:
            idx = self.combo_source.findData(source_doc)
            if idx >= 0:
                self.combo_source.setCurrentIndex(idx)
            else:
                self.combo_source.setCurrentIndex(0)
        else:
            self.combo_source.setCurrentIndex(0)
            
    def get_data(self) -> tuple[str, str, str, str]:
        # Vracíme text, datum, autora a (případně) cestu ke zdrojovému dokumentu
        text = self.text_edit.toPlainText().strip()
        date_str = self.date_edit.dateTime().toString("dd.MM.yyyy HH:mm")
        author = self.author_edit.text().strip()
        data = self.combo_source.currentData()
        source_doc = str(data) if data is not None else ""
        return text, date_str, author, source_doc

class MainWindow(QMainWindow):
    """Hlavní okno aplikace."""

    def _selected_question_ids(self) -> List[str]:
        ids: List[str] = []
        for it in self.tree.selectedItems():
            meta = it.data(0, Qt.UserRole) or {}
            if meta.get("kind") == "question":
                ids.append(meta.get("id"))
        return ids

    def _reselect_questions(self, ids: List[str]) -> None:
        if not ids:
            return
        wanted = set(ids)
        self.tree.clearSelection()

        def walk(item: QTreeWidgetItem):
            meta = item.data(0, Qt.UserRole) or {}
            if meta.get("kind") == "question" and meta.get("id") in wanted:
                item.setSelected(True)
            for i in range(item.childCount()):
                walk(item.child(i))

        for i in range(self.tree.topLevelItemCount()):
            walk(self.tree.topLevelItem(i))

        items = self.tree.selectedItems()
        if items:
            self.tree.setCurrentItem(items[0])
            self._on_tree_selection_changed()

    def __init__(self, data_path: Optional[Path] = None) -> None:
        super().__init__()
        # ZMĚNA: Titulek s verzí
        self.setWindowTitle(APP_NAME)
        self.resize(1800, 1600)
    
        self.project_root = Path.cwd()
        default_data_dir = self.project_root / "data"
        default_data_dir.mkdir(parents=True, exist_ok=True)
        self.data_path = data_path or (default_data_dir / "questions.json")
    
        self.images_dir = default_data_dir / "obrázky"
        self.images_dir.mkdir(parents=True, exist_ok=True)
    
        # Aplikace ikona (pokud existuje)
        icon_file = self.project_root / "icon" / "icon.png"
        if icon_file.exists():
            app_icon = QIcon(str(icon_file))
            self.setWindowIcon(app_icon)
            QApplication.instance().setWindowIcon(app_icon)
    
        self.root: RootData = RootData(groups=[])
    
        # NOVÉ: aby koš fungoval i když RootData nemá trash jako pole (minimal-change)
        if not hasattr(self.root, "trash") or not isinstance(getattr(self.root, "trash", None), list):
            self.root.trash = []
    
        self._current_question_id: Optional[str] = None
        self._current_node_kind: Optional[str] = None
    
        self._autosave_timer = QTimer(self)
        self._autosave_timer.setSingleShot(True)
        self._autosave_timer.setInterval(1200)
        self._autosave_timer.timeout.connect(self._autosave_current_question)
    
        self._build_ui()
        self._connect_signals()
        
        self.load_data()
    
        # NOVÉ: po load znovu zajistit koš (kdyby JSON byl starší bez trash)
        if not hasattr(self.root, "trash") or not isinstance(getattr(self.root, "trash", None), list):
            self.root.trash = []
    
        self._refresh_tree()
        self._refresh_funny_answers_tab()
    
        # NOVÉ: refresh koše po načtení
        self._refresh_trash_table()
    
        # ZMĚNA: Strom 60%, Editor 40% (cca 840px : 560px)
        self.splitter.setSizes([940, 860])

    def _init_trash_tab(self) -> None:
        self.tab_trash = QWidget()
        trash_layout = QVBoxLayout(self.tab_trash)
        trash_layout.setContentsMargins(4, 4, 4, 4)
        trash_layout.setSpacing(6)
    
        self.table_trash = QTableWidget(0, 5)
        self.table_trash.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table_trash.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table_trash.setSortingEnabled(True)
        self.table_trash.verticalHeader().setVisible(False)
    
        self.table_trash.setHorizontalHeaderLabels([
            "NÁZEV OTÁZKY",
            "TYP",
            "SMAZÁNO",
            "PŮVODNÍ SKUPINA",
            "PŮVODNÍ PODSKUPINA",
        ])
        header = self.table_trash.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(3, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(4, QHeaderView.ResizeToContents)
    
        trash_layout.addWidget(self.table_trash, 1)
    
        self.trash_detail = QTextEdit()
        self.trash_detail.setReadOnly(True)
        self.trash_detail.setPlaceholderText("Vyber otázku v koši pro zobrazení detailu…")
        self.trash_detail.setFixedHeight(220)
        trash_layout.addWidget(self.trash_detail)
    
        btns = QHBoxLayout()
        self.btn_trash_restore = QPushButton("Obnovit")
        self.btn_trash_delete = QPushButton("Trvale smazat")
        self.btn_trash_empty = QPushButton("Vysypat koš")
    
        self.btn_trash_delete.setStyleSheet("background-color: #d32f2f; color: white; font-weight: bold; padding: 4px 8px;")
        self.btn_trash_empty.setStyleSheet("background-color: #d32f2f; color: white; font-weight: bold; padding: 4px 8px;")
    
        btns.addWidget(self.btn_trash_restore)
        btns.addWidget(self.btn_trash_delete)
        btns.addStretch()
        btns.addWidget(self.btn_trash_empty)
        trash_layout.addLayout(btns)
    
        self.left_tabs.addTab(self.tab_trash, "Koš")
    
        self.btn_trash_restore.setEnabled(False)
        self.btn_trash_delete.setEnabled(False)
        self.btn_trash_empty.setEnabled(False)

    def _duplicate_question(self) -> None:
        kind, meta = self._selected_node()
        if kind != "question":
            return

        gid = meta.get("parent_group_id")
        sgid = meta.get("parent_subgroup_id")
        qid = meta.get("id")

        q_orig = self._find_question(gid, sgid, qid)
        if not q_orig:
            return

        # Vytvoření kopie
        data = asdict(q_orig)
        data["id"] = str(_uuid.uuid4())
        data["title"] = (q_orig.title or "Otázka") + " (kopie)"
        
        new_q = Question(**data)

        # Vložení do správné podskupiny
        target_sg = self._find_subgroup(gid, sgid)
        if target_sg:
            target_sg.questions.append(new_q)
            self._refresh_tree()
            self._select_question(new_q.id)
            self.save_data()
            self.statusBar().showMessage("Otázka byla duplikována.", 3000)

    def _build_ui(self) -> None:
        # Hlavní container
        main_widget = QWidget()
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
    
        self.splitter = QSplitter()
        self.splitter.setChildrenCollapsible(False)
        self.splitter.setHandleWidth(8)
    
        # LEVÝ PANEL
        left_panel_container = QWidget()
        left_container_layout = QVBoxLayout(left_panel_container)
        left_container_layout.setContentsMargins(0, 0, 0, 0)
        self.left_tabs = QTabWidget()
    
        # ZÁLOŽKA 1: OTÁZKY
        self.tab_questions = QWidget()
        questions_layout = QVBoxLayout(self.tab_questions)
        questions_layout.setContentsMargins(4, 4, 4, 4)
        questions_layout.setSpacing(6)
    
        filter_bar = QWidget()
        filter_layout = QHBoxLayout(filter_bar)
        filter_layout.setContentsMargins(0, 0, 0, 0)
        filter_layout.setSpacing(6)
        self.filter_edit = QLineEdit()
        self.filter_edit.setPlaceholderText("Filtr: název / obsah otázky…")
        self.btn_move_selected = QPushButton("Přesunout vybrané…")
        self.btn_delete_selected = QPushButton("Smazat vybrané")
        filter_layout.addWidget(self.filter_edit, 1)
        filter_layout.addWidget(self.btn_move_selected)
        filter_layout.addWidget(self.btn_delete_selected)
        questions_layout.addWidget(filter_bar)
    
        self.tree = DnDTree(self)
        self.tree.setContextMenuPolicy(Qt.CustomContextMenu)
        self.tree.customContextMenuRequested.connect(self._on_tree_context_menu)
        questions_layout.addWidget(self.tree, 1)
    
        # Legenda
        legend_box = QFrame()
        legend_box.setStyleSheet("background-color: #2d2d2d; border-radius: 4px;")
        legend_layout = QHBoxLayout(legend_box)
        legend_layout.setContentsMargins(8, 4, 8, 4)
        legend_layout.setSpacing(15)
    
        def add_legend_item(text, color_hex):
            lbl = QLabel(f"<span style='color:{color_hex}; font-size:14px;'>■</span> <span style='color:#cccccc;'>{text}</span>")
            lbl.setTextFormat(Qt.RichText)
            lbl.setStyleSheet("border: none; background: transparent;")
            legend_layout.addWidget(lbl)
    
        add_legend_item("Skupina", "#ff5252")
        add_legend_item("Podskupina", "#ff8a80")
        add_legend_item("Klasická", "#42a5f5")
        add_legend_item("BONUS", "#ffea00")
        legend_layout.addStretch()
    
        questions_layout.addWidget(legend_box)
    
        self.left_tabs.addTab(self.tab_questions, "Otázky")
    
        # ZÁLOŽKA 2: HISTORIE
        self.tab_history = QWidget()
        history_layout = QVBoxLayout(self.tab_history)
        history_layout.setContentsMargins(4, 4, 4, 4)
    
        self.table_history = QTableWidget(0, 2)
        self.table_history.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table_history.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table_history.setSortingEnabled(True)
        self.table_history.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table_history.customContextMenuRequested.connect(self._on_history_context_menu)
        history_layout.addWidget(self.table_history)
    
        h_btns = QHBoxLayout()
        btn_refresh_hist = QPushButton("Obnovit historii")
        btn_refresh_hist.clicked.connect(self._refresh_history_table)
    
        btn_clear_hist = QPushButton("Vymazat celou historii")
        btn_clear_hist.setStyleSheet("background-color: #d32f2f; color: white; font-weight: bold; padding: 4px 8px;")
        btn_clear_hist.clicked.connect(self._clear_all_history)
    
        h_btns.addWidget(btn_refresh_hist)
        h_btns.addStretch()
        h_btns.addWidget(btn_clear_hist)
        history_layout.addLayout(h_btns)
    
        self.left_tabs.addTab(self.tab_history, "Historie")
    
        # NOVÉ: ZÁLOŽKA 3: KOŠ
        self._init_trash_tab()
    
        self._init_funny_answers_tab()
        left_container_layout.addWidget(self.left_tabs)
    
        # PRAVÝ PANEL
        self.detail_stack = QWidget()
        self.detail_layout = QVBoxLayout(self.detail_stack)
        self.detail_layout.setContentsMargins(6, 6, 6, 6)
        self.detail_layout.setSpacing(8)
    
        self.editor_toolbar = QToolBar("Formát")
        self.editor_toolbar.setIconSize(QSize(18, 18))
    
        # --- STYLING EDITOR TOOLBARU ---
        self.editor_toolbar.setStyleSheet("""
            QToolBar {
                border-bottom: 1px solid #3e3e3e;
                background-color: #2d2d2d;
                spacing: 4px;
            }
            QToolButton {
                background-color: #383838;
                border: 1px solid #505050;
                border-radius: 3px;
                padding: 2px 4px;
                color: #e0e0e0;
            }
            QToolButton:checked {
                background-color: #4a90e2;
                border-color: #4a90e2;
                color: white;
            }
            QToolButton:hover {
                background-color: #454545;
                border-color: #707070;
            }
            QToolButton:pressed {
                background-color: #252525;
            }
            QToolBar::separator {
                background: #555;
                width: 1px;
                margin: 4px 4px;
            }
        """)
    
        self.action_bold = QAction("Tučné", self); self.action_bold.setCheckable(True); self.action_bold.setShortcut(QKeySequence.Bold)
        self.action_italic = QAction("Kurzíva", self); self.action_italic.setCheckable(True); self.action_italic.setShortcut(QKeySequence.Italic)
        self.action_underline = QAction("Podtržení", self); self.action_underline.setCheckable(True); self.action_underline.setShortcut(QKeySequence.Underline)
        self.action_color = QAction("Barva", self)
        self.action_bullets = QAction("Odrážky", self); self.action_bullets.setCheckable(True)
        self.action_indent_dec = QAction("< Odsadit", self); self.action_indent_dec.setToolTip("Zmenšit odsazení")
        self.action_indent_inc = QAction("> Odsadit", self); self.action_indent_inc.setToolTip("Zvětšit odsazení")
        self.action_align_left = QAction("Vlevo", self)
        self.action_align_center = QAction("Na střed", self)
        self.action_align_right = QAction("Vpravo", self)
        self.action_align_justify = QAction("Do bloku", self)
        self.align_group = QActionGroup(self)
        for a in (self.action_align_left, self.action_align_center, self.action_align_right, self.action_align_justify):
            a.setCheckable(True); self.align_group.addAction(a)
        self.editor_toolbar.addAction(self.action_bold)
        self.editor_toolbar.addAction(self.action_italic)
        self.editor_toolbar.addAction(self.action_underline)
        self.editor_toolbar.addSeparator()
        self.editor_toolbar.addAction(self.action_color)
        self.editor_toolbar.addSeparator()
        self.editor_toolbar.addAction(self.action_bullets)
        self.editor_toolbar.addSeparator()
        self.editor_toolbar.addAction(self.action_indent_dec)
        self.editor_toolbar.addAction(self.action_indent_inc)
        self.editor_toolbar.addSeparator()
        self.editor_toolbar.addAction(self.action_align_left)
        self.editor_toolbar.addAction(self.action_align_center)
        self.editor_toolbar.addAction(self.action_align_right)
        self.editor_toolbar.addAction(self.action_align_justify)
    
        self.form_layout = QFormLayout()
        self.form_layout.setLabelAlignment(Qt.AlignLeft)
        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText("Krátký název otázky…")
        self.combo_type = QComboBox(); self.combo_type.addItems(["Klasická", "BONUS"])
        self.spin_points = QSpinBox(); self.spin_points.setRange(-999, 999); self.spin_points.setValue(1)
        self.spin_bonus_correct = QDoubleSpinBox(); self.spin_bonus_correct.setDecimals(2); self.spin_bonus_correct.setSingleStep(0.01); self.spin_bonus_correct.setRange(-999.99, 999.99); self.spin_bonus_correct.setValue(1.00)
        self.spin_bonus_wrong = QDoubleSpinBox(); self.spin_bonus_wrong.setDecimals(2); self.spin_bonus_wrong.setSingleStep(0.01); self.spin_bonus_wrong.setRange(-999.99, 999.99); self.spin_bonus_wrong.setValue(0.00)
    
        self.image_path_edit = QLineEdit()
        self.image_path_edit.setPlaceholderText("Cesta k obrázku (volitelné)…")
        self.btn_choose_image = QPushButton("Vybrat…")
        self.btn_clear_image = QPushButton("Smazat")
        img_row = QWidget()
        img_row_l = QHBoxLayout(img_row)
        img_row_l.setContentsMargins(0, 0, 0, 0)
        img_row_l.setSpacing(6)
        img_row_l.addWidget(self.image_path_edit, 1)
        img_row_l.addWidget(self.btn_choose_image)
        img_row_l.addWidget(self.btn_clear_image)
    
        self.form_layout.addRow("Název otázky:", self.title_edit)
        self.form_layout.addRow("Typ otázky:", self.combo_type)
        self.form_layout.addRow("Obrázek:", img_row)
    
        # Aktuální velikost obrázku (cm) – pouze informativní (dle DPI v souboru / fallback)
        self.lbl_img_actual_size_label = QLabel("Aktuální velikost (cm):")
        self.lbl_img_actual_size = QLabel("")
        self.lbl_img_actual_size.setTextInteractionFlags(Qt.TextSelectableByMouse)
        self.lbl_img_actual_size.setStyleSheet("color: #cfcfcf;")
        self.form_layout.addRow(self.lbl_img_actual_size_label, self.lbl_img_actual_size)
    
        # Rozměry obrázku v DOCX (cm) – použije se při exportu
        self.spin_img_w_cm = QDoubleSpinBox()
        self.spin_img_w_cm.setRange(0.0, 200.0)
        self.spin_img_w_cm.setDecimals(2)
        self.spin_img_w_cm.setSingleStep(0.10)
        self.spin_img_w_cm.setSuffix(" cm")
        self.spin_img_w_cm.setValue(14.00)
    
        self.spin_img_h_cm = QDoubleSpinBox()
        self.spin_img_h_cm.setRange(0.0, 200.0)
        self.spin_img_h_cm.setDecimals(2)
        self.spin_img_h_cm.setSingleStep(0.10)
        self.spin_img_h_cm.setSuffix(" cm")
        self.spin_img_h_cm.setValue(0.00)
    
        img_size_row = QWidget()
        img_size_l = QHBoxLayout(img_size_row)
        img_size_l.setContentsMargins(0, 0, 0, 0)
        img_size_l.setSpacing(6)
        img_size_l.addWidget(QLabel("Šířka:"), 0)
        img_size_l.addWidget(self.spin_img_w_cm, 0)
        img_size_l.addSpacing(12)
        img_size_l.addWidget(QLabel("Výška:"), 0)
        img_size_l.addWidget(self.spin_img_h_cm, 0)
        img_size_l.addStretch(1)
    
        # defaultně deaktivované – aktivuje se jen když otázka má obrázek
        self.spin_img_w_cm.setEnabled(False)
        self.spin_img_h_cm.setEnabled(False)
    
        self.lbl_img_export_size = QLabel("Velikost pro export (cm):")
        self.img_size_row = img_size_row
        self.form_layout.addRow(self.lbl_img_export_size, self.img_size_row)
    
        self.chk_img_keep_aspect = QCheckBox("Zachovat poměr stran")
        self.chk_img_keep_aspect.setChecked(True)
        self.chk_img_keep_aspect.setEnabled(False)
        self.chk_img_keep_aspect.setVisible(False)
        self.form_layout.addRow(self.chk_img_keep_aspect)
    
        # Interní pomocné flagy pro synchronizaci rozměrů obrázku (aby nevznikaly smyčky signálů)
        self._img_size_sync_block = False
        self._img_ratio_hw = None  # height/width poměr z pixelů aktuálního obrázku
    
        # defaultně skryté – zobrazí se jen když otázka má obrázek
        self.lbl_img_actual_size_label.setVisible(False)
        self.lbl_img_actual_size.setVisible(False)
        self.lbl_img_export_size.setVisible(False)
        self.img_size_row.setVisible(False)
    
        self.form_layout.addRow("Body (klasická):", self.spin_points)
        self.form_layout.addRow("Body za správně (BONUS):", self.spin_bonus_correct)
        self.form_layout.addRow("Body za špatně (BONUS):", self.spin_bonus_wrong)
    
        self.edit_correct_answer = QTextEdit()
        self.edit_correct_answer.setPlaceholderText("Volitelný text správné odpovědi...")
        self.edit_correct_answer.setFixedHeight(60)
    
        self.funny_container = QWidget()
        fc_layout = QVBoxLayout(self.funny_container)
        fc_layout.setContentsMargins(0,0,0,0)
        self.table_funny = QTableWidget(0, 4)
        self.table_funny.setHorizontalHeaderLabels(["Odpověď", "Datum", "Jméno", "Zdroj"])
        self.table_funny.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.table_funny.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.table_funny.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.table_funny.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.table_funny.setFixedHeight(120)
        self.table_funny.setSelectionBehavior(QAbstractItemView.SelectRows)
        btns_layout = QHBoxLayout()
        self.btn_add_funny = QPushButton("Přidat vtipnou odpoveď")
        self.btn_rem_funny = QPushButton("Odebrat")
        btns_layout.addWidget(self.btn_add_funny)
        btns_layout.addWidget(self.btn_rem_funny)
        btns_layout.addStretch()
        fc_layout.addLayout(btns_layout)
        fc_layout.addWidget(self.table_funny)
    
        self.text_edit = QTextEdit()
        self.text_edit.setAcceptRichText(True)
        self.text_edit.setPlaceholderText("Sem napište znění otázky…\nPodporováno: tučné, kurzíva, podtržení, barva, odrážky, zarovnání.")
        self.text_edit.setMinimumHeight(200)
    
        self.btn_save_question = QPushButton("Uložit změny otázky"); self.btn_save_question.setDefault(True)
    
        self.rename_panel = QWidget()
        rename_layout = QFormLayout(self.rename_panel)
        self.rename_line = QLineEdit()
        self.btn_rename = QPushButton("Uložit název")
        rename_layout.addRow("Název:", self.rename_line)
        rename_layout.addRow(self.btn_rename)
    
        self.lbl_content = QLabel("<b>Obsah otázky:</b>")
        self.lbl_correct = QLabel("<b>Správná odpověď:</b>")
        self.lbl_funny = QLabel("<b>Vtipné odpovědi:</b>")
    
        self.detail_layout.addWidget(self.editor_toolbar)
        self.detail_layout.addLayout(self.form_layout)
        self.detail_layout.addWidget(self.lbl_content)
        self.detail_layout.addWidget(self.text_edit, 1)
        self.detail_layout.addWidget(self.lbl_correct)
        self.detail_layout.addWidget(self.edit_correct_answer)
        self.detail_layout.addWidget(self.lbl_funny)
        self.detail_layout.addWidget(self.funny_container)
        self.detail_layout.addWidget(self.btn_save_question)
        self.detail_layout.addWidget(self.rename_panel)
    
        self._set_editor_enabled(False)
        self.splitter.addWidget(left_panel_container)
        self.splitter.addWidget(self.detail_stack)
        self.splitter.setStretchFactor(1, 1)
        self.setCentralWidget(self.splitter)
    
        # -- TOOLBAR STYLING & MERGE --
    
        # 1. Odstranit starý toolbar "Import/Export"
        for child in self.children():
            if isinstance(child, QToolBar) and child.windowTitle() == "Import/Export":
                self.removeToolBar(child)
                child.deleteLater()
                break
    
        # 2. Najít nebo vytvořit "Hlavní" toolbar
        tb = None
        for child in self.children():
            if isinstance(child, QToolBar) and child.windowTitle() == "Hlavní":
                tb = child
                break
    
        if not tb:
            tb = self.addToolBar("Hlavní")
    
        tb.clear()  # Vyčistit
    
        # Stylizace Hlavního toolbaru
        tb.setIconSize(QSize(20, 20))
        tb.setMovable(False)
        tb.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
    
        tb.setStyleSheet("""
            QToolBar {
                background-color: #2d2d2d;
                border-bottom: 1px solid #3e3e3e;
                padding: 4px;
                spacing: 6px;
            }
            QToolButton {
                background-color: #383838;
                border: 1px solid #505050;
                border-radius: 3px;
                color: #e0e0e0;
                padding: 4px 8px;
                font-size: 12px;
                font-weight: bold;
            }
            QToolButton:hover {
                background-color: #454545;
                border-color: #606060;
            }
            QToolButton:pressed {
                background-color: #252525;
            }
        """)
    
        def get_gen_icon(char, color):
            return self._generate_icon(char, color, "rect")
    
        # --- ACTIONS ---
    
        # 1. Skupina
        if not hasattr(self, "act_add_group"):
            self.act_add_group = QAction("Skupina", self)
            self.act_add_group.setShortcut("Ctrl+G")
        self.act_add_group.setText("Skupina")
        self.act_add_group.setIcon(get_gen_icon("S", QColor("#ff5252")))
    
        # 2. Podskupina
        if not hasattr(self, "act_add_subgroup"):
            self.act_add_subgroup = QAction("Podskupina", self)
            self.act_add_subgroup.setShortcut("Ctrl+Shift+G")
        self.act_add_subgroup.setText("Podskupina")
        self.act_add_subgroup.setIcon(get_gen_icon("P", QColor("#ff8a80")))
    
        # 3. Otázka
        if not hasattr(self, "act_add_question"):
            self.act_add_question = QAction("Otázka", self)
            self.act_add_question.setShortcut(QKeySequence.New)
        self.act_add_question.setText("Otázka")
        self.act_add_question.setIcon(get_gen_icon("O", QColor("#42a5f5")))
    
        # 4. Smazat
        if not hasattr(self, "act_delete"):
            self.act_delete = QAction("Smazat", self)
            self.act_delete.setShortcut(QKeySequence.Delete)
        self.act_delete.setText("Smazat")
        self.act_delete.setIcon(self.style().standardIcon(QStyle.SP_TrashIcon))
    
        # 5. Import
        if not hasattr(self, "act_import_docx"):
            self.act_import_docx = QAction("Import", self)
            if hasattr(self, "_import_from_docx"):
                self.act_import_docx.triggered.connect(self._import_from_docx)
        self.act_import_docx.setText("Import")
        self.act_import_docx.setIcon(self.style().standardIcon(QStyle.SP_ArrowDown))
    
        # 6. Export
        if not hasattr(self, "act_export_docx"):
            self.act_export_docx = QAction("Export", self)
            if hasattr(self, "_export_docx_wizard"):
                self.act_export_docx.triggered.connect(self._export_docx_wizard)
        self.act_export_docx.setText("Export")
        self.act_export_docx.setIcon(self.style().standardIcon(QStyle.SP_ArrowUp))
    
        # --- ZMĚNA: NOVÉ JSON AKCE ---
        if not hasattr(self, "act_load_json"):
            self.act_load_json = QAction("Nahrát DB", self)
            self.act_load_json.triggered.connect(self._action_load_questions_json)
        self.act_load_json.setText("Nahrát DB")
        self.act_load_json.setIcon(self.style().standardIcon(QStyle.SP_DialogOpenButton))
    
        if not hasattr(self, "act_save_json"):
            self.act_save_json = QAction("Uložit DB", self)
            self.act_save_json.triggered.connect(self._action_export_questions_json)
        self.act_save_json.setText("Uložit DB")
        self.act_save_json.setIcon(self.style().standardIcon(QStyle.SP_DialogSaveButton))
        # -----------------------------
    
        # Přidání do Hlavního Toolbaru
        tb.addAction(self.act_add_group)
        tb.addAction(self.act_add_subgroup)
        tb.addAction(self.act_add_question)
        tb.addSeparator()
        tb.addAction(self.act_import_docx)
        tb.addAction(self.act_export_docx)
        tb.addSeparator()
    
        # ZMĚNA: Přidání JSON tlačítek do layoutu toolbaru
        tb.addAction(self.act_load_json)
        tb.addAction(self.act_save_json)
    
        tb.addSeparator()
        tb.addAction(self.act_delete)
    
        self.statusBar().showMessage(f"Datový soubor: {self.data_path}")
        self._refresh_history_table()
    
        self.left_tabs.currentChanged.connect(self._on_left_tab_changed)

    def _action_load_questions_json(self) -> None:
        """Načte kompletní strukturu otázek z vybraného JSON souboru."""
        path, _ = QFileDialog.getOpenFileName(
            self, 
            "Načíst databázi otázek", 
            str(self.project_root), 
            "JSON soubory (*.json)"
        )
        if not path:
            return

        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            if "groups" not in data:
                raise ValueError("JSON neobsahuje klíč 'groups'.")

            # --- Pomocné funkce pro rekurzivní převod ---

            def dict_to_funny_answers(raw_list: list) -> list:
                """Převede seznam slovníků (nebo stringů) na objekty FunnyAnswer."""
                res = []
                for item in raw_list:
                    if isinstance(item, dict):
                        res.append(FunnyAnswer(**item))
                    else:
                        # Fallback pro staré verze (pokud to byl jen string)
                        res.append(FunnyAnswer(text=str(item), date="", author="", source_doc=""))
                return res

            def dict_to_question(d: dict) -> Question:
                """Převede slovník na objekt Question."""
                # Zpracujeme funny answers zvlášť
                f_answers = dict_to_funny_answers(d.get("funny_answers", []))
                
                # Vytvoříme kopii dictu a aktualizujeme funny_answers
                q_args = d.copy()
                q_args["funny_answers"] = f_answers
                
                # Bezpečné vytvoření (ignoruje extra klíče, pokud by vadily, 
                # ale dataclass **kwargs by to měl zvládnout, pokud sedí pole.
                # Pro jistotu explicitní mapping u klíčových polí, pokud by JSON obsahoval balast)
                return Question(
                    id=q_args.get("id", str(_uuid.uuid4())),
                    type=q_args.get("type", "classic"),
                    text_html=q_args.get("text_html", ""),
                    title=q_args.get("title", ""),
                    points=int(q_args.get("points", 1)),
                    bonus_correct=float(q_args.get("bonus_correct", 0.0)),
                    bonus_wrong=float(q_args.get("bonus_wrong", 0.0)),
                    created_at=q_args.get("created_at", ""),
                    correct_answer=q_args.get("correct_answer", ""),
                    funny_answers=f_answers,
                    image_path=q_args.get("image_path", ""),
                    image_width_cm=float(q_args.get("image_width_cm", 0.0) or 0.0),
                    image_height_cm=float(q_args.get("image_height_cm", 0.0) or 0.0),
                    image_keep_aspect=bool(q_args.get("image_keep_aspect", True)),
                )

            def dict_to_subgroup(d: dict) -> Subgroup:
                """Rekurzivně převede slovník na objekt Subgroup."""
                # 1. Převedeme otázky
                qs = [dict_to_question(q) for q in d.get("questions", [])]
                
                # 2. Převedeme vnořené podskupiny (rekurze)
                subs = [dict_to_subgroup(s) for s in d.get("subgroups", [])]
                
                return Subgroup(
                    id=d.get("id", str(_uuid.uuid4())),
                    name=d.get("name", "Bez názvu"),
                    subgroups=subs,
                    questions=qs
                )

            # --- Hlavní smyčka převodu ---
            new_groups = []
            for g_data in data["groups"]:
                # Převedeme podskupiny v této skupině
                converted_subgroups = [dict_to_subgroup(sg) for sg in g_data.get("subgroups", [])]
                
                new_group = Group(
                    id=g_data.get("id", str(_uuid.uuid4())),
                    name=g_data.get("name", "Bez názvu"),
                    subgroups=converted_subgroups
                )
                new_groups.append(new_group)

            # Nahrazení dat v aplikaci
            self.root.groups = new_groups
            
            # Refresh UI
            self._refresh_tree()
            self.save_data() # Uložíme hned do pracovního souboru
            self.statusBar().showMessage(f"Databáze úspěšně načtena: {os.path.basename(path)}", 5000)

        except Exception as e:
            # Výpis chyby do konzole pro lepší debug
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Chyba načítání JSON", f"Nepodařilo se načíst soubor.\n\nDetail: {str(e)}")


    def _action_export_questions_json(self) -> None:
        default_name = f"questions_export_{QDateTime.currentDateTime().toString('yyyyMMdd_HHmm')}.json"
        path, _ = QFileDialog.getSaveFileName(
            self, "Exportovat databázi", str(self.project_root / default_name), "JSON soubory (*.json)"
        )
        if not path: return
        try:
            data_out = {"groups": [self._serialize_group(g) for g in self.root.groups]}
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data_out, f, indent=2, ensure_ascii=False)
            self.statusBar().showMessage(f"Exportováno: {os.path.basename(path)}", 5000)
        except Exception as e:
            QMessageBox.critical(self, "Chyba", f"Nepodařilo se exportovat:\n{e}")


    def _on_left_tab_changed(self, index: int) -> None:
        """Skrývá/zobrazuje pravý panel podle aktivní záložky."""
        current_widget = self.left_tabs.widget(index)
        
        if current_widget == self.tab_questions:
            self.detail_stack.setVisible(True)
        else:
            self.detail_stack.setVisible(False)
            
        if hasattr(self, "tab_funny") and current_widget == self.tab_funny:
            if hasattr(self, "_refresh_funny_answers_tab"):
                self._refresh_funny_answers_tab()


    def _on_tree_context_menu(self, pos: QPoint) -> None:
        """Kontextové menu stromu otázek (v6.7.2)."""
        item = self.tree.itemAt(pos)
        if not item:
            return
            
        self.tree.setCurrentItem(item)

        # Robustní získání metadat (podpora tuple i dict)
        raw_data = item.data(0, Qt.UserRole)
        kind = "unknown"
        
        if isinstance(raw_data, tuple) and len(raw_data) >= 1:
            kind = raw_data[0]
        elif isinstance(raw_data, dict):
            kind = raw_data.get("kind", "unknown")

        menu = QMenu(self.tree)
        has_action = False

        # 1. Přidat podskupinu (Group/Subgroup)
        if kind in ("group", "subgroup"):
            act = menu.addAction("Přidat podskupinu")
            act.triggered.connect(self._add_subgroup)
            has_action = True

        # 2. Přidat otázku (Subgroup)
        if kind == "subgroup":
            act = menu.addAction("Přidat otázku")
            act.triggered.connect(self._add_question)
            has_action = True
            
        # 3. Duplikovat otázku (Question)
        if kind == "question":
            act = menu.addAction("Duplikovat otázku")
            act.triggered.connect(self._duplicate_question)
            has_action = True

        if has_action:
            menu.addSeparator()

        # 4. Smazat (Vše) -> Použijeme existující metodu _delete_selected
        act_del = menu.addAction("Smazat")
        act_del.triggered.connect(self._delete_selected)

        menu.exec(self.tree.mapToGlobal(pos))

    def _change_indent(self, steps: int) -> None:
        """Změní odsazení aktuálního bloku nebo listu."""
        cursor = self.text_edit.textCursor()
        cursor.beginEditBlock()
        
        current_list = cursor.currentList()
        if current_list:
            # Pokud jsme v listu, měníme level (odsazení formátu listu)
            fmt = current_list.format()
            current_indent = fmt.indent()
            new_indent = max(1, current_indent + steps)
            fmt.setIndent(new_indent)
            current_list.setFormat(fmt)
        else:
            # Pokud je to běžný text, měníme margin bloku
            block_fmt = cursor.blockFormat()
            current_margin = block_fmt.leftMargin()
            # Krok odsazení např. 20px
            new_margin = max(0, current_margin + (steps * 20))
            block_fmt.setLeftMargin(new_margin)
            cursor.setBlockFormat(block_fmt)
            
        cursor.endEditBlock()
        self._autosave_schedule()
        self.text_edit.setFocus()


    def _on_format_bullets(self, checked: bool) -> None:
        """Přepne aktuální výběr na odrážky s lepším odsazením."""
        cursor = self.text_edit.textCursor()
        cursor.beginEditBlock()
        
        if checked:
            # Vytvoříme formát seznamu
            list_fmt = QTextListFormat()
            list_fmt.setStyle(QTextListFormat.ListDisc)
            list_fmt.setIndent(1) # Level 1
            
            # Nastavení odsazení (v pixelech/bodech) pro vizuální úpravu
            # NumberSuffix = mezera za odrážkou
            list_fmt.setNumberPrefix("")
            list_fmt.setNumberSuffix(" ") 
            
            cursor.createList(list_fmt)
            
            # Vynucení odsazení bloku (pro celý list item)
            block_fmt = cursor.blockFormat()
            block_fmt.setLeftMargin(15)  # Odsazení celého bloku zleva
            block_fmt.setTextIndent(-10) # Předsazení odrážky (aby byla vlevo od textu)
            cursor.setBlockFormat(block_fmt)
        else:
            # Zrušit list
            # Standardní blok bez listu
            block_fmt = QTextBlockFormat()
            block_fmt.setObjectIndex(-1) # Not a list
            cursor.setBlockFormat(block_fmt)
            
            # Reset odsazení
            cursor.setBlockFormat(QTextBlockFormat())

        cursor.endEditBlock()
        # Synchronizace stavu tlačítka (pokud by cursor change změnil stav zpět)
        self.text_edit.setFocus()

    def _refresh_history_table(self) -> None:
        """Načte historii exportů a zobrazí ji ve stylu 'Hall of Shame' / System Log."""
        import re
        
        history_file = self.project_root / "data" / "history.json"
        history = []
        if history_file.exists():
            try:
                with open(history_file, "r", encoding="utf-8") as f:
                    history = json.load(f)
            except Exception as e:
                print(f"Chyba při čtení historie: {e}")
        
        self.table_history.setStyleSheet("""
            QTableWidget {
                background-color: #121212;
                color: #e0e0e0;
                gridline-color: #333333;
                border: 1px solid #333;
                font-family: 'Consolas', 'Monospace', 'Courier New';
                selection-background-color: #c62828;
                selection-color: #ffffff;
            }
            QHeaderView::section {
                background-color: #1e1e1e;
                color: #9e9e9e;
                padding: 4px;
                border: 1px solid #333;
                font-weight: bold;
            }
        """)
        self.table_history.verticalHeader().setVisible(False)
        self.table_history.setShowGrid(False)
        self.table_history.setAlternatingRowColors(True)
        
        self.table_history.setColumnCount(4)
        self.table_history.setHorizontalHeaderLabels(["TYP", "CÍLOVÝ SOUBOR", "DIGITÁLNÍ OTISK (HASH)", "ČASOVÁ STOPA"])
        
        header = self.table_history.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(1, QHeaderView.Stretch)
        header.setSectionResizeMode(2, QHeaderView.Interactive)
        header.setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.table_history.setColumnWidth(2, 220)

        self.table_history.setRowCount(0)
        self.table_history.setSortingEnabled(False)
        
        font_mono = QFont("Courier New")
        font_mono.setStyleHint(QFont.Monospace)
        font_bold = QFont("Courier New"); font_bold.setBold(True)
        
        # Helper: Ikona balíku (Stack)
        def get_stack_icon(color):
            pix = QPixmap(24, 24)
            pix.fill(Qt.transparent)
            p = QPainter(pix)
            p.setRenderHint(QPainter.Antialiasing)
            p.setPen(QPen(QColor(30, 30, 30), 1)) # Dark border
            p.setBrush(color)
            
            # 3 dokumenty za sebou
            # Bottom
            p.drawRect(4, 4, 14, 16)
            # Middle
            p.drawRect(6, 6, 14, 16)
            # Top
            p.drawRect(8, 8, 14, 16)
            p.end()
            return QIcon(pix)

        for entry in history:
            row = self.table_history.rowCount()
            self.table_history.insertRow(row)
            
            raw_fn = entry.get("filename", "Neznámý soubor")
            h = entry.get("hash", "---")
            raw_date = entry.get("date", "")
            
            match = re.match(r"Balík (\d+) verzí: (.+)", raw_fn)
            if match:
                count = int(match.group(1))
                clean_fn = match.group(2)
                is_multi = True
            else:
                count = 1
                clean_fn = raw_fn
                is_multi = False
            
            date_str = raw_date
            try:
                if raw_date:
                    dt = datetime.fromisoformat(raw_date)
                    date_str = dt.strftime("%Y-%m-%d %H:%M")
            except: pass

            # 1. TYP (Ikona + Počet)
            if is_multi:
                icon = get_stack_icon(QColor("#ff9800")) # Oranžový balík
                text_type = f"{count:02d}x" # Zero pad
                color_type = QColor("#ff9800")
            else:
                icon = self.style().standardIcon(QStyle.SP_FileIcon)
                text_type = "01x"
                color_type = QColor("#9e9e9e")

            item_type = QTableWidgetItem(text_type)
            item_type.setIcon(icon)
            item_type.setForeground(QBrush(color_type))
            item_type.setFont(font_bold)
            item_type.setTextAlignment(Qt.AlignCenter)
            item_type.setFlags(item_type.flags() ^ Qt.ItemIsEditable)
            self.table_history.setItem(row, 0, item_type)

            # 2. SOUBOR
            item_fn = QTableWidgetItem(clean_fn)
            item_fn.setForeground(QBrush(QColor("#80d8ff")))
            item_fn.setFont(font_bold)
            item_fn.setFlags(item_fn.flags() ^ Qt.ItemIsEditable)
            self.table_history.setItem(row, 1, item_fn)
            
            # 3. HASH
            if len(h) > 24:
                display_hash = f"{h[:12]}...{h[-12:]}"
            else:
                display_hash = h
            item_hash = QTableWidgetItem(display_hash)
            item_hash.setForeground(QBrush(QColor("#ff5252")))
            item_hash.setFont(font_mono)
            item_hash.setTextAlignment(Qt.AlignCenter)
            item_hash.setToolTip(h)
            item_hash.setFlags(item_hash.flags() ^ Qt.ItemIsEditable)
            self.table_history.setItem(row, 2, item_hash)

            # 4. DATUM
            item_date = QTableWidgetItem(f"{date_str}")
            item_date.setForeground(QBrush(QColor("#757575")))
            item_date.setFont(font_mono)
            item_date.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            item_date.setFlags(item_date.flags() ^ Qt.ItemIsEditable)
            self.table_history.setItem(row, 3, item_date)
            
        self.table_history.setSortingEnabled(True)
        self.table_history.sortItems(1, Qt.AscendingOrder)

    def _on_history_context_menu(self, pos) -> None:
        """Zobrazí kontextové menu pro tabulku historie."""
        items = self.table_history.selectedItems()
        if not items:
            return
            
        menu = QMenu(self)
        act_del = QAction("Smazat záznam(y)", self)
        act_del.triggered.connect(self._delete_history_items)
        menu.addAction(act_del)
        menu.exec(self.table_history.mapToGlobal(pos))

    def _clear_all_history(self) -> None:
        """Smaže celou historii exportů."""
        history_file = self.project_root / "data" / "history.json"
        if not history_file.exists():
            return
            
        reply = QMessageBox.question(
            self, 
            "Vymazat historii", 
            "Opravdu chcete nenávratně smazat CELOU historii exportů?\n\nTato akce neodstraní vygenerované soubory, pouze záznamy v logu.",
            QMessageBox.Yes | QMessageBox.No, 
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            try:
                with open(history_file, "w", encoding="utf-8") as f:
                    json.dump([], f)
                self._refresh_history_table()
                QMessageBox.information(self, "Hotovo", "Historie byla vymazána.")
            except Exception as e:
                QMessageBox.critical(self, "Chyba", f"Nepodařilo se smazat historii:\n{e}")


    def _delete_history_items(self) -> None:
        """Smaže vybrané záznamy z historie."""
        rows = sorted(set(index.row() for index in self.table_history.selectedIndexes()), reverse=True)
        if not rows:
            return

        if QMessageBox.question(self, "Smazat", f"Opravdu smazat {len(rows)} záznamů z historie?") != QMessageBox.Yes:
            return
        
        # Získání identifikátorů (Filename na indexu 1, Hash na indexu 2)
        to_remove = [] 
        for r in rows:
            # UPDATE: Indexy posunuty kvůli sloupci "TYP"
            fn = self.table_history.item(r, 1).text() 
            h = self.table_history.item(r, 2).toolTip() # Hash bereme z tooltipu, protože v textu může být zkrácený ("abcd...1234")!
            # Fallback pokud tooltip není (což by měl být)
            if not h: h = self.table_history.item(r, 2).text()
            
            to_remove.append((fn, h))

        history_file = self.project_root / "data" / "history.json"
        history = []
        if history_file.exists():
            try:
                with open(history_file, "r", encoding="utf-8") as f:
                    history = json.load(f)
            except Exception:
                pass

        # Filtrace
        new_history = []
        for entry in history:
            match = False
            for r_fn, r_h in to_remove:
                # Porovnání: Hash musí sedět přesně.
                # Filename: U balíků jsme v tabulce zobrazili "clean_fn" (bez "Balík X verzí: "), ale v JSONu je "raw_fn".
                # Musíme porovnat chytřeji.
                # Hash je unikátní identifikátor (pro daný export). 
                # Pokud máme hash, stačí hash. Pokud ne (starší záznamy?), musíme řešit filename.
                
                json_h = entry.get("hash", "")
                json_fn = entry.get("filename", "")
                
                # Pokud hash v JSONu existuje a sedí, je to shoda.
                if json_h and r_h and json_h == r_h:
                    match = True
                    break
                
                # Fallback: Pokud hash není, zkusíme filename (obsahuje)
                # V tabulce máme "soubor.docx", v jsonu "Balík X: soubor.docx"
                # Zkusíme: json_fn.endswith(r_fn)
                if not json_h and r_fn and r_fn in json_fn:
                    match = True
                    break
                    
            if not match:
                new_history.append(entry)
        
        # Uložení
        try:
            with open(history_file, "w", encoding="utf-8") as f:
                json.dump(new_history, f, indent=2, ensure_ascii=False)
        except Exception as e:
            QMessageBox.warning(self, "Chyba", f"Nelze uložit historii:\n{e}")

        self._refresh_history_table()

    def _init_funny_answers_tab(self) -> None:
        """Inicializuje záložku s přehledem vtipných odpovědí."""
        self.tab_funny = QWidget()
        layout = QVBoxLayout(self.tab_funny)
        layout.setContentsMargins(4, 4, 4, 4)

        self.tree_funny = QTreeWidget()
        
        # -- STYLIZACE "HALL OF SHAME" --
        self.tree_funny.setStyleSheet("""
            QTreeWidget {
                background-color: #121212;
                color: #e0e0e0;
                border: 1px solid #333;
                font-family: 'Consolas', 'Monospace', 'Courier New';
            }
            QHeaderView::section {
                background-color: #1e1e1e;
                color: #9e9e9e;
                padding: 4px;
                border: 1px solid #333;
                font-weight: bold;
            }
            QTreeWidget::item:hover {
                background-color: #1e1e1e;
            }
        """)
        
        self.tree_funny.setColumnCount(4)
        self.tree_funny.setHeaderLabels(["OTÁZKA / ODPOVĚĎ (HŘÍCH)", "DATUM", "PACHATEL", "ZDROJ"])
        self.tree_funny.setRootIsDecorated(True)
        self.tree_funny.setIndentation(18)
        self.tree_funny.setUniformRowHeights(True)
        self.tree_funny.setAlternatingRowColors(False)

        header = self.tree_funny.header()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(3, QHeaderView.ResizeToContents)

        self.tree_funny.setSelectionMode(QAbstractItemView.NoSelection)
        self.tree_funny.setFocusPolicy(Qt.NoFocus)

        layout.addWidget(self.tree_funny)
        self.left_tabs.addTab(self.tab_funny, "Hall of Shame - legendární odpovědi")
        
        # DŮLEŽITÉ: Načíst data ihned po inicializaci
        self._refresh_funny_answers_tab()

    def _on_funny_tree_select(self):
        """Zobrazí detail vybrané vtipné odpovědi ze stromu."""
        selected = self.tree_funny.selectedItems()
        if not selected:
            self.text_funny_detail.clear()
            return
            
        item = selected[0]
        # Pokud je to top-level item (otázka), nic nezobrazujeme (nebo název otázky)
        if item.parent() is None:
            self.text_funny_detail.clear()
            return
            
        # Získáme plný text z UserRole
        full_text = item.data(0, Qt.UserRole)
        if full_text:
            self.text_funny_detail.setText(full_text)
        else:
            # Fallback na text položky, pokud data chybí
            self.text_funny_detail.setText(item.text(0))

    def _filter_funny_answers(self, text: str):
        """Filtruje položky ve stromu vtipných odpovědí."""
        search_text = text.lower()
        
        # Projdeme všechny top-level položky (otázky)
        root = self.tree_funny.invisibleRootItem()
        for i in range(root.childCount()):
            q_item = root.child(i)
            
            # Projdeme odpovědi (děti)
            has_visible_child = False
            for j in range(q_item.childCount()):
                child = q_item.child(j)
                child_text = child.text(0).lower()
                
                if not search_text or search_text in child_text:
                    child.setHidden(False)
                    has_visible_child = True
                else:
                    child.setHidden(True)
            
            # Pokud otázka nemá viditelné odpovědi a sama neodpovídá filtru, skryjeme ji
            # (Zde pro jednoduchost filtrujeme jen podle obsahu odpovědí)
            q_item.setHidden(not has_visible_child)
            
            # Pokud filtr není prázdný a našli jsme shodu, rozbalíme
            if search_text and has_visible_child:
                q_item.setExpanded(True)
            elif not search_text:
                q_item.setExpanded(False)

    def _refresh_funny_answers_tab(self) -> None:
        """Znovu vygeneruje strom 'Seznam vtipných odpovědí' ze struktury otázek."""
        if not hasattr(self, "tree_funny"):
            return

        self.tree_funny.clear()

        root = getattr(self, "root", None)
        if root is None or not root.groups:
            return
        
        from pathlib import Path

        font_mono = QFont("Courier New"); font_mono.setStyleHint(QFont.Monospace)
        font_bold = QFont("Courier New"); font_bold.setBold(True)
        
        color_q_text = QBrush(QColor("#ff9800")) 
        color_ans_text = QBrush(QColor("#80d8ff")) 
        color_name = QBrush(QColor("#ff5252")) 
        color_date = QBrush(QColor("#757575")) 
        color_source = QBrush(QColor("#9e9e9e"))
        
        icon_q = self.style().standardIcon(QStyle.SP_MessageBoxInformation)

        questions_with_funny = []
        
        def collect_questions_recursive(subgroups):
            for sg in subgroups:
                for q in sg.questions:
                    if hasattr(q, "funny_answers") and q.funny_answers:
                        questions_with_funny.append(q)
                if sg.subgroups:
                    collect_questions_recursive(sg.subgroups)

        for g in root.groups:
            collect_questions_recursive(g.subgroups)
            
        questions_with_funny.sort(key=lambda q: q.title.lower() if q.title else "")

        # Helper pro převod HTML na Plain Text (pro tooltip)
        def html_to_plain(html_text):
            if not html_text: return ""
            doc = QTextDocument()
            doc.setHtml(html_text)
            return doc.toPlainText().strip()

        for q in questions_with_funny:
            q_title = q.title if q.title else "(bez názvu)"
            q_item = QTreeWidgetItem([q_title])
            q_item.setIcon(0, icon_q)
            q_item.setForeground(0, color_q_text)
            q_item.setFont(0, font_bold)
            
            # Čistý text do tooltipu
            plain_text = html_to_plain(q.text_html)
            # Omezíme délku tooltipu, aby nebyl přes celou obrazovku
            if len(plain_text) > 300:
                plain_text = plain_text[:300] + "..."
            q_item.setToolTip(0, plain_text)
            
            self.tree_funny.addTopLevelItem(q_item)
            q_item.setExpanded(True)
            
            answers = sorted(q.funny_answers, key=lambda x: x.date, reverse=True)
            
            for fa in answers:
                text = fa.text
                date = fa.date
                author = getattr(fa, "author", "Neznámý")
                full_source = getattr(fa, "source_doc", "")
                
                source_display = Path(full_source).name if full_source else ""
                
                child = QTreeWidgetItem([text, date, author, source_display])
                
                child.setForeground(0, color_ans_text); child.setFont(0, font_mono)
                child.setToolTip(0, text) 
                
                child.setForeground(1, color_date); child.setFont(1, font_mono)
                child.setForeground(2, color_name); child.setFont(2, font_bold)
                
                child.setForeground(3, color_source); child.setFont(3, font_mono)
                child.setToolTip(3, full_source)
                
                q_item.addChild(child)

    def register_export(self, filename: str, k_hash: str) -> None:
        """Zaznamená nový export a obnoví tabulku."""
        history_file = self.project_root / "data" / "history.json"
        history = []
        if history_file.exists():
            try:
                with open(history_file, "r", encoding="utf-8") as f:
                    history = json.load(f)
            except Exception:
                pass # Ignorujeme chyby čtení, vytvoříme nový seznam
        
        # Přidání záznamu
        history.append({
            "filename": filename, 
            "hash": k_hash, 
            "date": datetime.now().isoformat()
        })
        
        # Uložení
        try:
            with open(history_file, "w", encoding="utf-8") as f:
                json.dump(history, f, indent=2, ensure_ascii=False)
        except Exception as e:
            QMessageBox.warning(self, "Chyba historie", f"Nepodařilo se uložit historii exportu:\n{e}")
            
        self._refresh_history_table()


    def _set_editor_enabled(self, enabled: bool) -> None:
        self.editor_toolbar.setEnabled(enabled)
        self.title_edit.setEnabled(enabled)
        self.combo_type.setEnabled(enabled)
        self.image_path_edit.setEnabled(enabled)
        self.btn_choose_image.setEnabled(enabled)
        self.btn_clear_image.setEnabled(enabled)
        # Rozměry obrázku povolíme jen pokud je nastaven a existuje
        full_img = getattr(self, "_current_image_full_path", "") or ""
        has_img = bool(full_img and os.path.exists(full_img))
        if hasattr(self, "spin_img_w_cm"):
            self.spin_img_w_cm.setEnabled(enabled and has_img)
        if hasattr(self, "spin_img_h_cm"):
            self.spin_img_h_cm.setEnabled(enabled and has_img)
        if hasattr(self, "chk_img_keep_aspect"):
            self.chk_img_keep_aspect.setEnabled(enabled and has_img)
        self.spin_points.setEnabled(enabled)
        self.spin_bonus_correct.setEnabled(enabled and self.combo_type.currentIndex() == 1)
        self.spin_bonus_wrong.setEnabled(enabled and self.combo_type.currentIndex() == 1)
        self.text_edit.setEnabled(enabled)
        self.btn_save_question.setEnabled(enabled)

    def _connect_signals(self) -> None:
        self.tree.itemSelectionChanged.connect(self._on_tree_selection_changed)
        # self.tree.itemChanged.connect(self._on_tree_item_changed) # REMOVED previously
    
        self.btn_save_question.clicked.connect(self._on_save_question_clicked)
        self.btn_rename.clicked.connect(self._on_rename_clicked)
        self.btn_choose_image.clicked.connect(self._choose_question_image)
        self.btn_clear_image.clicked.connect(self._clear_question_image)
    
        # Tree actions context menu
        self.act_add_group.triggered.connect(self._add_group)
        self.act_add_subgroup.triggered.connect(self._add_subgroup)
        self.act_add_question.triggered.connect(self._add_question)
    
        # Delete shortcut
        self.act_delete.triggered.connect(self._bulk_delete_selected)
        self.btn_delete_selected.clicked.connect(self._bulk_delete_selected)
    
        # Autosave triggers
        self.title_edit.textChanged.connect(self._autosave_schedule)
        self.image_path_edit.textChanged.connect(self._autosave_schedule)
        self.spin_img_w_cm.valueChanged.connect(self._autosave_schedule)
        self.spin_img_h_cm.valueChanged.connect(self._autosave_schedule)
        # Synchronizace rozměrů obrázku (poměr stran)
        self.spin_img_w_cm.valueChanged.connect(self._on_export_img_w_changed)
        self.spin_img_h_cm.valueChanged.connect(self._on_export_img_h_changed)
        if hasattr(self, "chk_img_keep_aspect"):
            self.chk_img_keep_aspect.stateChanged.connect(self._on_keep_aspect_changed)
        if hasattr(self, "chk_img_keep_aspect"):
            self.chk_img_keep_aspect.stateChanged.connect(self._autosave_schedule)
        self.combo_type.currentIndexChanged.connect(self._on_type_changed_ui)
        self.spin_points.valueChanged.connect(self._autosave_schedule)
        self.spin_bonus_correct.valueChanged.connect(self._autosave_schedule)
        self.spin_bonus_wrong.valueChanged.connect(self._autosave_schedule)
        self.text_edit.textChanged.connect(self._autosave_schedule)
    
        # Autosave triggers (New fields)
        self.edit_correct_answer.textChanged.connect(self._autosave_schedule)
        self.table_funny.itemChanged.connect(self._autosave_schedule)
    
        # NOVÉ: Kontextové menu pro vtipné odpovědi (Editace)
        self.table_funny.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table_funny.customContextMenuRequested.connect(self._on_funny_context_menu)
    
        # Formátování
        self.action_bold.triggered.connect(lambda: self._toggle_format("bold"))
        self.action_italic.triggered.connect(lambda: self._toggle_format("italic"))
        self.action_underline.triggered.connect(lambda: self._toggle_format("underline"))
        self.action_color.triggered.connect(self._choose_color)
        self.action_bullets.triggered.connect(self._toggle_bullets)
    
        self.action_align_left.triggered.connect(lambda: self._apply_alignment(Qt.AlignLeft))
        self.action_align_center.triggered.connect(lambda: self._apply_alignment(Qt.AlignHCenter))
        self.action_align_right.triggered.connect(lambda: self._apply_alignment(Qt.AlignRight))
        self.action_align_justify.triggered.connect(lambda: self._apply_alignment(Qt.AlignJustify))
    
        self.text_edit.cursorPositionChanged.connect(self._sync_toolbar_to_cursor)
    
        # Filter
        self.filter_edit.textChanged.connect(self._apply_filter)
    
        # Drag Drop Move (btn)
        self.btn_move_selected.clicked.connect(self._move_selected_dialog)
    
        # Tlačítka pro vtipné odpovědi
        self.btn_add_funny.clicked.connect(self._add_funny_row)
        self.btn_rem_funny.clicked.connect(self._remove_funny_row)
    
        # NOVÉ: Koš
        if hasattr(self, "table_trash"):
            self.table_trash.itemSelectionChanged.connect(self._on_trash_selection_changed)
        if hasattr(self, "btn_trash_restore"):
            self.btn_trash_restore.clicked.connect(self._trash_restore_selected)
        if hasattr(self, "btn_trash_delete"):
            self.btn_trash_delete.clicked.connect(self._trash_delete_selected)
        if hasattr(self, "btn_trash_empty"):
            self.btn_trash_empty.clicked.connect(self._trash_empty)

    def _add_funny_row(self) -> None:
        # Předáváme self.project_root pro vyhledání souborů
        dlg = FunnyAnswerDialog(self, project_root=self.project_root)

        if dlg.exec() == QDialog.Accepted:
            text, date_str, author, source_doc = dlg.get_data()

            row = self.table_funny.rowCount()
            self.table_funny.insertRow(row)

            self.table_funny.setItem(row, 0, QTableWidgetItem(text))
            self.table_funny.setItem(row, 1, QTableWidgetItem(date_str))
            self.table_funny.setItem(row, 2, QTableWidgetItem(author))

            # Ve sloupci "Zdroj" zobrazíme jen název souboru,
            # ale do UserRole uložíme plnou cestu
            display_source = os.path.basename(source_doc) if source_doc else ""
            source_item = QTableWidgetItem(display_source)
            source_item.setData(Qt.UserRole, source_doc)
            self.table_funny.setItem(row, 3, source_item)

            self._autosave_schedule()

    def _remove_funny_row(self) -> None:
        rows = sorted(set(index.row() for index in self.table_funny.selectedIndexes()), reverse=True)
        for r in rows:
            self.table_funny.removeRow(r)
        if rows:
            self._autosave_schedule()

    def _on_funny_context_menu(self, pos) -> None:
        """Kontextové menu pro tabulku vtipných odpovědí."""
        item = self.table_funny.itemAt(pos)
        if not item:
            return
        
        menu = QMenu(self)
        act_edit = QAction("Upravit odpověď", self)
        act_edit.triggered.connect(lambda: self._edit_funny_row(item.row()))
        menu.addAction(act_edit)
        
        act_del = QAction("Smazat odpověď", self)
        act_del.triggered.connect(self._remove_funny_row) # Použije selected items
        menu.addAction(act_del)
        
        menu.exec(self.table_funny.mapToGlobal(pos))

    def _edit_funny_row(self, row: int) -> None:
        """Otevře dialog pro editaci vtipné odpovědi na daném řádku."""
        # Načtení dat z tabulky
        text_item = self.table_funny.item(row, 0)
        date_item = self.table_funny.item(row, 1)
        author_item = self.table_funny.item(row, 2)
        source_item = self.table_funny.item(row, 3) if self.table_funny.columnCount() > 3 else None

        if not text_item or not date_item or not author_item:
            return

        old_text = text_item.text()
        old_date = date_item.text()
        old_author = author_item.text()

        # Plná cesta je v UserRole, pokud není, použijeme text
        if source_item is not None:
            data = source_item.data(Qt.UserRole)
            if isinstance(data, str) and data:
                old_source = data
            else:
                old_source = source_item.text()
        else:
            old_source = ""

        # Otevření dialogu
        dlg = FunnyAnswerDialog(self, project_root=self.project_root)
        dlg.setWindowTitle("Upravit vtipnou odpověď")
        dlg.set_data(old_text, old_date, old_author, old_source)

        if dlg.exec() == QDialog.Accepted:
            new_text, new_date, new_author, new_source = dlg.get_data()

            # Uložení zpět do tabulky
            self.table_funny.setItem(row, 0, QTableWidgetItem(new_text))
            self.table_funny.setItem(row, 1, QTableWidgetItem(new_date))
            self.table_funny.setItem(row, 2, QTableWidgetItem(new_author))

            display_source = os.path.basename(new_source) if new_source else ""
            new_source_item = QTableWidgetItem(display_source)
            new_source_item.setData(Qt.UserRole, new_source)
            self.table_funny.setItem(row, 3, new_source_item)

            self._autosave_schedule()
        
    def _move_selected_dialog(self) -> None:
        """Otevře dialog pro přesun vybraných otázek do jiné skupiny/podskupiny."""
        ids = self._selected_question_ids()
        if not ids:
            QMessageBox.information(self, "Přesun", "Vyberte otázky k přesunu.")
            return

        # MoveTargetDialog musí být definován v souboru (byl vidět v původním výpisu)
        dlg = MoveTargetDialog(self)
        if dlg.exec() != QDialog.Accepted:
            return

        gid, sgid = dlg.selected_target()
        if not gid:
            return

        # Nalezení cílové podskupiny
        target_sg: Optional[Subgroup] = None
        if sgid:
            target_sg = self._find_subgroup(gid, sgid)
        else:
            # Cíl je skupina -> zkusíme najít první podskupinu nebo vytvoříme Default
            g = self._find_group(gid)
            if g:
                if g.subgroups:
                    target_sg = g.subgroups[0]
                else:
                    # Vytvoření defaultní podskupiny, pokud skupina žádnou nemá
                    new_sg = Subgroup(id=str(_uuid.uuid4()), name="Default", subgroups=[], questions=[])
                    g.subgroups.append(new_sg)
                    target_sg = new_sg

        if not target_sg:
            QMessageBox.warning(self, "Chyba", "Cílová skupina/podskupina nebyla nalezena.")
            return

        # PROVEDENÍ PŘESUNU
        # 1. Najdeme a vyjmeme otázky z původních umístění
        moved_questions: List[Question] = []

        def remove_from_list(sgs: List[Subgroup]):
            for sg in sgs:
                # Ponecháme jen ty, které NEJSOU v seznamu k přesunu
                # Ty co JSOU, si uložíme
                to_keep = []
                for q in sg.questions:
                    if q.id in ids:
                        moved_questions.append(q)
                    else:
                        to_keep.append(q)
                sg.questions = to_keep
                
                # Rekurze
                remove_from_list(sg.subgroups)

        for g in self.root.groups:
            remove_from_list(g.subgroups)

        # 2. Vložíme je do cíle
        # (Otázky se přidají na konec cílové podskupiny)
        target_sg.questions.extend(moved_questions)

        # 3. Uložit a obnovit
        self._refresh_tree()
        self._reselect_questions(ids) # Zkusíme znovu označit přesunuté
        self.save_data()
        self.statusBar().showMessage(f"Přesunuto {len(moved_questions)} otázek.", 3000)


    # -------------------- Práce s daty (JSON) --------------------

    def default_root_obj(self) -> RootData:
        root = RootData(groups=[])
        if not hasattr(root, "trash") or not isinstance(getattr(root, "trash", None), list):
            root.trash = []
        return root

    def load_data(self) -> None:
        if self.data_path.exists():
            try:
                with self.data_path.open("r", encoding="utf-8") as f:
                    raw = json.load(f)
    
                groups: List[Group] = []
                for g in raw.get("groups", []):
                    groups.append(self._parse_group(g))
    
                trash_raw = raw.get("trash", [])
                if not isinstance(trash_raw, list):
                    trash_raw = []
    
                self.root = RootData(groups=[])
                self.root.groups = groups
    
                # KOŠ: nastavíme až po vytvoření RootData (ne přes constructor)
                if not hasattr(self.root, "trash") or not isinstance(getattr(self.root, "trash", None), list):
                    self.root.trash = []
                self.root.trash = trash_raw
    
            except Exception as e:
                QMessageBox.warning(
                    self,
                    "Načtení selhalo",
                    f"Soubor {self.data_path} nelze načíst: {e}\nVytvořen prázdný projekt."
                )
                self.root = self.default_root_obj()
        else:
            self.root = self.default_root_obj()

    def save_data(self) -> None:
        self._apply_editor_to_current_question(silent=True)
        self.data_path.parent.mkdir(parents=True, exist_ok=True)
        data = {"groups": [self._serialize_group(g) for g in self.root.groups], "trash": getattr(self.root, "trash", [])}
        try:
            sf = QSaveFile(str(self.data_path))
            sf.open(QSaveFile.WriteOnly)
            payload = json.dumps(data, ensure_ascii=False, indent=2)
            sf.write(QByteArray(payload.encode("utf-8")))
            sf.commit()
            self.statusBar().showMessage(f"Uloženo: {self.data_path}", 1500)
        except Exception as e:
            QMessageBox.critical(self, "Uložení selhalo", f"Chyba při ukládání do {self.data_path}:\n{e}")

    def _parse_group(self, g: dict) -> Group:
        subgroups = [self._parse_subgroup(sg) for sg in g.get("subgroups", [])]
        return Group(id=g["id"], name=g["name"], subgroups=subgroups)

    def _parse_subgroup(self, sg: dict) -> Subgroup:
        subgroups_raw = sg.get("subgroups", [])
        subgroups = [self._parse_subgroup(s) for s in subgroups_raw]
        questions = [self._parse_question(q) for q in sg.get("questions", [])]
        return Subgroup(id=sg["id"], name=sg["name"], subgroups=subgroups, questions=questions)

    def _parse_question(self, q: dict) -> Question:
        title = q.get("title") or self._derive_title_from_html(
            q.get("text_html") or "<p></p>",
            prefix=("BONUS: " if q.get("type") == "bonus" else ""),
        )
        bc_default = 1.0 if q.get("type") == "bonus" else 0.0
        bw_default = 0.0
        try:
            bc = round(float(q.get("bonus_correct", bc_default)), 2)
        except Exception:
            bc = round(float(bc_default), 2)
        try:
            bw = round(float(q.get("bonus_wrong", bw_default)), 2)
        except Exception:
            bw = round(float(bw_default), 2)

        # Deserializace vtipných odpovědí (včetně zdrojového dokumentu, pokud je uložen)
        f_answers_raw = q.get("funny_answers", [])
        f_answers: List[FunnyAnswer] = []
        for item in f_answers_raw:
            if isinstance(item, dict):
                f_answers.append(
                    FunnyAnswer(
                        text=item.get("text", ""),
                        author=item.get("author", ""),
                        date=item.get("date", ""),
                        source_doc=item.get("source_doc", ""),
                    )
                )

        return Question(
            id=q.get("id", ""),
            type=q.get("type", "classic"),
            text_html=q.get("text_html", "<p><br></p>"),
            title=title,
            points=int(q.get("points", 1)),
            bonus_correct=bc,
            bonus_wrong=bw,
            created_at=q.get("created_at", ""),
            correct_answer=q.get("correct_answer", ""),
            funny_answers=f_answers,
            image_path=q.get("image_path", ""),
            image_width_cm=float(q.get("image_width_cm", 0.0) or 0.0),
            image_height_cm=float(q.get("image_height_cm", 0.0) or 0.0),
            image_keep_aspect=bool(q.get("image_keep_aspect", True)),
        )

    def _serialize_group(self, g: Group) -> dict:
        return {"id": g.id, "name": g.name, "subgroups": [self._serialize_subgroup(sg) for sg in g.subgroups]}

    def _serialize_subgroup(self, sg: Subgroup) -> dict:
        return {"id": sg.id, "name": sg.name, "subgroups": [self._serialize_subgroup(s) for s in sg.subgroups], "questions": [asdict(q) for q in sg.questions]}

    # -------------------- Tree helpery --------------------

    def _bonus_points_label(self, q: Question) -> str:
        return f"+{q.bonus_correct:.2f}/ {q.bonus_wrong:.2f}"
    
    def _generate_icon(self, text: str, color: QColor, shape: str = "circle") -> QIcon:
        """Vygeneruje jednoduchou ikonu s textem/symbolem."""
        pix = QPixmap(16, 16)
        pix.fill(Qt.transparent)
        painter = QPainter(pix)
        painter.setRenderHint(QPainter.Antialiasing)
        
        painter.setBrush(color)
        painter.setPen(Qt.NoPen)
        
        if shape == "circle":
            painter.drawEllipse(1, 1, 14, 14)
        elif shape == "star":
            # Jednoduchá hvězda/diamant (kosočtverec)
            path = QPainterPath()
            path.moveTo(8, 0)
            path.lineTo(16, 8)
            path.lineTo(8, 16)
            path.lineTo(0, 8)
            path.closeSubpath()
            painter.drawPath(path)
        else:
            painter.drawRect(1, 1, 14, 14)
            
        painter.setPen(QColor("black")) # Text černý pro kontrast na světlé barvě
        font = painter.font()
        font.setBold(True)
        font.setPointSize(9)
        painter.setFont(font)
        painter.drawText(pix.rect(), Qt.AlignCenter, text)
        painter.end()
        return QIcon(pix)


    def _get_image_px_size(self, path: str) -> Optional[tuple[int, int]]:
        """Vrátí (šířka_px, výška_px) obrázku, nebo None."""
        try:
            reader = QImageReader(path)
            size = reader.size()
            if size.isValid():
                return (int(size.width()), int(size.height()))
        except Exception:
            pass
        try:
            img = QImage(path)
            if not img.isNull():
                return (int(img.width()), int(img.height()))
        except Exception:
            pass
        return None


    def _get_image_actual_size_cm(self, path: str) -> Optional[tuple[float, float, int, int, bool, float, float]]:
        """
        Vrátí (w_cm, h_cm, w_px, h_px, has_dpi, dpi_x, dpi_y) pro informativní zobrazení.
        Pokud obrázek neobsahuje DPI, použije se fallback 96 DPI.
        """
        try:
            img = QImage(path)
            if img.isNull():
                return None
            w_px = int(img.width())
            h_px = int(img.height())
            dpmx = int(img.dotsPerMeterX())
            dpmy = int(img.dotsPerMeterY())
            if dpmx > 0 and dpmy > 0:
                # cm = px * 100 / dots_per_meter
                w_cm = (w_px * 100.0) / float(dpmx)
                h_cm = (h_px * 100.0) / float(dpmy)
                dpi_x = float(dpmx) * 0.0254
                dpi_y = float(dpmy) * 0.0254
                return (float(w_cm), float(h_cm), w_px, h_px, True, float(dpi_x), float(dpi_y))
            # fallback
            dpi_x = 96.0
            dpi_y = 96.0
            w_cm = (w_px / dpi_x) * 2.54
            h_cm = (h_px / dpi_y) * 2.54
            return (float(w_cm), float(h_cm), w_px, h_px, False, float(dpi_x), float(dpi_y))
        except Exception:
            return None

    def _set_image_size_rows_visible(self, has_img: bool) -> None:
        """Zobrazí/skryje informativní a exportní rozměry obrázku v editoru."""
        if hasattr(self, "lbl_img_actual_size_label"):
            self.lbl_img_actual_size_label.setVisible(bool(has_img))
        if hasattr(self, "lbl_img_actual_size"):
            self.lbl_img_actual_size.setVisible(bool(has_img))
        if hasattr(self, "lbl_img_export_size"):
            self.lbl_img_export_size.setVisible(bool(has_img))
        if hasattr(self, "img_size_row"):
            self.img_size_row.setVisible(bool(has_img))
        if hasattr(self, "chk_img_keep_aspect"):
            self.chk_img_keep_aspect.setVisible(bool(has_img))


    def _update_current_image_ratio_cache(self, path: str) -> None:
        """Aktualizuje cache poměru stran (height/width) pro aktuální obrázek."""
        self._img_ratio_hw = None
        try:
            px = self._get_image_px_size(path) if path else None
            if px and px[0] > 0 and px[1] > 0:
                self._img_ratio_hw = float(px[1]) / float(px[0])
        except Exception:
            self._img_ratio_hw = None

    def _on_keep_aspect_changed(self, _state: int) -> None:
        """Po zapnutí 'Zachovat poměr stran' dopočítá druhý rozměr podle obrázku."""
        if getattr(self, "_img_size_sync_block", False):
            return
        if not hasattr(self, "chk_img_keep_aspect") or not self.chk_img_keep_aspect.isChecked():
            return
        full_img = getattr(self, "_current_image_full_path", "") or ""
        if not (full_img and os.path.exists(full_img)):
            return

        self._update_current_image_ratio_cache(full_img)
        if not self._img_ratio_hw or self._img_ratio_hw <= 0:
            return

        w = float(self.spin_img_w_cm.value())
        h = float(self.spin_img_h_cm.value())

        # pokud nejsou rozměry nastavené, nastavíme rozumný default
        if w <= 0.0 and h <= 0.0:
            w = 14.0
            h = w * self._img_ratio_hw

        # preferujeme šířku jako řídící
        if w > 0.0:
            self._img_size_sync_block = True
            try:
                self.spin_img_h_cm.blockSignals(True)
                self.spin_img_h_cm.setValue(round(float(w * self._img_ratio_hw), 2))
                self.spin_img_h_cm.blockSignals(False)
            finally:
                self._img_size_sync_block = False
        elif h > 0.0:
            self._img_size_sync_block = True
            try:
                self.spin_img_w_cm.blockSignals(True)
                self.spin_img_w_cm.setValue(round(float(h / self._img_ratio_hw), 2))
                self.spin_img_w_cm.blockSignals(False)
            finally:
                self._img_size_sync_block = False

    def _on_export_img_w_changed(self, value: float) -> None:
        """Když uživatel změní šířku, a je zapnutý poměr stran, dopočítá výšku."""
        if getattr(self, "_img_size_sync_block", False):
            return
        if not hasattr(self, "chk_img_keep_aspect") or not self.chk_img_keep_aspect.isChecked():
            return

        full_img = getattr(self, "_current_image_full_path", "") or ""
        if not (full_img and os.path.exists(full_img)):
            return

        self._update_current_image_ratio_cache(full_img)
        if not self._img_ratio_hw or self._img_ratio_hw <= 0:
            return

        w = float(value)
        h = 0.0 if w <= 0.0 else (w * self._img_ratio_hw)

        self._img_size_sync_block = True
        try:
            self.spin_img_h_cm.blockSignals(True)
            self.spin_img_h_cm.setValue(round(float(h), 2))
            self.spin_img_h_cm.blockSignals(False)
        finally:
            self._img_size_sync_block = False

    def _on_export_img_h_changed(self, value: float) -> None:
        """Když uživatel změní výšku, a je zapnutý poměr stran, dopočítá šířku."""
        if getattr(self, "_img_size_sync_block", False):
            return
        if not hasattr(self, "chk_img_keep_aspect") or not self.chk_img_keep_aspect.isChecked():
            return

        full_img = getattr(self, "_current_image_full_path", "") or ""
        if not (full_img and os.path.exists(full_img)):
            return

        self._update_current_image_ratio_cache(full_img)
        if not self._img_ratio_hw or self._img_ratio_hw <= 0:
            return

        h = float(value)
        w = 0.0 if h <= 0.0 else (h / self._img_ratio_hw)

        self._img_size_sync_block = True
        try:
            self.spin_img_w_cm.blockSignals(True)
            self.spin_img_w_cm.setValue(round(float(w), 2))
            self.spin_img_w_cm.blockSignals(False)
        finally:
            self._img_size_sync_block = False



    def _apply_question_item_visuals(self, item: QTreeWidgetItem, q_type: str, has_image: bool = False) -> None:
        """Aplikuje vizuální styl na položku otázky (ikona, barva, font)."""
        color_classic_bg = QColor("#42a5f5")  # Modrá
        color_bonus_bg = QColor("#ffea00")    # Žlutá
        
        is_bonus = str(q_type).lower() == "bonus" or q_type == 1

        if is_bonus:
            base_icon_char = "B"
            base_color = color_bonus_bg
            shape = "star"
            item.setForeground(0, QBrush(color_bonus_bg))
            item.setForeground(1, QBrush(color_bonus_bg))
            f = item.font(0); f.setBold(True); item.setFont(0, f)
        else:
            base_icon_char = "Q"
            base_color = color_classic_bg
            shape = "circle"
            item.setForeground(0, QBrush(color_classic_bg))
            item.setForeground(1, QBrush(color_classic_bg))
            f = item.font(0); f.setBold(False); item.setFont(0, f)

        # Generování ikony (případně kompozitní s indikátorem obrázku)
        if has_image:
            # Vytvoříme širší pixmapu pro dvě ikony vedle sebe [IMG][Q]
            pix = QPixmap(34, 16)
            pix.fill(Qt.transparent)
            painter = QPainter(pix)
            painter.setRenderHint(QPainter.Antialiasing)
            
            # 1. Ikona obrázku (vlevo)
            # Malý obdélník s naznačením "obrázku"
            painter.setBrush(QColor("#ab47bc")) # Fialová pro odlišení
            painter.setPen(Qt.NoPen)
            painter.drawRoundedRect(0, 2, 14, 12, 2, 2)
            # Symbol (kolečko uvnitř jako čočka)
            painter.setBrush(QColor("white"))
            painter.drawEllipse(4, 5, 6, 6)

            # 2. Standardní ikona (vpravo, posunutá o 18px)
            painter.setBrush(base_color)
            painter.translate(18, 0)
            
            if shape == "circle":
                painter.drawEllipse(1, 1, 14, 14)
            elif shape == "star":
                path = QPainterPath()
                path.moveTo(8, 0); path.lineTo(16, 8); path.lineTo(8, 16); path.lineTo(0, 8); path.closeSubpath()
                painter.drawPath(path)
            
            painter.setPen(QColor("black"))
            font = painter.font()
            font.setBold(True)
            font.setPointSize(9)
            painter.setFont(font)
            painter.drawText(QRect(0, 0, 16, 16), Qt.AlignCenter, base_icon_char)
            
            painter.end()
            item.setIcon(0, QIcon(pix))
        else:
            # Standardní ikona bez obrázku
            icon = self._generate_icon(base_icon_char, base_color, shape)
            item.setIcon(0, icon)

    def _refresh_tree(self) -> None:
        """Obnoví strom otázek podle self.root."""
        self.tree.clear()
        if not self.root:
            return

        sorted_groups = sorted(self.root.groups, key=lambda g: g.name.lower())
        
        color_group = QBrush(QColor("#ff5252")) 
        icon_group = self._generate_icon("S", QColor("#ff5252"), "rect")

        for g in sorted_groups:
            g_item = QTreeWidgetItem([g.name, ""])
            g_item.setData(0, Qt.UserRole, {"kind": "group", "id": g.id})
            g_item.setIcon(0, icon_group)
            
            g_item.setForeground(0, color_group)
            g_item.setForeground(1, color_group)
            f = g_item.font(0); f.setBold(True); f.setPointSize(13); g_item.setFont(0, f)
            
            self.tree.addTopLevelItem(g_item)
            g_item.setExpanded(True)

            if g.subgroups:
                sorted_subgroups = sorted(g.subgroups, key=lambda s: s.name.lower())
                self._add_subgroups_to_item(g_item, g.id, sorted_subgroups)

        # DŮLEŽITÉ: Aktualizujeme také Hall of Shame, pokud existuje
        if hasattr(self, "_refresh_funny_answers_tab"):
            self._refresh_funny_answers_tab()

    def _add_subgroups_to_item(self, parent_item: QTreeWidgetItem, group_id: str, subgroups: List[Subgroup]) -> None:
        color_subgroup = QBrush(QColor("#ff8a80"))
        # Ikona Podskupiny: Světle červený kruh/čtverec s "P"
        icon_subgroup = self._generate_icon("P", QColor("#ff8a80"), "rect")

        for sg in subgroups:
            parent_meta = parent_item.data(0, Qt.UserRole) or {}
            parent_sub_id = parent_meta.get("id") if parent_meta.get("kind") == "subgroup" else None

            sg_item = QTreeWidgetItem([sg.name, ""])
            sg_item.setData(0, Qt.UserRole, {
                "kind": "subgroup",
                "id": sg.id,
                "parent_group_id": group_id,
                "parent_subgroup_id": parent_sub_id
            })
            sg_item.setIcon(0, icon_subgroup)
            sg_item.setForeground(0, color_subgroup)
            sg_item.setForeground(1, color_subgroup)
            f = sg_item.font(0); f.setBold(True); sg_item.setFont(0, f)
            parent_item.addChild(sg_item)

            # 1. Otázky
            sorted_questions = sorted(sg.questions, key=lambda q: (q.title or "").lower())
            for q in sorted_questions:
                label = "Klasická" if str(q.type).lower() != "bonus" else "BONUS"
                is_bonus = (label == "BONUS")
                if is_bonus:
                    pts = f"+{q.bonus_correct}/{q.bonus_wrong} b."
                else:
                    pts = f"{q.points} b."

                q_item = QTreeWidgetItem([q.title or "Otázka", f"{label} | {pts}"])
                q_item.setData(0, Qt.UserRole, {
                    "kind": "question",
                    "id": q.id,
                    "parent_group_id": group_id,
                    "parent_subgroup_id": sg.id
                })
                
                # NOVÉ: Předáváme informaci o existenci obrázku
                has_img = bool(getattr(q, "image_path", "") and os.path.exists(q.image_path))
                self._apply_question_item_visuals(q_item, q.type, has_image=has_img)

                sg_item.addChild(q_item)

            # 2. Rekurze
            if sg.subgroups:
                sorted_nested_subgroups = sorted(sg.subgroups, key=lambda s: s.name.lower())
                self._add_subgroups_to_item(sg_item, group_id, sorted_nested_subgroups)
                
            sg_item.setExpanded(True)

    def _selected_node(self):
        """Vrátí (kind, meta) pro vybranou položku ve stromu."""
        sel = self.tree.selectedItems()
        if not sel:
            return None, None
        item = sel[0]
        data = item.data(0, Qt.UserRole)
        
        # Podpora pro tuple (kind, meta) - nový formát
        if isinstance(data, tuple) and len(data) >= 2:
            return data[0], data[1]
            
        # Podpora pro starý formát (dict s klíčem 'kind')
        if isinstance(data, dict):
            return data.get("kind"), data
            
        return None, None


    def _sync_model_from_tree(self) -> None:
        group_map = {g.id: g for g in self.root.groups}
        subgroup_map: dict[str, Subgroup] = {}
        question_map: dict[str, Question] = {}

        def scan_subgroups(lst: List[Subgroup]):
            for sg in lst:
                subgroup_map[sg.id] = sg
                for q in sg.questions:
                    question_map[q.id] = q
                scan_subgroups(sg.subgroups)

        for g in self.root.groups:
            scan_subgroups(g.subgroups)

        new_groups: List[Group] = []

        def build_from_item(item: QTreeWidgetItem, container) -> None:
            for i in range(item.childCount()):
                ch = item.child(i)
                meta = ch.data(0, Qt.UserRole) or {}
                kind = meta.get("kind")
                if kind == "subgroup":
                    old = subgroup_map.get(meta["id"])
                    if not old:
                        continue
                    new_sg = Subgroup(id=old.id, name=ch.text(0), subgroups=[], questions=[])
                    container.subgroups.append(new_sg)
                    build_from_item(ch, new_sg)
                elif kind == "question":
                    q = question_map.get(meta["id"])
                    if not q:
                        continue
                    if isinstance(container, Group):
                        if not container.subgroups:
                            container.subgroups.append(Subgroup(id=str(_uuid.uuid4()), name="Default", subgroups=[], questions=[]))
                        container.subgroups[0].questions.append(q)
                    else:
                        container.questions.append(q)

        for i in range(self.tree.topLevelItemCount()):
            gi = self.tree.topLevelItem(i)
            meta = gi.data(0, Qt.UserRole) or {}
            if meta.get("kind") != "group":
                continue
            old_g = group_map.get(meta["id"])
            if not old_g:
                old_g = Group(id=meta["id"], name=gi.text(0), subgroups=[])
            new_g = Group(id=old_g.id, name=gi.text(0), subgroups=[])
            build_from_item(gi, new_g)
            new_groups.append(new_g)

        self.root.groups = new_groups

    # -------------------- Akce: přidání/mazání/přejmenování --------------------

    def _add_group(self) -> None:
        from PySide6.QtWidgets import QInputDialog
        name, ok = QInputDialog.getText(self, "Nová skupina", "Název skupiny:")
        if not ok or not name.strip():
            return
        g = Group(id=str(_uuid.uuid4()), name=name.strip(), subgroups=[])
        self.root.groups.append(g)
        self._refresh_tree()
        self.save_data()

    def _add_subgroup(self) -> None:
        kind, meta = self._selected_node()
        if kind not in ("group", "subgroup"):
            QMessageBox.information(self, "Výběr", "Vyberte skupinu (nebo podskupinu) pro přidání podskupiny.")
            return
        from PySide6.QtWidgets import QInputDialog
        name, ok = QInputDialog.getText(self, "Nová podskupina", "Název podskupiny:")
        if not ok or not name.strip():
            return

        if kind == "group":
            g = self._find_group(meta["id"])
            if not g:
                return
            g.subgroups.append(Subgroup(id=str(_uuid.uuid4()), name=name.strip(), subgroups=[], questions=[]))
        else:
            parent_sg = self._find_subgroup(meta["parent_group_id"], meta["id"])
            if not parent_sg:
                return
            parent_sg.subgroups.append(Subgroup(id=str(_uuid.uuid4()), name=name.strip(), subgroups=[], questions=[]))

        self._refresh_tree()
        self.save_data()

    def _add_question(self) -> None:
        kind, meta = self._selected_node()
        if kind not in ("group", "subgroup"):
            QMessageBox.information(self, "Výběr", "Vyberte skupinu nebo podskupinu, do které chcete přidat otázku.")
            return

        target_sg: Optional[Subgroup] = None
        if kind == "group":
            g = self._find_group(meta["id"])
            if not g:
                return
            if not g.subgroups:
                sg = Subgroup(id=str(_uuid.uuid4()), name="Default", subgroups=[], questions=[])
                g.subgroups.append(sg)
                target_sg = sg
            else:
                target_sg = g.subgroups[0]
        else:
            target_sg = self._find_subgroup(meta["parent_group_id"], meta["id"])
        if not target_sg:
            return

        q = Question.new_default("classic")
        target_sg.questions.append(q)
        self._refresh_tree()
        self._select_question(q.id)
        self.save_data()

    def _delete_selected(self) -> None:
        """Deprecated: Redirects to bulk delete."""
        self._bulk_delete_selected()

    def _bulk_delete_selected(self) -> None:
        """Hromadné mazání vybraných položek (otázky, podskupiny, skupiny)."""
        items = self.tree.selectedItems()
        if not items:
            QMessageBox.information(self, "Smazat", "Vyberte položky ke smazání.")
            return
    
        count = len(items)
        msg = f"Opravdu smazat {count} vybraných položek?\n(Včetně obsahu skupin/podskupin)"
        if QMessageBox.question(self, "Smazat vybrané", msg) != QMessageBox.Yes:
            return
    
        # --- KOŠ: Nejdřív si uložíme všechny otázky, které se budou mazat ---
        if not hasattr(self.root, "trash") or not isinstance(getattr(self.root, "trash", None), list):
            self.root.trash = []
    
        now_iso = datetime.now().isoformat(timespec="seconds")
        trash_records: List[dict] = []
        seen_qids: set[str] = set()
    
        selected_question_metas: List[dict] = []
        selected_subgroup_metas: List[dict] = []
        selected_group_ids: List[str] = []
    
        # Sběr IDček k smazání
        to_delete_q_ids = set()      # otázky
        to_delete_sg_ids = set()     # podskupiny
        to_delete_g_ids = set()      # skupiny
    
        for it in items:
            meta = it.data(0, Qt.UserRole) or {}
            kind = meta.get("kind")
            if kind == "question":
                to_delete_q_ids.add(meta.get("id"))
                selected_question_metas.append(meta)
            elif kind == "subgroup":
                to_delete_sg_ids.add(meta.get("id"))
                selected_subgroup_metas.append(meta)
            elif kind == "group":
                to_delete_g_ids.add(meta.get("id"))
                selected_group_ids.append(meta.get("id"))
    
        def add_question_to_trash(q: Question, gid: str, sgid: str, gname: str, sgname: str) -> None:
            if not q or not q.id:
                return
            if q.id in seen_qids:
                return
            seen_qids.add(q.id)
            trash_records.append({
                "question": asdict(q),
                "deleted_at": now_iso,
                "source_group_id": gid or "",
                "source_group_name": gname or "",
                "source_subgroup_id": sgid or "",
                "source_subgroup_name": sgname or "",
            })
    
        def collect_questions_under_subgroups(subgroups: List[Subgroup], gid: str, gname: str) -> None:
            for sg in subgroups:
                for q in sg.questions:
                    add_question_to_trash(q, gid, sg.id, gname, sg.name)
                if sg.subgroups:
                    collect_questions_under_subgroups(sg.subgroups, gid, gname)
    
        # 1) Skupiny
        for gid in selected_group_ids:
            g = self._find_group(gid)
            if not g:
                continue
            collect_questions_under_subgroups(g.subgroups, g.id, g.name)
    
        # 2) Podskupiny (včetně vnořených)
        for meta in selected_subgroup_metas:
            gid = meta.get("parent_group_id") or ""
            sgid = meta.get("id") or ""
            g = self._find_group(gid)
            gname = g.name if g else ""
            sg = self._find_subgroup(gid, sgid)
            if not sg:
                continue
            collect_questions_under_subgroups([sg], gid, gname)
    
        # 3) Konkrétní otázky
        for meta in selected_question_metas:
            gid = meta.get("parent_group_id") or ""
            sgid = meta.get("parent_subgroup_id") or ""
            qid = meta.get("id") or ""
            if not qid:
                continue
            q = self._find_question(gid, sgid, qid)
            if not q:
                continue
            g = self._find_group(gid)
            gname = g.name if g else ""
            sg = self._find_subgroup(gid, sgid)
            sgname = sg.name if sg else ""
            add_question_to_trash(q, gid, sgid, gname, sgname)
    
        if trash_records:
            self.root.trash.extend(trash_records)
    
        # 1. Filtrace skupin (nejvyšší úroveň)
        # Pokud mažeme skupinu, zmizí vše pod ní, takže nemusíme řešit její podskupiny/otázky
        self.root.groups = [g for g in self.root.groups if g.id not in to_delete_g_ids]
    
        # 2. Procházení zbytku a mazání podskupin a otázek
        for g in self.root.groups:
            # Filtrace podskupin v této skupině
            g.subgroups = [sg for sg in g.subgroups if sg.id not in to_delete_sg_ids]
            
            # Rekurzivní čištění uvnitř podskupin (pro otázky a vnořené podskupiny)
            self._clean_subgroups_recursive(g.subgroups, to_delete_sg_ids, to_delete_q_ids)
    
        self._refresh_tree()
        self._refresh_trash_table()
        self._clear_editor()
        self.save_data()
        self.statusBar().showMessage(f"Smazáno {count} položek.", 4000)

    def _clean_subgroups_recursive(self, subgroups: List[Subgroup], delete_sg_ids: set, delete_q_ids: set) -> None:
        """Pomocná metoda pro rekurzivní čištění."""
        for sg in subgroups:
            # Smazání otázek v aktuální podskupině
            if delete_q_ids:
                sg.questions = [q for q in sg.questions if q.id not in delete_q_ids]
            
            # Filtrace vnořených podskupin (pokud existují)
            if sg.subgroups:
                sg.subgroups = [s for s in sg.subgroups if s.id not in delete_sg_ids]
                # Rekurze do hloubky
                self._clean_subgroups_recursive(sg.subgroups, delete_sg_ids, delete_q_ids)

    def _on_rename_clicked(self) -> None:
        kind, meta = self._selected_node()
        if kind not in ("group", "subgroup"):
            return
        new_name = self.rename_line.text().strip()
        if not new_name:
            return
        if kind == "group":
            g = self._find_group(meta["id"])
            if g:
                g.name = new_name
        else:
            sg = self._find_subgroup(meta["parent_group_id"], meta["id"])
            if sg:
                sg.name = new_name
        self._refresh_tree()
        self.save_data()

    def _on_tree_selection_changed(self) -> None:
        # Zajistíme, že během načítání nebude aktivní žádné ID, aby se nespouštěl autosave
        self._current_question_id = None
        
        kind, meta = self._selected_node()
        self._current_node_kind = kind

        if kind == "question":
            q = self._find_question(meta["parent_group_id"], meta["parent_subgroup_id"], meta["id"])
            if q:
                # Načteme data do editoru (ID zůstává None)
                self._load_question_to_editor(q)
                # Nastavíme viditelnost (to může vyvolat změny hodnot přes _on_type_changed_ui, ale ID je None, takže se neuloží)
                self._set_question_editor_visible(True)
                self.rename_panel.hide()
                self._set_editor_enabled(True)
                # Teprve až je vše načteno a UI nastaveno, povolíme ID pro budoucí uživatelské úpravy
                self._current_question_id = q.id

        elif kind in ("group", "subgroup"):
            name = ""
            if kind == "group":
                g = self._find_group(meta["id"]); name = g.name if g else ""
            else:
                sg = self._find_subgroup(meta["parent_group_id"], meta["id"]); name = sg.name if sg else ""
            
            self.rename_line.blockSignals(True)
            self.rename_line.setText(name)
            self.rename_line.blockSignals(False)
            
            self._set_question_editor_visible(False)
            self.rename_panel.show()
            self._set_editor_enabled(False)

        else:
            self._clear_editor()
            self._set_question_editor_visible(False)
            self.rename_panel.hide()

    def _clear_editor(self) -> None:
        self._current_question_id = None
        
        widgets = [
            self.text_edit,
            self.spin_points,
            self.spin_bonus_correct,
            self.spin_bonus_wrong,
            self.combo_type,
            self.title_edit,
            self.image_path_edit,
            self.edit_correct_answer,
            self.table_funny
        ]
        
        for w in widgets:
            w.blockSignals(True)

        try:
            self.text_edit.clear()
            self.spin_points.setValue(1)
            self.spin_bonus_correct.setValue(1.00)
            self.spin_bonus_wrong.setValue(0.00)
            self.combo_type.setCurrentIndex(0)
            self.title_edit.clear()
            self.image_path_edit.clear()
            self.image_path_edit.setToolTip("")
            if hasattr(self, "spin_img_w_cm"):
                self.spin_img_w_cm.setValue(0.0)
                self.spin_img_h_cm.setValue(0.0)
                self.spin_img_w_cm.setEnabled(False)
                self.spin_img_h_cm.setEnabled(False)
                if hasattr(self, "chk_img_keep_aspect"):
                    self.chk_img_keep_aspect.setChecked(True)
                    self.chk_img_keep_aspect.setEnabled(False)
            if hasattr(self, "lbl_img_actual_size"):
                self.lbl_img_actual_size.setText("")
            if hasattr(self, "_set_image_size_rows_visible"):
                self._set_image_size_rows_visible(False)
            self.edit_correct_answer.clear() 
            self.table_funny.setRowCount(0) 
        finally:
            for w in widgets:
                w.blockSignals(False)
                
        self._set_editor_enabled(False)

    def _set_question_editor_visible(self, visible: bool) -> None:
        """Zobrazí nebo skryje kompletní editor otázky (toolbar, formulář, text)."""
        self.editor_toolbar.setVisible(visible)
        self.text_edit.setVisible(visible)
        self.btn_save_question.setVisible(visible)

        # Skrytí/Zobrazení prvků formuláře
        widgets = [
            self.title_edit,
            self.image_path_edit,
            self.btn_choose_image,
            self.btn_clear_image,
            self.combo_type,
            self.spin_points,
            self.spin_bonus_correct,
            self.spin_bonus_wrong,
            self.edit_correct_answer,
            self.funny_container,
            self.lbl_content,
            self.lbl_correct,
            self.lbl_funny
        ]
        
        # Přidáme i label náhledu, pokud existuje
        if hasattr(self, "lbl_image_preview"):
            # Zobrazíme jen pokud je 'visible' True A máme co zobrazit (kontroluje se v load_question)
            # Ale pro jistotu: pokud visible=False, skryjeme určitě.
            if not visible:
                self.lbl_image_preview.setVisible(False)
            # Pokud visible=True, stav labelu řídí _load_question_to_editor, takže neměníme.

        for w in widgets:
            if hasattr(self, w.objectName()) or w in widgets:
                w.setVisible(visible)

        # Skrytí labelů ve form layoutu
        for i in range(self.form_layout.rowCount()):
            item = self.form_layout.itemAt(i, QFormLayout.LabelRole)
            if item and item.widget():
                item.widget().setVisible(visible)
            item = self.form_layout.itemAt(i, QFormLayout.FieldRole)
            if item and item.widget():
                # Specifická výjimka pro náhled - ten si řídí viditelnost sám podle obsahu
                if hasattr(self, "lbl_image_preview") and item.widget() == self.lbl_image_preview:
                    if not visible:
                        item.widget().setVisible(False)
                    continue
                item.widget().setVisible(visible)

        if visible:
            self._on_type_changed_ui()
            # Zajistit, že informace/rozměry obrázku se zobrazí jen pokud otázka skutečně má obrázek
            full_img = getattr(self, "_current_image_full_path", "") or ""
            self._set_image_size_rows_visible(bool(full_img and os.path.exists(full_img)))

    def _load_question_to_editor(self, q: Question) -> None:
        # ID nulujeme i zde pro jistotu
        self._current_question_id = None
        
        # Uložení plné cesty k obrázku bokem (protože v GUI ukazujeme jen název)
        self._current_image_full_path = getattr(q, "image_path", "") or ""

        widgets = [
            self.combo_type,
            self.spin_points,
            self.spin_bonus_correct,
            self.spin_bonus_wrong,
            self.text_edit,
            self.title_edit,
            self.image_path_edit,
            self.spin_img_w_cm,
            self.spin_img_h_cm,
            self.chk_img_keep_aspect,
            self.edit_correct_answer,
            self.table_funny
        ]

        for w in widgets:
            w.blockSignals(True)

        try:
            self.combo_type.setCurrentIndex(0 if q.type == "classic" else 1)
            self.spin_points.setValue(int(q.points))
            self.spin_bonus_correct.setValue(float(q.bonus_correct))
            self.spin_bonus_wrong.setValue(float(q.bonus_wrong))
            self.text_edit.setHtml(q.text_html or "")
            self.title_edit.setText(q.title or self._derive_title_from_html(q.text_html))
            
            # NOVÉ: Zobrazit jen název souboru
            full_path = getattr(q, "image_path", "") or ""
            if full_path:
                self.image_path_edit.setText(os.path.basename(full_path))
                self.image_path_edit.setToolTip(full_path) # Full path v tooltipu
            else:
                self.image_path_edit.clear()
                self.image_path_edit.setToolTip("")

            # Rozměry obrázku (cm) – pokud nejsou uložené, nastavíme defaultně 14 cm na šířku a dopočítáme výšku dle poměru
            w_cm = float(getattr(q, "image_width_cm", 0.0) or 0.0)
            h_cm = float(getattr(q, "image_height_cm", 0.0) or 0.0)
            if full_path and os.path.exists(full_path):
                if hasattr(self, "_update_current_image_ratio_cache"):
                    self._update_current_image_ratio_cache(full_path)
                px = self._get_image_px_size(full_path)
                if w_cm <= 0.0:
                    w_cm = 14.0
                if h_cm <= 0.0 and px and px[0] > 0:
                    h_cm = round(w_cm * (px[1] / px[0]), 2)

                self.spin_img_w_cm.setValue(w_cm)
                self.spin_img_h_cm.setValue(h_cm)
                self.spin_img_w_cm.setEnabled(True)
                self.spin_img_h_cm.setEnabled(True)
                if hasattr(self, "chk_img_keep_aspect"):
                    self.chk_img_keep_aspect.setEnabled(True)
                    self.chk_img_keep_aspect.setChecked(bool(getattr(q, "image_keep_aspect", True)))

                # Informativní: aktuální (fyzická) velikost dle DPI v souboru (nebo fallback)
                if hasattr(self, "lbl_img_actual_size"):
                    info = self._get_image_actual_size_cm(full_path)
                    if info:
                        w_act, h_act, w_px, h_px, has_dpi, dpi_x, dpi_y = info
                        src = f"DPI {dpi_x:.0f}×{dpi_y:.0f}" if has_dpi else "DPI default 96"
                        self.lbl_img_actual_size.setText(f"{w_act:.2f} × {h_act:.2f} cm ({w_px}×{h_px} px, {src})")
                    else:
                        self.lbl_img_actual_size.setText("")

                self._set_image_size_rows_visible(True)
            else:
                self.spin_img_w_cm.setValue(0.0)
                self.spin_img_h_cm.setValue(0.0)
                self.spin_img_w_cm.setEnabled(False)
                self.spin_img_h_cm.setEnabled(False)

                if hasattr(self, "lbl_img_actual_size"):
                    self.lbl_img_actual_size.setText("")

                self._set_image_size_rows_visible(False)

            self.edit_correct_answer.setPlainText(q.correct_answer or "")

            self.table_funny.setRowCount(0)
            f_answers = getattr(q, "funny_answers", []) or []
            for fa in f_answers:
                if isinstance(fa, FunnyAnswer):
                    text = fa.text; date = fa.date; author = fa.author; source_doc = fa.source_doc
                else:
                    text = fa.get("text", ""); date = fa.get("date", ""); author = fa.get("author", ""); source_doc = fa.get("source_doc", "")
                
                row = self.table_funny.rowCount()
                self.table_funny.insertRow(row)
                self.table_funny.setItem(row, 0, QTableWidgetItem(text))
                self.table_funny.setItem(row, 1, QTableWidgetItem(date))
                self.table_funny.setItem(row, 2, QTableWidgetItem(author))
                display_source = os.path.basename(source_doc) if source_doc else ""
                source_item = QTableWidgetItem(display_source)
                source_item.setData(Qt.UserRole, source_doc)
                self.table_funny.setItem(row, 3, source_item)

            # NOVÉ: Náhled obrázku
            # Lazy init labelu pro náhled, pokud neexistuje
            if not hasattr(self, "lbl_image_preview"):
                self.lbl_image_preview = QLabel()
                self.lbl_image_preview.setAlignment(Qt.AlignCenter)
                self.lbl_image_preview.setStyleSheet("border: 1px solid #444; background: #222; border-radius: 4px; margin-top: 5px;")
                self.lbl_image_preview.setMinimumHeight(150)
                # Přidáme do form layoutu (nakonec)
                self.form_layout.addRow("Náhled:", self.lbl_image_preview)

            if full_path and os.path.exists(full_path):
                pix = QPixmap(full_path)
                if not pix.isNull():
                    # Škálování s poměrem stran
                    scaled = pix.scaled(QSize(400, 200), Qt.KeepAspectRatio, Qt.SmoothTransformation)
                    self.lbl_image_preview.setPixmap(scaled)
                    self.lbl_image_preview.setVisible(True)
                    # Zvětšíme okno, pokud je malé, aby se náhled vešel
                    if self.height() < 850:
                        self.resize(self.width(), 850)
                else:
                    self.lbl_image_preview.setText("Chyba načítání")
                    self.lbl_image_preview.setVisible(True)
            else:
                self.lbl_image_preview.clear()
                self.lbl_image_preview.hide()

            self._set_editor_enabled(True)
            # Toto volání aktualizuje UI, ale protože ID je None, neuloží se nic
            self._on_type_changed_ui()

        finally:
            for w in widgets:
                w.blockSignals(False)

    def _apply_editor_to_current_question(self, silent: bool = False) -> None:
        if not self._current_question_id:
            return

        def apply_in(sgs: List[Subgroup]) -> bool:
            for sg in sgs:
                # 1. Prohledání otázek v aktuální podskupině
                for i, q in enumerate(sg.questions):
                    if q.id == self._current_question_id:
                        q.type = "classic" if self.combo_type.currentIndex() == 0 else "bonus"
                        q.text_html = self.text_edit.toHtml()
                        q.title = (
                            self.title_edit.text().strip()
                            or self._derive_title_from_html(
                                q.text_html, 
                                prefix=("BONUS: " if q.type == "bonus" else "")
                            )
                        )
                        
                        # Uložení bodů
                        if q.type == "classic":
                            q.points = int(self.spin_points.value())
                            q.bonus_correct = 0.0
                            q.bonus_wrong = 0.0
                        else:
                            q.points = 0
                            q.bonus_correct = round(float(self.spin_bonus_correct.value()), 2)
                            q.bonus_wrong = round(float(self.spin_bonus_wrong.value()), 2)

                        # Uložení správné odpovědi
                        q.correct_answer = self.edit_correct_answer.toPlainText()

                        # Uložení cesty k obrázku (s kontrolou na basename)
                        editor_txt = self.image_path_edit.text().strip()
                        stored_full = getattr(self, "_current_image_full_path", "")
                        
                        final_path = ""
                        # Pokud uživatel nezměnil text (je stále basename původní cesty), zachováme full path
                        if stored_full and editor_txt == os.path.basename(stored_full):
                            final_path = stored_full
                        else:
                            # Uživatel něco napsal/vybral -> použijeme to
                            final_path = editor_txt
                            # Aktualizujeme stored path pro příští uložení
                            self._current_image_full_path = final_path
                        
                        q.image_path = final_path

                        # Rozměry obrázku (cm) pro export do DOCX
                        if hasattr(self, "chk_img_keep_aspect"):
                            q.image_keep_aspect = bool(self.chk_img_keep_aspect.isChecked())
                        else:
                            q.image_keep_aspect = True
                        if final_path and os.path.exists(final_path):
                            q.image_width_cm = float(self.spin_img_w_cm.value())
                            q.image_height_cm = float(self.spin_img_h_cm.value())
                        else:
                            q.image_width_cm = 0.0
                            q.image_height_cm = 0.0

                        # --- NOVÉ: OKAMŽITÁ AKTUALIZACE NÁHLEDU ---
                        if hasattr(self, "lbl_image_preview"):
                            if final_path and os.path.exists(final_path):
                                pix = QPixmap(final_path)
                                if not pix.isNull():
                                    scaled = pix.scaled(QSize(400, 200), Qt.KeepAspectRatio, Qt.SmoothTransformation)
                                    self.lbl_image_preview.setPixmap(scaled)
                                    self.lbl_image_preview.setVisible(True)
                                else:
                                    self.lbl_image_preview.setText("Chyba načítání")
                                    self.lbl_image_preview.setVisible(True)
                            else:
                                self.lbl_image_preview.clear()
                                self.lbl_image_preview.hide()
                        # ------------------------------------------

                        # NOVÉ: Informativní rozměr + viditelnost řádků (při ruční editaci cesty)
                        has_img = bool(final_path and os.path.exists(final_path))
                        if hasattr(self, "lbl_img_actual_size"):
                            if has_img:
                                info = self._get_image_actual_size_cm(final_path)
                                if info:
                                    w_act, h_act, w_px, h_px, has_dpi, dpi_x, dpi_y = info
                                    src = f"DPI {dpi_x:.0f}×{dpi_y:.0f}" if has_dpi else "DPI default 96"
                                    self.lbl_img_actual_size.setText(f"{w_act:.2f} × {h_act:.2f} cm ({w_px}×{h_px} px, {src})")
                                else:
                                    self.lbl_img_actual_size.setText("")
                            else:
                                self.lbl_img_actual_size.setText("")
                        self._set_image_size_rows_visible(has_img)

                        # Uložení vtipných odpovědí z tabulky
                        new_funny: List[FunnyAnswer] = []
                        for r in range(self.table_funny.rowCount()):
                            t_item = self.table_funny.item(r, 0)
                            if not t_item: continue
                            d_item = self.table_funny.item(r, 1)
                            a_item = self.table_funny.item(r, 2)
                            s_item = self.table_funny.item(r, 3) if self.table_funny.columnCount() > 3 else None
                            
                            text = t_item.text()
                            date = d_item.text() if d_item else ""
                            author = a_item.text() if a_item else ""
                            
                            if s_item is not None:
                                data = s_item.data(Qt.UserRole)
                                if isinstance(data, str) and data:
                                    source_doc = data
                                else:
                                    source_doc = s_item.text()
                            else:
                                source_doc = ""
                            
                            new_funny.append(
                                FunnyAnswer(
                                    text=text,
                                    author=author,
                                    date=date,
                                    source_doc=source_doc
                                )
                            )

                        q.funny_answers = new_funny
                        sg.questions[i] = q

                        label = "Klasická" if q.type == "classic" else "BONUS"
                        pts = q.points if q.type == "classic" else self._bonus_points_label(q)
                        
                        self._update_selected_question_item_title(q.title)
                        self._update_selected_question_item_subtitle(f"{label} | {pts}")
                        
                        items = self.tree.selectedItems()
                        if items:
                            # Aktualizace vizuálu položky (ikona)
                            has_img = bool(q.image_path and os.path.exists(q.image_path))
                            self._apply_question_item_visuals(items[0], q.type, has_image=has_img)

                        if not silent:
                            self.statusBar().showMessage("Změny otázky uloženy (lokálně).", 1200)
                        
                        self._refresh_funny_answers_tab()
                        return True
            
                # 2. Rekurzivní hledání v podskupinách
                if apply_in(sg.subgroups):
                    return True

            return False

        for g in self.root.groups:
            if apply_in(g.subgroups):
                break

    def _clear_editor(self) -> None:
        self._current_question_id = None
        self._current_image_full_path = ""  # Vyčistit cache cesty

        widgets = [
            self.text_edit,
            self.spin_points,
            self.spin_bonus_correct,
            self.spin_bonus_wrong,
            self.combo_type,
            self.title_edit,
            self.image_path_edit,
            self.edit_correct_answer,
            self.table_funny
        ]

        for w in widgets:
            w.blockSignals(True)

        try:
            self.text_edit.clear()
            self.spin_points.setValue(1)
            self.spin_bonus_correct.setValue(1.00)
            self.spin_bonus_wrong.setValue(0.00)
            self.combo_type.setCurrentIndex(0)
            self.title_edit.clear()
            self.image_path_edit.clear()
            self.edit_correct_answer.clear()
            self.table_funny.setRowCount(0)
            
            # NOVÉ: Vyčistit náhled
            if hasattr(self, "lbl_image_preview"):
                self.lbl_image_preview.clear()
                self.lbl_image_preview.hide()

            # NOVÉ: Vyčistit rozměry obrázku (info + export)
            self.image_path_edit.setToolTip("")
            if hasattr(self, "spin_img_w_cm"):
                self.spin_img_w_cm.setValue(0.0)
                self.spin_img_h_cm.setValue(0.0)
                self.spin_img_w_cm.setEnabled(False)
                self.spin_img_h_cm.setEnabled(False)
            if hasattr(self, "lbl_img_actual_size"):
                self.lbl_img_actual_size.setText("")
            self._set_image_size_rows_visible(False)

        finally:
            for w in widgets:
                w.blockSignals(False)
        
        self._set_editor_enabled(False)
            
    def _autosave_schedule(self) -> None:
        if not self._current_question_id:
            return
        self._autosave_timer.stop(); self._autosave_timer.start()

    def _autosave_current_question(self) -> None:
        if not self._current_question_id:
            return
        self._apply_editor_to_current_question(silent=True)
        self.save_data()

    def _on_save_question_clicked(self) -> None:
        self._apply_editor_to_current_question(silent=True)
        self.save_data()
        self.statusBar().showMessage("Otázka uložena.", 1500)

    def _update_selected_question_item_title(self, text: str) -> None:
        items = self.tree.selectedItems()
        if not items:
            return

        item = items[0]
        item.setText(0, text or "Otázka")

        # Po změně názvu otázky znovu seřadíme otázky v rámci stejné podskupiny
        parent = item.parent()
        if parent is not None:
            parent.sortChildren(0, Qt.AscendingOrder)

    def _update_selected_question_item_subtitle(self, text: str) -> None:
        items = self.tree.selectedItems()
        if items: items[0].setText(1, text)

    # -------------------- Vyhledávače --------------------

    def _find_group(self, gid: str) -> Optional[Group]:
        for g in self.root.groups:
            if g.id == gid:
                return g
        return None

    def _find_subgroup(self, gid: str, sgid: Optional[str]) -> Optional[Subgroup]:
        g = self._find_group(gid)
        if not g: return None
        def rec(lst: List[Subgroup]) -> Optional[Subgroup]:
            for sg in lst:
                if sg.id == sgid:
                    return sg
                found = rec(sg.subgroups)
                if found: return found
            return None
        return rec(g.subgroups)

    def _find_question(self, gid: str, sgid: Optional[str], qid: str) -> Optional[Question]:
        sg = self._find_subgroup(gid, sgid)
        if not sg: return None
        for q in sg.questions:
            if q.id == qid: return q
        return None

    def _find_question_by_id(self, qid: str) -> Optional[Question]:
        def walk(lst: List[Subgroup]) -> Optional[Question]:
            for sg in lst:
                for q in sg.questions:
                    if q.id == qid:
                        return q
                r = walk(sg.subgroups)
                if r: return r
            return None
        for g in self.root.groups:
            r = walk(g.subgroups)
            if r: return r
        return None

    def _select_question(self, qid: str) -> None:
        def _walk(item: QTreeWidgetItem) -> Optional[QTreeWidgetItem]:
            meta = item.data(0, Qt.UserRole)
            if meta and meta.get("kind") == "question" and meta.get("id") == qid:
                return item
            for i in range(item.childCount()):
                found = _walk(item.child(i))
                if found: return found
            return None
        for i in range(self.tree.topLevelItemCount()):
            found = _walk(self.tree.topLevelItem(i))
            if found:
                self.tree.setCurrentItem(found); self.tree.scrollToItem(found)
                break

    # -------------------- Formátování Rich text --------------------

    def _merge_format_on_selection(self, fmt: QTextCharFormat) -> None:
        cursor = self.text_edit.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.WordUnderCursor)
        cursor.mergeCharFormat(fmt)
        self.text_edit.mergeCurrentCharFormat(fmt)

    def _toggle_format(self, which: str) -> None:
        fmt = QTextCharFormat()
        if which == "bold":
            new_weight = QFont.Weight.Bold if self.action_bold.isChecked() else QFont.Weight.Normal
            fmt.setFontWeight(int(new_weight))
        elif which == "italic":
            fmt.setFontItalic(self.action_italic.isChecked())
        elif which == "underline":
            fmt.setFontUnderline(self.action_underline.isChecked())
        self._merge_format_on_selection(fmt); self._autosave_schedule()

    def _choose_color(self) -> None:
        color = QColorDialog.getColor(parent=self, title="Vyberte barvu textu")
        if color.isValid():
            fmt = QTextCharFormat(); fmt.setForeground(color)
            self._merge_format_on_selection(fmt); self._autosave_schedule()

    def _toggle_bullets(self) -> None:
        """Přepne aktuální výběr na odrážky s lepším odsazením."""
        cursor = self.text_edit.textCursor()
        cursor.beginEditBlock()
        
        # Zjistíme, zda už jsme v listu (podle prvního bloku ve výběru)
        current_list = cursor.currentList()
        
        if current_list:
            # Zrušit list (nastavit styl na Undefined, což ho odstraní)
            block_fmt = QTextBlockFormat()
            block_fmt.setObjectIndex(-1) # Zruší vazbu na list
            cursor.setBlockFormat(block_fmt)
        else:
            # Vytvořit nový list s lepším formátováním
            list_fmt = QTextListFormat()
            list_fmt.setStyle(QTextListFormat.ListDisc)
            list_fmt.setIndent(1) # Level 1
            
            # Odsazení čísla/odrážky
            # Ve Wordu/Docx to odpovídá Hanging Indent
            
            cursor.createList(list_fmt)
            
            # Aplikovat odsazení bloku pro vizuální shodu
            bf = cursor.blockFormat()
            # Nastavíme levý margin (celé odsuneme) a text indent (první řádek vrátíme zpět pro odrážku)
            # Hodnoty jsou v px/pt (záleží na DPI, ale 20/-15 je rozumný start)
            bf.setLeftMargin(20)
            bf.setTextIndent(-15) 
            cursor.setBlockFormat(bf)
            
        cursor.endEditBlock()
        self._autosave_schedule()
        self.text_edit.setFocus()


    def _apply_alignment(self, align_flag: Qt.AlignmentFlag) -> None:
        cursor = self.text_edit.textCursor()
        bf = QTextBlockFormat(); bf.setAlignment(align_flag)
        if cursor.hasSelection():
            start = cursor.selectionStart(); end = cursor.selectionEnd()
            cursor.beginEditBlock()
            tmp = QTextCursor(self.text_edit.document()); tmp.setPosition(start)
            while tmp.position() <= end:
                block_cursor = QTextCursor(tmp.block())
                block_format = block_cursor.blockFormat()
                block_format.setAlignment(align_flag)
                block_cursor.setBlockFormat(block_format)
                if tmp.block().position() + tmp.block().length() > end:
                    break
                tmp.movePosition(QTextCursor.NextBlock)
            cursor.endEditBlock()
        else:
            cursor.mergeBlockFormat(bf)
        self.text_edit.setTextCursor(cursor)
        self._sync_toolbar_to_cursor()
        self._autosave_schedule()

    def _sync_toolbar_to_cursor(self) -> None:
        fmt = self.text_edit.currentCharFormat()
        try:
            self.action_bold.setChecked(int(fmt.fontWeight()) == int(QFont.Weight.Bold))
        except Exception:
            self.action_bold.setChecked(False)
        self.action_italic.setChecked(fmt.fontItalic())
        self.action_underline.setChecked(fmt.fontUnderline())
        in_list = self.text_edit.textCursor().block().textList() is not None
        self.action_bullets.setChecked(in_list)
        align = self.text_edit.textCursor().blockFormat().alignment()
        self.action_align_left.setChecked(bool(align & Qt.AlignLeft))
        self.action_align_center.setChecked(bool(align & Qt.AlignHCenter))
        self.action_align_right.setChecked(bool(align & Qt.AlignRight))
        self.action_align_justify.setChecked(bool(align & Qt.AlignJustify))

    def _on_type_changed_ui(self) -> None:
        is_classic = self.combo_type.currentIndex() == 0

        # Skrytí/zobrazení polí a jejich popisků
        self.spin_points.setVisible(is_classic)
        lbl_points = self.form_layout.labelForField(self.spin_points)
        if lbl_points: 
            lbl_points.setVisible(is_classic)

        self.spin_bonus_correct.setVisible(not is_classic)
        lbl_bonus_correct = self.form_layout.labelForField(self.spin_bonus_correct)
        if lbl_bonus_correct: 
            lbl_bonus_correct.setVisible(not is_classic)

        self.spin_bonus_wrong.setVisible(not is_classic)
        lbl_bonus_wrong = self.form_layout.labelForField(self.spin_bonus_wrong)
        if lbl_bonus_wrong: 
            lbl_bonus_wrong.setVisible(not is_classic)

        self._autosave_schedule()


    # -------------------- Výběr datového souboru --------------------

    def _choose_data_file(self) -> None:
        new_path, _ = QFileDialog.getSaveFileName(self, "Zvolit/uložit JSON s otázkami", str(self.data_path), "JSON (*.json)")
        if new_path:
            self.data_path = Path(new_path)
            self.statusBar().showMessage(f"Datový soubor změněn na: {self.data_path}", 4000)
            self.load_data(); self._refresh_tree()

    def _choose_question_image(self) -> None:
        """Vybere obrázek k aktuální otázce (uloží se pouze cesta)."""
        if not getattr(self, "_current_question_id", None):
            return
        start_dir = str(getattr(self, "images_dir", self.project_root))
        Path(start_dir).mkdir(parents=True, exist_ok=True)
        fn, _ = QFileDialog.getOpenFileName(
            self,
            "Vybrat obrázek k otázce",
            start_dir,
            "Images (*.jpg *.jpeg *.png *.bmp *.tif *.tiff *.heic);;All files (*)",
        )
        if fn:
            # Uložíme full-path i bokem (v editoru se drží jen basename)
            self._current_image_full_path = fn
            if hasattr(self, "_update_current_image_ratio_cache"):
                self._update_current_image_ratio_cache(fn)

            # V editoru zobrazíme jen název souboru, plnou cestu necháme v tooltipu
            self.image_path_edit.setText(os.path.basename(fn))
            self.image_path_edit.setToolTip(fn)

            # Informativní: aktuální velikost dle DPI v souboru (nebo fallback)
            if hasattr(self, "lbl_img_actual_size"):
                info = self._get_image_actual_size_cm(fn)
                if info:
                    w_act, h_act, w_px, h_px, has_dpi, dpi_x, dpi_y = info
                    src = f"DPI {dpi_x:.0f}×{dpi_y:.0f}" if has_dpi else "DPI default 96"
                    self.lbl_img_actual_size.setText(f"{w_act:.2f} × {h_act:.2f} cm ({w_px}×{h_px} px, {src})")
                else:
                    self.lbl_img_actual_size.setText("")

            self._set_image_size_rows_visible(True)

            # Default: 14 cm na šířku, výšku dopočítáme dle poměru stran (aby to hned dávalo smysl)
            px = self._get_image_px_size(fn)
            w_cm = 14.0
            h_cm = 0.0
            if px and px[0] > 0:
                h_cm = round(w_cm * (px[1] / px[0]), 2)

            if hasattr(self, "spin_img_w_cm"):
                self.spin_img_w_cm.setEnabled(True)
                self.spin_img_h_cm.setEnabled(True)
                if hasattr(self, "chk_img_keep_aspect"):
                    self.chk_img_keep_aspect.setEnabled(True)
                self.spin_img_w_cm.setValue(w_cm)
                self.spin_img_h_cm.setValue(h_cm)

            self._autosave_schedule()

    def _clear_question_image(self) -> None:
        """Odebere obrázek z aktuální otázky."""
        if not getattr(self, "_current_question_id", None):
            return
        self._current_image_full_path = ""
        self.image_path_edit.clear()
        self.image_path_edit.setToolTip("")

        if hasattr(self, "spin_img_w_cm"):
            self.spin_img_w_cm.setValue(0.0)
            self.spin_img_h_cm.setValue(0.0)
            self.spin_img_w_cm.setEnabled(False)
            self.spin_img_h_cm.setEnabled(False)
        if hasattr(self, "chk_img_keep_aspect"):
            self.chk_img_keep_aspect.setChecked(True)
            self.chk_img_keep_aspect.setEnabled(False)

        if hasattr(self, "lbl_img_actual_size"):
            self.lbl_img_actual_size.setText("")
        self._set_image_size_rows_visible(False)

        self._autosave_schedule()

    # -------------------- Filtr --------------------

    def _apply_filter(self, text: str) -> None:
        pat = (text or '').strip().lower()
        def question_matches(qid: str) -> bool:
            q = None
            for g in self.root.groups:
                stack = list(g.subgroups)
                while stack:
                    sg = stack.pop()
                    for qq in sg.questions:
                        if qq.id == qid:
                            q = qq; break
                    if q: break
                    stack.extend(sg.subgroups)
                if q: break
            if not q: return False
            plain = re.sub(r'<[^>]+>', ' ', q.text_html)
            plain = _html.unescape(plain).lower()
            title = (q.title or '').lower()
            return (pat in title) or (pat in plain)
        def apply_item(item) -> bool:
            meta = item.data(0, Qt.UserRole) or {}
            kind = meta.get('kind')
            any_child = False
            for i in range(item.childCount()):
                if apply_item(item.child(i)):
                    any_child = True
            self_match = False
            if not pat:
                self_match = True
            elif kind in ('group', 'subgroup'):
                self_match = pat in item.text(0).lower()
            elif kind == 'question':
                self_match = question_matches(meta.get('id'))
            show = self_match or any_child
            item.setHidden(not show)
            return show
        for i in range(self.tree.topLevelItemCount()):
            apply_item(self.tree.topLevelItem(i))

    # -------------------- Import z DOCX --------------------

    def _read_numbering_map(self, z: zipfile.ZipFile) -> tuple[dict[int, int], dict[tuple[int, int], str]]:
        num_to_abstract: dict[int, int] = {}
        fmt_map: dict[tuple[int, int], str] = {}
        try:
            with z.open("word/numbering.xml") as f:
                xml = f.read()
        except KeyError:
            return num_to_abstract, fmt_map
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = ET.fromstring(xml)
        for num in root.findall(".//w:num", ns):
            nid = num.find("w:abstractNumId", ns)
            val = nid.get("{%s}val" % ns["w"]) if nid is not None else None
            num_id_el = num.find("w:numId", ns)
            num_id_val = num_id_el.get("{%s}val" % ns["w"]) if num_id_el is not None else None
            if val is not None and num_id_val is not None:
                try:
                    num_to_abstract[int(num_id_val)] = int(val)
                except Exception:
                    pass
        for absn in root.findall(".//w:abstractNum", ns):
            aid_attr = absn.get("{%s}abstractNumId" % ns["w"])
            if aid_attr is None:
                continue
            try:
                abs_id = int(aid_attr)
            except Exception:
                continue
            for lvl in absn.findall("w:lvl", ns):
                ilvl_attr = lvl.get("{%s}ilvl" % ns["w"])
                if ilvl_attr is None:
                    continue
                try:
                    ilvl = int(ilvl_attr)
                except Exception:
                    ilvl = 0
                numfmt = lvl.find("w:numFmt", ns)
                fmt = numfmt.get("{%s}val" % ns["w"]) if numfmt is not None else "decimal"
                fmt_map[(abs_id, ilvl)] = fmt
        return num_to_abstract, fmt_map

    def _extract_paragraphs_from_docx(self, path: Path) -> List[dict]:
        if 'docx' not in sys.modules:
            QMessageBox.critical(self, "Chyba", "Knihovna python-docx není nainstalována.")
            return []

        doc = docx.Document(path)
        out = []

        # Cache pro numbering definitions
        numbering_cache = {} # numId -> (abstractId, {ilvl: fmt})

        # Přednačtení definic z XML, abychom to nemuseli lovit per paragraph
        try:
            numbering_part = doc.part.numbering_part
            if numbering_part:
                # Mapování abstractNumId -> {ilvl: fmt}
                abstract_formats = {}
                for abstract_id, abstract in numbering_part.numbering_definitions._abstract_nums.items():
                    levels = {}
                    for lvl in abstract.levels:
                        levels[lvl.ilvl] = lvl.num_fmt
                    abstract_formats[abstract_id] = levels
                
                # Mapování numId -> abstractNumId
                for num_id, num in numbering_part.numbering_definitions._nums.items():
                    if num.abstractNumId in abstract_formats:
                        numbering_cache[num_id] = (num.abstractNumId, abstract_formats[num.abstractNumId])
        except Exception:
            pass # Pokud selže přístup k internals, pojedeme bez formátů

        def get_numbering(p: Paragraph):
            p_element = p._p
            pPr = p_element.pPr
            if pPr is None: return None, None, None
            numPr = pPr.numPr
            if numPr is None: return None, None, None
            
            val_ilvl = numPr.ilvl.val if numPr.ilvl is not None else 0
            val_numId = numPr.numId.val if numPr.numId is not None else None
            
            fmt = "decimal" # default fallback
            
            if val_numId in numbering_cache:
                _, levels = numbering_cache[val_numId]
                if val_ilvl in levels:
                    fmt = levels[val_ilvl]
            
            return val_numId, val_ilvl, fmt

        for p in doc.paragraphs:
            txt = p.text.strip()
            numId, ilvl, fmt = get_numbering(p)
            is_numbered = (numId is not None)
            
            out.append({
                "text": txt,
                "is_numbered": is_numbered,
                "ilvl": ilvl,
                "num_fmt": fmt,
                "num_id": numId # Ukládáme si i ID seznamu pro detekci změny kontextu
            })
        
        return out

    def _parse_questions_from_paragraphs(self, paragraphs: List[dict]) -> List[Question]:
        out: List[Question] = []
        n = len(paragraphs)
        i = 0

        # Regexy
        rx_bonus_start = re.compile(r'^\s*Otázka\s+\d+.*BONUS', re.IGNORECASE)

        # Podpora "1. Text ..." / "1) Text ..." (ručně psané číslování v textu)
        rx_plain_numbered_question = re.compile(r'^\s*\d{1,3}\s*[\.)]\s+\S')

        # Přísnější definice otázky: Musí začínat velkým písmenem
        rx_looks_like_question_start = re.compile(r'^[A-ZŽŠČŘĎŤŇ]')

        rx_visual_bullet = re.compile(r'^\s*[\-–—•]')
        rx_not_question = re.compile(r'^(Slovník|Tabulka|Obrázek|Příklad|Body|Poznámka)', re.IGNORECASE)

        def html_escape(s: str) -> str:
            return _html.escape(s or "")

        def strip_manual_numbering(s: str) -> str:
            # Odstraní pouze ručně napsané "1." / "1)" na začátku
            return re.sub(r'^\s*\d{1,3}\s*[\.)]\s+', '', (s or ""))

        # --- JÁDRO LOGIKY ---
        def is_start_of_new_question(p: dict) -> bool:
            text = p["text"].strip()
            if not text:
                return False

            # 0. Ručně psané číslování "1." / "1)" -> nová otázka
            if rx_plain_numbered_question.match(text):
                return True

            # 1. Vizuální odrážka (pomlčka atd.) -> NIKDY není nová otázka
            if rx_visual_bullet.match(text):
                return False

            # 2. Začíná malým písmenem -> NIKDY není nová otázka
            if text[0].islower():
                return False

            # 3. Explicitní "Otázka X" -> VŽDY je nová otázka
            if "Otázka" in text and re.search(r'Otázka\s+\d', text):
                return True

            # 4. Word číslovaný seznam (level 0)
            if p["is_numbered"] and p["ilvl"] == 0 and p["num_fmt"] != "bullet":
                if rx_looks_like_question_start.match(text) and not rx_not_question.match(text):
                    return True

            return False

        while i < n:
            p = paragraphs[i]
            txt = p["text"].strip()

            # Hledáme start
            if not is_start_of_new_question(p):
                i += 1
                continue

            # -- START OTÁZKY --
            is_bonus = bool(rx_bonus_start.search(txt))
            q_type = "bonus" if is_bonus else "classic"

            # Pokud je číslování ručně psané ("1."/"1)"), odstraň ho z titulku i z první věty v HTML
            first_line_txt = strip_manual_numbering(txt)

            clean_title = first_line_txt
            if len(clean_title) > 60:
                clean_title = clean_title[:57] + "..."

            html_parts = [f"<p>{html_escape(first_line_txt)}</p>"]

            # -- ČTENÍ OBSAHU (Hladový režim, ale s respektem k nové otázce) --
            j = i + 1
            while j < n:
                next_p = paragraphs[j]
                next_txt = next_p["text"].strip()

                # Pokud narazíme na něco, co splňuje definici NOVÉ otázky, končíme.
                if is_start_of_new_question(next_p):
                    break

                # Jinak je to součást této otázky
                if next_txt:
                    html_parts.append(f"<p>{html_escape(next_txt)}</p>")

                j += 1

            full_html = "".join(html_parts)

            q = Question.new_default(q_type)
            q.title = ("BONUS: " + clean_title) if is_bonus else clean_title
            q.text_html = full_html
            q.points = (0 if is_bonus else 1)
            q.bonus_correct = (1.0 if is_bonus else 0.0)

            out.append(q)
            i = j

        return out

    def _ensure_unassigned_group(self) -> tuple[str, Optional[str]]:
        """Zajistí existenci skupiny 'Neroztříděné'. Vrací (group_id, None)."""
        name = "Neroztříděné"
        g = next((g for g in self.root.groups if g.name == name), None)
        if not g:
            g = Group(id=str(_uuid.uuid4()), name=name, subgroups=[])
            self.root.groups.append(g)
        # Nevytváříme "Default" podskupinu automaticky, pokud není potřeba.
        # V importu si vytvoříme "Klasické" a "Bonusové" specificky.
        return g.id, None


    def _import_from_docx(self) -> None:
        # Výchozí složka pro import
        import_dir = self.project_root / "data" / "Staré písemky"
        import_dir.mkdir(parents=True, exist_ok=True)

        paths, _ = QFileDialog.getOpenFileNames(self, "Import z DOCX", str(import_dir), "Word dokument (*.docx)")
        if not paths:
            return

        # 1. Výběr cílové skupiny/podskupiny pomocí dialogu (NOVÉ)
        dlg = MoveTargetDialog(self)
        dlg.setWindowTitle("Importovat do...")
        if dlg.exec() != QDialog.Accepted:
            return
        
        target_gid, target_sgid = dlg.selected_target()
        if not target_gid:
             return

        # Získání reference na cílovou podskupinu
        target_sg = None
        if target_sgid:
            target_sg = self._find_subgroup(target_gid, target_sgid)
        else:
            # Pokud vybral jen skupinu, zkusíme najít/vytvořit první podskupinu
            g = self._find_group(target_gid)
            if g:
                if g.subgroups:
                    target_sg = g.subgroups[0]
                else:
                    # Vytvoření defaultní podskupiny, pokud skupina žádnou nemá
                    new_sg = Subgroup(id=str(_uuid.uuid4()), name="Default", subgroups=[], questions=[])
                    g.subgroups.append(new_sg)
                    target_sg = new_sg

        if not target_sg:
            QMessageBox.warning(self, "Chyba", "Nepodařilo se určit cílovou podskupinu.")
            return

        # 2. Vytvoření indexu existujících otázek pro kontrolu duplicit
        #    Jako klíč použijeme ostripovaný HTML obsah.
        existing_hashes = set()

        def index_questions(node):
            # Rekurzivně projít strom
            if isinstance(node, RootData):
                for gr in node.groups: index_questions(gr)
            elif isinstance(node, Group):
                for sgr in node.subgroups: index_questions(sgr)
            elif isinstance(node, Subgroup):
                for q in node.questions:
                    if q.text_html:
                        existing_hashes.add(q.text_html.strip())
                for sub in node.subgroups:
                    index_questions(sub)

        index_questions(self.root)

        total_imported = 0
        total_duplicates = 0

        for p in paths:
            try:
                paras = self._extract_paragraphs_from_docx(Path(p))
                qs = self._parse_questions_from_paragraphs(paras)
                
                if not qs:
                    # Info, ale nepovažujeme za chybu přerušující ostatní soubory
                    continue

                file_imported_count = 0
                
                for q in qs:
                    content_hash = (q.text_html or "").strip()
                    
                    # Kontrola duplicit
                    if content_hash in existing_hashes:
                        total_duplicates += 1
                        continue
                    
                    # Pokud není duplicitní, přidáme do DB a aktualizujeme hashset
                    existing_hashes.add(content_hash)
                    
                    # Vložíme přímo do vybrané cílové podskupiny (nerozlišujeme složky Klasické/Bonusové)
                    target_sg.questions.append(q)
                    
                    file_imported_count += 1

                total_imported += file_imported_count

            except Exception as e:
                QMessageBox.warning(self, "Import – chyba", f"Soubor: {p}\n{e}")

        self._refresh_tree()
        self.save_data()

        msg = f"Import dokončen do: {target_sg.name}\n\nÚspěšně importováno: {total_imported}\nDuplicitních (přeskočeno): {total_duplicates}"
        QMessageBox.information(self, "Výsledek importu", msg)



    # -------------------- Přesun otázky --------------------

    def _move_question(self) -> None:
        kind, meta = self._selected_node()
        if kind != "question":
            QMessageBox.information(self, "Přesun", "Vyberte nejprve otázku ve stromu."); return
        dlg = MoveTargetDialog(self)
        if dlg.exec() != QDialog.Accepted: return
        g_id, sg_id = dlg.selected_target()
        if not g_id: return
        g = self._find_group(g_id); 
        if not g: return
        target_sg = self._find_subgroup(g_id, sg_id) if sg_id else None
        if not target_sg:
            if not g.subgroups:
                g.subgroups.append(Subgroup(id=str(_uuid.uuid4()), name="Default", subgroups=[], questions=[]))
            target_sg = g.subgroups[0]
        src_gid = meta["parent_group_id"]; src_sgid = meta["parent_subgroup_id"]; qid = meta["id"]
        src_sg = self._find_subgroup(src_gid, src_sgid); q = self._find_question(src_gid, src_sgid, qid)
        if not (src_sg and q): return
        src_sg.questions = [qq for qq in src_sg.questions if qq.id != qid]
        target_sg.questions.append(q)
        self._refresh_tree(); self.save_data()
        g_name = g.name if g else ""; sg_name = target_sg.name if target_sg else "Default"
        self.statusBar().showMessage(f"Otázka přesunuta do {g_name} / {sg_name}.", 4000)

    def _bulk_move_selected(self) -> None:
        items = [it for it in self.tree.selectedItems() if (it.data(0, Qt.UserRole) or {}).get('kind') == 'question']
        if not items:
            QMessageBox.information(self, "Přesun", "Vyberte ve stromu alespoň jednu otázku."); return
        dlg = MoveTargetDialog(self)
        if dlg.exec() != QDialog.Accepted: return
        g_id, sg_id = dlg.selected_target()
        if not g_id: return
        g = self._find_group(g_id); 
        if not g: return
        target_sg = self._find_subgroup(g_id, sg_id) if sg_id else None
        if not target_sg:
            if not g.subgroups:
                g.subgroups.append(Subgroup(id=str(_uuid.uuid4()), name="Default", subgroups=[], questions=[]))
            target_sg = g.subgroups[0]
        moved = 0
        for it in items:
            meta = it.data(0, Qt.UserRole) or {}
            qid = meta.get('id')
            sg = self._find_subgroup(meta.get('parent_group_id'), meta.get('parent_subgroup_id'))
            q = self._find_question(meta.get('parent_group_id'), meta.get('parent_subgroup_id'), qid)
            if sg and q:
                sg.questions = [qq for qq in sg.questions if qq.id != qid]
                target_sg.questions.append(q); moved += 1
        self._refresh_tree(); self.save_data()
        g_name = g.name if g else ""; sg_name = target_sg.name if target_sg else "Default"
        self.statusBar().showMessage(f"Přesunuto {moved} otázek do {g_name} / {sg_name}.", 4000)

    # -------------------- Export DOCX --------------------

    def _all_questions_by_type(self, qtype: str) -> List[Question]:
        out: List[Question] = []
        def walk_subs(lst: List[Subgroup]):
            for sg in lst:
                for q in sg.questions:
                    if q.type == qtype: out.append(q)
                walk_subs(sg.subgroups)
        for g in self.root.groups:
            walk_subs(g.subgroups)
        return out

    def _export_docx_wizard(self):
        # PŮVODNÍ (ŠPATNĚ):
        # wiz = ExportWizard(self)
        # wiz.le_output.setText(str(self.project_root / "test_vystup.docx")) <--- TOTO SMAZAT!
        # wiz.exec()

        # NOVÉ (SPRÁVNĚ):
        wiz = ExportWizard(self)
        wiz.exec()

    def _convert_docx_to_pdf(self, docx_path: Path) -> Optional[Path]:
        """
        Konvertuje DOCX na PDF pomocí LibreOffice (hledá spustitelný soubor i v /Applications).
        Vrací cestu k PDF souboru nebo None pokud selhalo.
        """
        import shutil
        
        # 1. Hledání spustitelného souboru LibreOffice
        lo_candidates = [
            "libreoffice",                                            # Standard Linux/PATH
            "soffice",                                                # Generic bin
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",   # macOS standard path
            "/usr/bin/libreoffice",
            "/usr/local/bin/libreoffice",
            r"C:\Program Files\LibreOffice\program\soffice.exe",      # Windows x64
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe" # Windows x86
        ]
        
        lo_executable = None
        for cand in lo_candidates:
            # shutil.which hledá v PATH, Path(cand).exists() hledá konkrétní soubor
            if shutil.which(cand) or Path(cand).exists():
                lo_executable = cand
                break
        
        if not lo_executable:
             QMessageBox.warning(
                self, 
                "LibreOffice nenalezen",
                "Nemohu najít nainstalovaný LibreOffice.\n"
                "Pokud jej máte nainstalovaný, ujistěte se, že je ve standardní složce "
                "(/Applications/LibreOffice.app na macOS)."
            )
             return None

        # 2. Samotná konverze
        try:
            pdf_path = docx_path.with_suffix('.pdf')
            
            cmd = [
                lo_executable,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(pdf_path.parent),
                str(docx_path)
            ]
            
            # Spuštění procesu
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                print(f"LibreOffice chyba (Code {result.returncode}):\nSTDERR: {result.stderr}\nSTDOUT: {result.stdout}")
            
            # Ověření, že PDF byl vytvořen
            if pdf_path.exists():
                return pdf_path
            else:
                # Někdy se stane, že returncode je 0, ale soubor nikde (např. sandbox issues)
                return None
                
        except subprocess.TimeoutExpired:
            QMessageBox.warning(self, "Chyba konverze", "Konverze trvala příliš dlouho (timeout).")
            return None
        except Exception as e:
            QMessageBox.warning(self, "Chyba konverze", f"Neočekávaná chyba při konverzi PDF:\n{e}")
            return None

    def _merge_pdfs(self, pdf_paths: List[Path], output_path: Path, cleanup: bool = True) -> bool:
        """
        Spojí více PDF souborů do jednoho.
        Zkouší: PyPDF2 -> macOS join.py -> Ghostscript.
        Pokud vše selže, vyhodí chybovou hlášku s instrukcemi.
        """
        success = False
        missing_tools = []
        
        # 1. Zkusíme PyPDF2 (Preferované)
        try:
            from PyPDF2 import PdfMerger
            merger = PdfMerger()
            for pdf_path in pdf_paths:
                if pdf_path.exists():
                    merger.append(str(pdf_path))
            merger.write(str(output_path))
            merger.close()
            success = True
        except ImportError:
            missing_tools.append("PyPDF2 (pip install PyPDF2)")
            # Pokračujeme na fallback
        except Exception as e:
            print(f"Chyba PyPDF2: {e}")
            # Pokračujeme na fallback

        if not success:
            # 2. Fallback: macOS Native Script nebo Ghostscript
            success = self._merge_pdfs_subprocess(pdf_paths, output_path)
            if not success:
                missing_tools.append("Ghostscript (brew install ghostscript)")

        # Pokud selhalo úplně
        if not success:
            msg = (
                "Nepodařilo se sloučit PDF soubory.\n"
                "Chybí potřebné nástroje.\n\n"
                "Řešení (vyberte jedno):\n"
                "1. Nainstalujte python knihovnu:  pip install PyPDF2\n"
                "2. Nainstalujte Ghostscript:      brew install ghostscript\n\n"
                "Jednotlivé PDF soubory byly ponechány ve složce."
            )
            QMessageBox.warning(self, "Chyba slučování PDF", msg)
            # Vracíme False a NEPROVÁDÍME cleanup, aby uživateli zůstaly aspoň jednotlivé soubory
            return False

        # Vyčištění (pouze při úspěchu)
        if success and cleanup:
            for pdf_path in pdf_paths:
                try:
                    pdf_path.unlink()
                except Exception as e:
                    print(f"Nemohu smazat dočasný PDF {pdf_path}: {e}")
        
        return True


    def _merge_pdfs_subprocess(self, pdf_paths: List[Path], output_path: Path) -> bool:
        """Fallback: slučování PDF pomocí externích nástrojů (macOS join.py nebo GS)."""
        import sys
        
        # A. macOS Built-in Script (Automator)
        # Tento skript je standardně přítomen na macOS
        macos_join_script = "/System/Library/Automator/Combine PDF Pages.action/Contents/Resources/join.py"
        if sys.platform == "darwin" and Path(macos_join_script).exists():
            try:
                cmd = [
                    "python3",  # Použijeme systémový python nebo ten co běží
                    macos_join_script,
                    "-o", str(output_path)
                ]
                cmd.extend([str(p) for p in pdf_paths if p.exists()])
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                if result.returncode == 0:
                    return True
                else:
                    print(f"macOS join.py failed: {result.stderr}")
            except Exception as e:
                print(f"macOS join.py exception: {e}")

        # B. Ghostscript
        try:
            cmd = ['gs', '-q', '-dNOPAUSE', '-dBATCH', '-dSAFER', '-sDEVICE=pdfwrite', f'-sOutputFile={output_path}']
            cmd.extend([str(p) for p in pdf_paths if p.exists()])
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            return result.returncode == 0
            
        except FileNotFoundError:
            # GS není nainstalován
            print("Ghostscript (gs) nebyl nalezen.")
            return False
        except Exception as e:
            print(f"Ghostscript error: {e}")
            return False

    def _generate_docx_from_template(self, template_path: Path, output_path: Path,
                                     simple_repl: Dict[str, str], rich_repl_html: Dict[str, object]) -> None:
        """
        Generuje DOCX. rich_repl_html může být:
         - Dict[str, str] -> {placeholder: html_content}
         - Dict[str, tuple] -> {placeholder: (html_content, image_path[, image_width_cm, image_height_cm])}
        """
        import docx
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        import os
        
        for ph, val in rich_repl_html.items():
            if isinstance(val, tuple):
                html_part, img_path = val[0], val[1]
                if img_path:
                    img_check = Path(img_path).exists() if isinstance(img_path, str) else False

        # Parsing HTML
        def parse_html_to_paragraphs(html):
            # Zde voláme vaši existující třídu HTMLToDocxParser
            parser = HTMLToDocxParser() 
            parser.feed(html)
            return parser.paragraphs

        try:
            doc = docx.Document(template_path)
        except Exception as e:
            QMessageBox.critical(self, "Export chyba", f"Nelze otevřít šablonu pomocí python-docx:\n{e}")
            print(f"[ERROR] Nelze otevřít DOCX: {e}")
            return

        # -- HELPER: Extrakce a obnova Page Breaks --
        def extract_page_breaks(paragraph):
            breaks = []
            p_elem = paragraph._p
            for run_elem in p_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                for br_elem in run_elem.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
                    if br_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                        breaks.append(br_elem)
            return breaks
        
        def restore_page_breaks(paragraph, breaks):
            if not breaks:
                return
            p_elem = paragraph._p
            new_run = OxmlElement('w:r')
            for br in breaks:
                br_copy = br.__copy__()
                new_run.append(br_copy)
            p_elem.append(new_run)

        # -- Helper: Vložení Rich Text bloku + Obrázku --
        def insert_rich_question_block(paragraph, html_content, image_path=None, image_w_cm: float = 0.0, image_h_cm: float = 0.0):
            
            # 1. Parse HTML
            paras_data = parse_html_to_paragraphs(html_content)
            
            # Pokud je obsah prázdný a není ani obrázek -> vyčistit a konec
            if not paras_data and not image_path: 
                breaks = extract_page_breaks(paragraph)
                paragraph.clear()
                restore_page_breaks(paragraph, breaks)
                return
            
            p_insert = paragraph._p
            first_p_used = False

            # --- VLOŽENÍ TEXTU ---
            if paras_data:
                for i, p_data in enumerate(paras_data):
                    if i == 0:
                        new_p = paragraph
                        breaks = extract_page_breaks(new_p)
                        new_p.clear()
                        first_p_used = True
                    else:
                        new_p = doc.add_paragraph()
                        p_insert.addnext(new_p._p)
                        p_insert = new_p._p

                    # Zarovnání
                    align = p_data.get('align', 'left')
                    if align == 'center': new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align == 'right': new_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif align == 'justify': new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else: new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Mezerování
                    new_p.paragraph_format.space_before = Pt(0)
                    new_p.paragraph_format.space_after = Pt(0)

                    # Prefix (odrážky)
                    if p_data.get('prefix'):
                        new_p.paragraph_format.left_indent = Pt(48)
                        new_p.paragraph_format.first_line_indent = Pt(-24)
                        new_p.add_run(p_data['prefix'])
                    
                    # Obsah (Runs)
                    for r_data in p_data['runs']:
                        text_content = r_data['text']
                        parts = text_content.split('\n')
                        for idx, part in enumerate(parts):
                            if part:
                                run = new_p.add_run(part)
                                if r_data.get('b'): run.bold = True
                                if r_data.get('i'): run.italic = True
                                if r_data.get('u'): run.underline = True
                                if r_data.get('color'):
                                    try:
                                        rgb = r_data['color']
                                        run.font.color.rgb = RGBColor(int(rgb[:2], 16), int(rgb[2:4], 16), int(rgb[4:], 16))
                                    except: pass
                            
                            if idx < len(parts) - 1:
                                run = new_p.add_run()
                                run.add_break()
            
            # --- VLOŽENÍ OBRÁZKU ---
            if image_path:
                import shutil
                import subprocess
                import tempfile
                
                img_path_obj = Path(image_path) if isinstance(image_path, str) else image_path
                final_img_path = img_path_obj

                if not img_path_obj.exists():
                    print(f"[ERROR]   Obrázek NEEXISTUJE! {img_path_obj}")
                    return

                # --- KONVERZE HEIC -> JPG (pro macOS) ---
                is_heic = img_path_obj.suffix.lower() in ('.heic', '.heif')
                temp_jpg = None
                
                if is_heic:
                    try:
                        # Vytvoříme dočasný soubor
                        fd, temp_jpg = tempfile.mkstemp(suffix=".jpg")
                        os.close(fd)
                        
                        # Použijeme systémový nástroj sips (macOS) nebo ImageMagick
                        # Zkusíme sips (je na každém macu)
                        cmd = ["sips", "-s", "format", "jpeg", str(img_path_obj), "--out", temp_jpg]
                        subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                        
                        final_img_path = Path(temp_jpg)
                    except Exception as e:
                        print(f"[ERROR]   Chyba konverze HEIC: {e}")
                        # Fallback - zkusíme vložit originál, i když to asi selže
                        final_img_path = img_path_obj

                # Vytvoříme nový odstavec pro obrázek
                img_p = doc.add_paragraph()
                if not paras_data:
                    paragraph.clear()
                    img_p = paragraph
                else:
                    p_insert.addnext(img_p._p)
                    p_insert = img_p._p

                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(6)
                img_p.paragraph_format.space_after = Pt(6)
                
                try:
                    run = img_p.add_run()
                    # Pokud je 0, zachováme původní default (14 cm). Pokud je výška 0, necháme Word dopočítat poměr.
                    w_use = float(image_w_cm or 0.0)
                    h_use = float(image_h_cm or 0.0)
                    if w_use <= 0.0 and h_use <= 0.0:
                        w_use = 14.0

                    if w_use > 0.0 and h_use > 0.0:
                        run.add_picture(str(final_img_path), width=Cm(w_use), height=Cm(h_use))
                    elif w_use > 0.0:
                        run.add_picture(str(final_img_path), width=Cm(w_use))
                    else:
                        run.add_picture(str(final_img_path), height=Cm(h_use)) 
                except Exception as e:
                    print(f"[ERROR] Chyba při vkládání obrázku: {e}")
                finally:
                    # Úklid dočasného souboru
                    if temp_jpg and os.path.exists(temp_jpg):
                        try:
                            os.remove(temp_jpg)
                        except: pass

            # --- OBNOVA PAGE BREAKS ---
            if 'breaks' in locals() and breaks:
                break_p = doc.add_paragraph()
                p_insert.addnext(break_p._p)
                restore_page_breaks(break_p, breaks)


        # -- Helper: Zpracování jednoho odstavce (Inline i Block) --
        def process_paragraph(p):
            full_text = p.text
            if not ("<" in full_text or "{" in full_text):
                return

            # 1. BLOCK CHECK
            txt_clean = full_text.strip()
            matched_rich = None
            
            for ph, val in rich_repl_html.items():
                if isinstance(val, tuple):
                    html_content, img_path = val[0], val[1]
                else:
                    html_content, img_path = val, None

                if txt_clean == f"<{ph}>" or txt_clean == f"{{{ph}}}":
                    matched_rich = (html_content, img_path, float((val[2] if len(val) > 2 else 0.0) or 0.0), float((val[3] if len(val) > 3 else 0.0) or 0.0))
                    break
            
            if matched_rich:
                insert_rich_question_block(p, matched_rich[0], matched_rich[1], matched_rich[2], matched_rich[3])
                return

            # 2. INLINE CHECK (zkráceno, bez obrázků)
            keys_found = []
            for k in simple_repl.keys():
                if f"<{k}>" in full_text or f"{{{k}}}" in full_text:
                    keys_found.append(k)
            for k in rich_repl_html.keys():
                if f"<{k}>" in full_text or f"{{{k}}}" in full_text:
                    keys_found.append(k)
            
            if not keys_found:
                return

            # Zbytek inline logiky (zkráceno pro délku)...
            segments = [full_text]
            all_repl_data = {}
            
            for k, v in simple_repl.items(): 
                all_repl_data[k] = {'type': 'simple', 'val': v}
            for k, val in rich_repl_html.items(): 
                html = val[0] if isinstance(val, tuple) else val
                all_repl_data[k] = {'type': 'rich', 'val': html}
            
            for k in keys_found:
                info = all_repl_data[k]
                tokens = [f"<{k}>", f"{{{k}}}"]
                for token in tokens:
                    new_segments = []
                    for seg in segments:
                        if isinstance(seg, str):
                            parts = seg.split(token)
                            for i, part in enumerate(parts):
                                if part: new_segments.append(part)
                                if i < len(parts) - 1: new_segments.append(info)
                        else:
                            new_segments.append(seg)
                    segments = new_segments

            base_font_name = None; base_font_size = None; base_bold = None
            if p.runs:
                r0 = p.runs[0]
                base_font_name = r0.font.name
                base_font_size = r0.font.size
                base_bold = r0.bold

            page_breaks = extract_page_breaks(p)
            p.clear()
            
            for seg in segments:
                if isinstance(seg, str):
                    run = p.add_run(seg)
                    if base_font_name: run.font.name = base_font_name
                    if base_font_size: run.font.size = base_font_size
                elif isinstance(seg, dict):
                    val = seg['val']
                    if seg['type'] == 'simple':
                        run = p.add_run(str(val))
                        if base_font_name: run.font.name = base_font_name
                        if base_font_size: run.font.size = base_font_size
                        if base_bold is not None: run.bold = base_bold
                    elif seg['type'] == 'rich':
                        paras = parse_html_to_paragraphs(val)
                        for p_idx, p_data in enumerate(paras):
                            if p_idx > 0: p.add_run().add_break()
                            for r_data in p_data['runs']:
                                text_content = r_data['text']
                                parts = text_content.split('\n')
                                for idx_part, part in enumerate(parts):
                                    if part:
                                        run = p.add_run(part)
                                        if r_data.get('b'): run.bold = True
                                        if r_data.get('i'): run.italic = True
                                        if r_data.get('u'): run.underline = True
                                        if r_data.get('color'):
                                             try:
                                                rgb = r_data['color']
                                                run.font.color.rgb = RGBColor(int(rgb[:2], 16), int(rgb[2:4], 16), int(rgb[4:], 16))
                                             except: pass
                                        if base_font_name: run.font.name = base_font_name
                                        if base_font_size: run.font.size = base_font_size
                                    if idx_part < len(parts) - 1:
                                        p.add_run().add_break()
            
            restore_page_breaks(p, page_breaks)


        # --- HLAVNÍ SMYČKA ---
        
        # 1. Body
        for i, p in enumerate(doc.paragraphs):
            
            process_paragraph(p)
            
        # 2. Tables in Body
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        process_paragraph(p)
        
        # 3. Headers / Footers
        for section in doc.sections:
            for h in [section.header, section.first_page_header]:
                if h:
                    for p in h.paragraphs: process_paragraph(p)
            for f in [section.footer, section.first_page_footer]:
                if f:
                    for p in f.paragraphs: process_paragraph(p)

        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
        except Exception as e:
            print(f"[ERROR] Chyba uložení: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Chyba uložení", f"Nelze uložit DOCX:\n{e}")
            
    def default_root_obj(self) -> RootData:
        return RootData(groups=[], trash=[])
    
    def _ensure_restored_subgroup(self) -> Subgroup:
        """Zajistí existenci cíle pro obnovené otázky a vrátí tuto podskupinu."""
        restored_group_name = "Obnovené ze smazaných"
        restored_subgroup_name = "Obnovené otázky"
    
        g = None
        for gg in self.root.groups:
            if gg.name == restored_group_name:
                g = gg
                break
        if g is None:
            g = Group(id=str(_uuid.uuid4()), name=restored_group_name, subgroups=[])
            self.root.groups.append(g)
    
        sg = None
        for s in g.subgroups:
            if s.name == restored_subgroup_name:
                sg = s
                break
        if sg is None:
            sg = Subgroup(id=str(_uuid.uuid4()), name=restored_subgroup_name, subgroups=[], questions=[])
            g.subgroups.append(sg)
    
        return sg
    
    def _trash_delete_selected(self) -> None:
        if not hasattr(self, "table_trash"):
            return
    
        sel = self.table_trash.selectionModel().selectedRows() if self.table_trash.selectionModel() else []
        if not sel:
            QMessageBox.information(self, "Koš", "Vyberte otázky pro trvalé smazání.")
            return
    
        if QMessageBox.question(self, "Koš", "Trvale smazat vybrané otázky z koše?") != QMessageBox.Yes:
            return
    
        trash_list = getattr(self.root, "trash", [])
        if not isinstance(trash_list, list):
            return
    
        qids: set[str] = set()
        for r in sel:
            it = self.table_trash.item(r.row(), 0)
            if not it:
                continue
            qid = it.data(Qt.UserRole)
            if isinstance(qid, str) and qid:
                qids.add(qid)
    
        if not qids:
            return
    
        new_trash: List[dict] = []
        for rec in trash_list:
            if not isinstance(rec, dict):
                continue
            qd = rec.get("question", {})
            if isinstance(qd, dict) and qd.get("id", "") in qids:
                continue
            new_trash.append(rec)
    
        self.root.trash = new_trash
        self._refresh_trash_table()
        self.save_data()
        self.statusBar().showMessage("Vybrané položky byly trvale smazány z koše.", 3000)
            
    def _trash_empty(self) -> None:
        trash_list = getattr(self.root, "trash", [])
        if not isinstance(trash_list, list) or not trash_list:
            return
    
        if QMessageBox.question(self, "Koš", "Vysypat koš? (Trvale smaže všechny otázky v koši)") != QMessageBox.Yes:
            return
    
        self.root.trash = []
        self._refresh_trash_table()
        self.save_data()
        self.statusBar().showMessage("Koš vysypán.", 2500)
    
    def _trash_restore_selected(self) -> None:
        if not hasattr(self, "table_trash"):
            return
    
        sel = self.table_trash.selectionModel().selectedRows() if self.table_trash.selectionModel() else []
        if not sel:
            QMessageBox.information(self, "Koš", "Vyberte otázky pro obnovení.")
            return
    
        trash_list = getattr(self.root, "trash", [])
        if not isinstance(trash_list, list):
            QMessageBox.warning(self, "Koš", "Koš není dostupný v datech.")
            return
    
        qids: List[str] = []
        for r in sel:
            it = self.table_trash.item(r.row(), 0)
            if not it:
                continue
            qid = it.data(Qt.UserRole)
            if isinstance(qid, str) and qid:
                qids.append(qid)
    
        if not qids:
            return
    
        remaining: List[dict] = []
        restored_count = 0
    
        for rec in trash_list:
            if not isinstance(rec, dict):
                continue
            qd = rec.get("question", {})
            if not isinstance(qd, dict):
                remaining.append(rec)
                continue
    
            qid = qd.get("id", "")
            if qid not in qids:
                remaining.append(rec)
                continue
    
            gid = rec.get("source_group_id", "") or ""
            sgid = rec.get("source_subgroup_id", "") or ""
    
            target_sg = None
    
            # obnovit do původního místa pokud existuje (spolehne se na tvé existující _find_group/_find_subgroup)
            if gid and sgid and hasattr(self, "_find_group") and hasattr(self, "_find_subgroup"):
                try:
                    g = self._find_group(gid)
                    sg = self._find_subgroup(gid, sgid)
                    if g is not None and sg is not None:
                        target_sg = sg
                except Exception:
                    target_sg = None
    
            if target_sg is None:
                _, target_sg = self._ensure_restored_targets()
    
            if not hasattr(self, "_parse_question"):
                QMessageBox.warning(self, "Koš", "Chybí _parse_question – nelze obnovit.")
                remaining.append(rec)
                continue
    
            q_obj = self._parse_question(qd)
            target_sg.questions.append(q_obj)
            restored_count += 1
    
        self.root.trash = remaining
        self._refresh_tree()
        self._refresh_trash_table()
        self.save_data()
        self.statusBar().showMessage(f"Obnoveno {restored_count} otázek.", 3000)
    
    def _on_trash_selection_changed(self) -> None:
        if not hasattr(self, "table_trash"):
            return
    
        sel = self.table_trash.selectionModel().selectedRows() if self.table_trash.selectionModel() else []
        has_sel = bool(sel)
    
        trash_list = getattr(self.root, "trash", [])
        has_any = isinstance(trash_list, list) and len(trash_list) > 0
    
        if hasattr(self, "btn_trash_restore"):
            self.btn_trash_restore.setEnabled(has_sel)
        if hasattr(self, "btn_trash_delete"):
            self.btn_trash_delete.setEnabled(has_sel)
        if hasattr(self, "btn_trash_empty"):
            self.btn_trash_empty.setEnabled(has_any)
    
        if not hasattr(self, "trash_detail"):
            return
    
        if not has_sel:
            self.trash_detail.setPlainText("")
            return
    
        row = sel[0].row()
        it = self.table_trash.item(row, 0)
        qid = it.data(Qt.UserRole) if it else ""
        if not isinstance(qid, str) or not qid:
            self.trash_detail.setPlainText("")
            return
    
        rec = None
        for r in trash_list if isinstance(trash_list, list) else []:
            if not isinstance(r, dict):
                continue
            qd = r.get("question", {})
            if isinstance(qd, dict) and qd.get("id", "") == qid:
                rec = r
                break
    
        if not rec:
            self.trash_detail.setPlainText("")
            return
    
        qd = rec.get("question", {})
        if not isinstance(qd, dict):
            qd = {}
    
        title = qd.get("title", "")
        qtype = qd.get("type", "classic")
        deleted_at = rec.get("deleted_at", "")
        gname = rec.get("source_group_name", "")
        sgname = rec.get("source_subgroup_name", "")
    
        text_html = qd.get("text_html", "") or ""
        correct_answer = qd.get("correct_answer", "") or ""
    
        txt = []
        txt.append(f"Název: {title}")
        txt.append(f"Typ: {'BONUS' if qtype == 'bonus' else 'Klasická'}")
        txt.append(f"Smazáno: {deleted_at}")
        txt.append(f"Původní skupina: {gname}")
        txt.append(f"Původní podskupina: {sgname}")
        txt.append("")
        txt.append("Obsah (HTML uložené):")
        txt.append(text_html)
        if correct_answer:
            txt.append("")
            txt.append("Správná odpověď:")
            txt.append(correct_answer)
    
        self.trash_detail.setPlainText("\n".join(txt))
        
    def _ensure_restored_targets(self) -> Tuple["Group", "Subgroup"]:
        restored_group_name = "Obnovené ze smazaných"
        restored_subgroup_name = "Obnovené otázky"
    
        g = None
        for gg in self.root.groups:
            if gg.name == restored_group_name:
                g = gg
                break
        if g is None:
            g = Group(id=str(_uuid.uuid4()), name=restored_group_name, subgroups=[])
            self.root.groups.append(g)
    
        sg = None
        for s in g.subgroups:
            if s.name == restored_subgroup_name:
                sg = s
                break
        if sg is None:
            sg = Subgroup(id=str(_uuid.uuid4()), name=restored_subgroup_name, subgroups=[], questions=[])
            g.subgroups.append(sg)
    
        return g, sg

    def _refresh_trash_table(self) -> None:
        if not hasattr(self, "table_trash"):
            return
    
        trash_list = getattr(self.root, "trash", [])
        if not isinstance(trash_list, list):
            trash_list = []
    
        def sort_key(r: dict) -> str:
            if isinstance(r, dict):
                return str(r.get("deleted_at", ""))
            return ""
    
        rows = sorted(trash_list, key=sort_key, reverse=True)
    
        self.table_trash.setSortingEnabled(False)
        self.table_trash.setRowCount(0)
    
        for rec in rows:
            if not isinstance(rec, dict):
                continue
            qd = rec.get("question", {})
            if not isinstance(qd, dict):
                qd = {}
    
            qid = qd.get("id", "")
            title = qd.get("title") or "(bez názvu)"
            qtype = qd.get("type", "classic")
            type_txt = "BONUS" if qtype == "bonus" else "Klasická"
            deleted_at = rec.get("deleted_at", "")
            gname = rec.get("source_group_name", "")
            sgname = rec.get("source_subgroup_name", "")
    
            row = self.table_trash.rowCount()
            self.table_trash.insertRow(row)
    
            it_title = QTableWidgetItem(title)
            it_title.setData(Qt.UserRole, qid)
    
            it_type = QTableWidgetItem(type_txt)
            it_deleted = QTableWidgetItem(deleted_at)
            it_g = QTableWidgetItem(gname)
            it_sg = QTableWidgetItem(sgname)
    
            self.table_trash.setItem(row, 0, it_title)
            self.table_trash.setItem(row, 1, it_type)
            self.table_trash.setItem(row, 2, it_deleted)
            self.table_trash.setItem(row, 3, it_g)
            self.table_trash.setItem(row, 4, it_sg)
    
        self.table_trash.setSortingEnabled(True)
        self._on_trash_selection_changed()

    # -------------------- Pomocné --------------------

    def _derive_title_from_html(self, html: str, prefix: str = "") -> str:
        import re as _re, html as _h
        txt = _re.sub(r'<[^>]+>', ' ', html or '')
        txt = _h.unescape(txt).strip()
        if not txt: return (prefix + "Otázka").strip()
        parts = _re.split(r'[.!?]\s', txt)
        head = parts[0] if parts and parts[0] else txt
        head = head.strip()
        if len(head) > 80: head = head[:77].rstrip() + '…'
        return (prefix + head).strip()


# --------------------------- main ---------------------------

def main() -> int:
    app = QApplication(sys.argv)
    apply_dark_theme(app)

    # Nastavení ikony na úrovni aplikace (pokud existuje)
    project_root = Path.cwd()
    icon_file = project_root / "icon" / "icon.png"
    if icon_file.exists():
        app.setWindowIcon(QIcon(str(icon_file)))

    w = MainWindow()
    w.show()
    return app.exec()


if __name__ == "__main__":
    sys.exit(main())
